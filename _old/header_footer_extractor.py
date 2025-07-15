#!/usr/bin/env python3
"""
Header and Footer Extractor Module

This module provides functionality to extract header, footer, margin, and comment information
from Word documents (.docx). It combines functionality from the rip scripts into a reusable module.

Features:
- Extract header content (paragraphs, tables, text boxes)
- Extract footer content (paragraphs, tables, text boxes)
- Extract margin settings
- Extract comments with metadata
- Save data to JSON and TXT formats

Usage:
    from header_footer_extractor import HeaderFooterExtractor
    
    extractor = HeaderFooterExtractor()
    data = extractor.extract_all(docx_path)
"""

from docx import Document
import json
import os
from typing import Dict, List, Any, Optional
from datetime import datetime

class HeaderFooterExtractor:
    """Extracts header, footer, margin, and comment information from Word documents"""
    
    def __init__(self):
        pass
    
    def extract_all(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract all header, footer, margin, and comment information from a document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            Dictionary containing header, footer, margin, and comment data
        """
        try:
            doc = Document(docx_path)
            
            # Extract all components
            header_footer_data = self.extract_header_footer_margins(docx_path)
            comments = self.extract_comments(docx_path)
            
            return {
                "header": header_footer_data["header"],
                "footer": header_footer_data["footer"],
                "margins": header_footer_data["margins"],
                "comments": comments
            }
            
        except Exception as e:
            print(f"Error extracting header/footer data: {e}")
            return {
                "header": {"paragraphs": [], "tables": [], "text_boxes": []},
                "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
                "margins": {},
                "comments": []
            }
    
    def extract_header_footer_margins(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract header, footer, and margin information from a Word document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            Dictionary containing header, footer, and margin data
        """
        try:
            doc = Document(docx_path)
            sec = doc.sections[0]
            
            # Extract margin settings
            margins = self._extract_margins(sec)
            
            # Extract header content
            header_content = self._extract_header_content(sec)
            
            # Extract footer content
            footer_content = self._extract_footer_content(sec)
            
            return {
                "header": header_content,
                "footer": footer_content,
                "margins": margins
            }
            
        except Exception as e:
            print(f"Error extracting header/footer/margins: {e}")
            return {
                "header": {"paragraphs": [], "tables": [], "text_boxes": []},
                "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
                "margins": {}
            }
    
    def _extract_margins(self, section) -> Dict[str, float]:
        """Extract margin settings from a document section"""
        margins = {}
        try:
            if section.top_margin:
                margins["top_margin"] = section.top_margin.inches
            if section.bottom_margin:
                margins["bottom_margin"] = section.bottom_margin.inches
            if section.left_margin:
                margins["left_margin"] = section.left_margin.inches
            if section.right_margin:
                margins["right_margin"] = section.right_margin.inches
            if section.header_distance:
                margins["header_distance"] = section.header_distance.inches
            if section.footer_distance:
                margins["footer_distance"] = section.footer_distance.inches
        except Exception as e:
            print(f"Warning: Could not extract margin settings: {e}")
        
        return margins
    
    def _extract_header_content(self, section) -> Dict[str, List]:
        """Extract header content from a document section"""
        header_content = {"paragraphs": [], "tables": [], "text_boxes": []}
        try:
            if section.header:
                header_element = section.header._element
                header_content = self._extract_content_from_section(header_element, header_element.nsmap)
        except Exception as e:
            print(f"Warning: Could not extract header content: {e}")
        
        return header_content
    
    def _extract_footer_content(self, section) -> Dict[str, List]:
        """Extract footer content from a document section"""
        footer_content = {"paragraphs": [], "tables": [], "text_boxes": []}
        try:
            if section.footer:
                footer_element = section.footer._element
                footer_content = self._extract_content_from_section(footer_element, footer_element.nsmap)
        except Exception as e:
            print(f"Warning: Could not extract footer content: {e}")
        
        return footer_content
    
    def _extract_content_from_section(self, section_element, nsmap) -> Dict[str, List]:
        """
        Extract content from header or footer section
        
        Args:
            section_element: XML element representing the section
            nsmap: Namespace mapping for XML parsing
            
        Returns:
            Dictionary containing paragraphs, tables, and text boxes
        """
        content = {
            "paragraphs": [],
            "tables": [],
            "text_boxes": []
        }
        
        # Extract paragraphs
        for p in section_element.findall('.//w:p', namespaces=nsmap):
            text = self._extract_text_from_element(p, nsmap)
            if text:
                content["paragraphs"].append(text)
        
        # Extract tables
        for tbl in section_element.findall('.//w:tbl', namespaces=nsmap):
            table_data = []
            for row in tbl.findall('.//w:tr', namespaces=nsmap):
                row_data = []
                for cell in row.findall('.//w:tc', namespaces=nsmap):
                    cell_text = self._extract_text_from_element(cell, nsmap)
                    row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            if table_data:
                content["tables"].append(table_data)
        
        # Extract text boxes
        for drawing in section_element.findall('.//w:txbxContent', namespaces=nsmap):
            textbox_data = []
            for p in drawing.findall('.//w:p', namespaces=nsmap):
                text = self._extract_text_from_element(p, nsmap)
                if text:
                    textbox_data.append(text)
            if textbox_data:
                content["text_boxes"].append(textbox_data)
        
        return content
    
    def _extract_text_from_element(self, element, nsmap) -> str:
        """
        Extract all text from an element and its children
        
        Args:
            element: XML element to extract text from
            nsmap: Namespace mapping for XML parsing
            
        Returns:
            Extracted text as string
        """
        texts = []
        for text_elem in element.findall('.//w:t', namespaces=nsmap):
            if text_elem.text:
                texts.append(text_elem.text)
        return ''.join(texts).strip()
    
    def extract_comments(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract comments from a Word document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            List of comment dictionaries with metadata
        """
        try:
            doc = Document(docx_path)
            comments = []
            
            # Check if the document has comments
            if hasattr(doc.part, '_comments_part') and doc.part._comments_part is not None:
                for c in doc.part._comments_part.comments:
                    # Assemble the full comment text
                    full_text = "\n".join(p.text for p in c.paragraphs).strip()
                    
                    comment_data = {
                        "text": full_text,
                        "ref": None
                    }
                    
                    # Get comment metadata
                    try:
                        comment_data["author"] = str(c.author) if c.author else None
                    except:
                        comment_data["author"] = None
                    
                    try:
                        comment_data["timestamp"] = str(c.timestamp) if c.timestamp else None
                    except:
                        comment_data["timestamp"] = None
                    
                    try:
                        comment_data["comment_id"] = str(c.comment_id) if c.comment_id else None
                    except:
                        comment_data["comment_id"] = None
                    
                    try:
                        comment_data["initials"] = str(c.initials) if c.initials else None
                    except:
                        comment_data["initials"] = None
                    
                    comments.append(comment_data)
            else:
                print("No comments found in the document")
            
            return comments
            
        except Exception as e:
            print(f"Error extracting comments: {e}")
            return []
    
    def save_to_json(self, data: Dict[str, Any], output_path: str):
        """
        Save header/footer data to JSON file
        
        Args:
            data: Dictionary containing header/footer data
            output_path: Path to save the JSON file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"Header/footer data saved to JSON: {output_path}")
    
    def save_to_txt(self, data: Dict[str, Any], output_path: str):
        """
        Save header/footer data to TXT file
        
        Args:
            data: Dictionary containing header/footer data
            output_path: Path to save the TXT file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("HEADER, FOOTER, AND MARGIN DATA FROM DOCUMENT\n")
            f.write("=" * 60 + "\n\n")
            
            # Header information
            f.write("HEADER CONTENT:\n")
            f.write("-" * 20 + "\n")
            if data.get('header'):
                header = data['header']
                if header.get('paragraphs'):
                    f.write("Paragraphs:\n")
                    for i, para in enumerate(header['paragraphs'], 1):
                        f.write(f"  {i}. {para}\n")
                
                if header.get('tables'):
                    f.write("\nTables:\n")
                    for i, table in enumerate(header['tables'], 1):
                        f.write(f"  Table {i}:\n")
                        for j, row in enumerate(table, 1):
                            f.write(f"    Row {j}: {row}\n")
                
                if header.get('text_boxes'):
                    f.write("\nText Boxes:\n")
                    for i, textbox in enumerate(header['text_boxes'], 1):
                        f.write(f"  Text Box {i}:\n")
                        for j, para in enumerate(textbox, 1):
                            f.write(f"    {j}. {para}\n")
            else:
                f.write("No header content found\n")
            
            # Footer information
            f.write("\n\nFOOTER CONTENT:\n")
            f.write("-" * 20 + "\n")
            if data.get('footer'):
                footer = data['footer']
                if footer.get('paragraphs'):
                    f.write("Paragraphs:\n")
                    for i, para in enumerate(footer['paragraphs'], 1):
                        f.write(f"  {i}. {para}\n")
                
                if footer.get('tables'):
                    f.write("\nTables:\n")
                    for i, table in enumerate(footer['tables'], 1):
                        f.write(f"  Table {i}:\n")
                        for j, row in enumerate(table, 1):
                            f.write(f"    Row {j}: {row}\n")
                
                if footer.get('text_boxes'):
                    f.write("\nText Boxes:\n")
                    for i, textbox in enumerate(footer['text_boxes'], 1):
                        f.write(f"  Text Box {i}:\n")
                        for j, para in enumerate(textbox, 1):
                            f.write(f"    {j}. {para}\n")
            else:
                f.write("No footer content found\n")
            
            # Margin settings
            f.write("\n\nMARGIN SETTINGS:\n")
            f.write("-" * 20 + "\n")
            if data.get('margins'):
                for key, value in data['margins'].items():
                    f.write(f"{key.replace('_', ' ').title()}: {value} inches\n")
            else:
                f.write("No margin information found\n")
            
            # Comments
            f.write("\n\nCOMMENTS:\n")
            f.write("-" * 20 + "\n")
            if data.get('comments'):
                for i, comment in enumerate(data['comments'], 1):
                    f.write(f"Comment {i}:\n")
                    f.write(f"  Text: {comment['text']}\n")
                    if comment.get('author'):
                        f.write(f"  Author: {comment['author']}\n")
                    if comment.get('timestamp'):
                        f.write(f"  Timestamp: {comment['timestamp']}\n")
                    if comment.get('initials'):
                        f.write(f"  Initials: {comment['initials']}\n")
                    if comment.get('comment_id'):
                        f.write(f"  Comment ID: {comment['comment_id']}\n")
                    f.write("-" * 30 + "\n")
            else:
                f.write("No comments found\n")
        
        print(f"Header/footer data saved to TXT: {output_path}")

def main():
    """Main function for standalone usage"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python header_footer_extractor.py <docx_file>")
        print("Example: python header_footer_extractor.py 'SECTION 26 05 00.docx'")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not os.path.exists(docx_file):
        print(f"Error: File '{docx_file}' not found.")
        sys.exit(1)
    
    try:
        extractor = HeaderFooterExtractor()
        data = extractor.extract_all(docx_file)
        
        if data:
            # Generate output filenames
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            json_file = f"{base_name}_header_footer.json"
            txt_file = f"{base_name}_header_footer.txt"
            
            # Save to JSON
            extractor.save_to_json(data, json_file)
            
            # Save to TXT
            extractor.save_to_txt(data, txt_file)
            
            print(f"\nExtracted header, footer, and margin data from {docx_file}")
            print("Files created:")
            print(f"  - {json_file}")
            print(f"  - {txt_file}")
            
            # Show a preview of the extracted data
            print("\nData Preview:")
            
            # Header preview
            if data.get('header'):
                header = data['header']
                if header.get('paragraphs'):
                    print(f"  Header paragraphs: {len(header['paragraphs'])}")
                    for i, para in enumerate(header['paragraphs'][:2], 1):
                        print(f"    {i}. {para[:50]}...")
            
            # Footer preview
            if data.get('footer'):
                footer = data['footer']
                if footer.get('paragraphs'):
                    print(f"  Footer paragraphs: {len(footer['paragraphs'])}")
                    for i, para in enumerate(footer['paragraphs'][:2], 1):
                        print(f"    {i}. {para[:50]}...")
            
            # Comments preview
            if data.get('comments'):
                print(f"  Comments: {len(data['comments'])}")
                for i, comment in enumerate(data['comments'][:2], 1):
                    print(f"    {i}. {comment['text'][:50]}...")
            
            # Margins preview
            if data.get('margins'):
                print(f"  Margins: {len(data['margins'])} settings extracted")
        else:
            print("No data extracted")
            
    except Exception as e:
        print(f"Error processing document: {e}")

if __name__ == "__main__":
    main() 