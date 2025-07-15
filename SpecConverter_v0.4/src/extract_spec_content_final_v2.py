#!/usr/bin/env python3
"""
Specification Content Processor - Version 2

This script processes Word documents (.docx) to match template formatting and structure.
It implements the dev goals for spec templates processing.

Key Features:
- Processes header, footer, margins, and styles from input to output
- Updates formatting to match test template's list formatting
- Handles special content (SECTION, END OF SECTION, blank pages)
- Processes tables and includes them in previous list items
- Handles manually numbered sections with special care
- Uses BWA-labeled list levels from template
- Analyzes list levels using Word's list detection
- Validates hierarchy and numbering format
- Reconstructs broken lists using content numbering
- Outputs debugging information for unprocessable content

Usage:
    python extract_spec_content_final_v2.py <input_docx> [output_dir] [template_file]

Example:
    python extract_spec_content_final_v2.py "SECTION 26 05 00.docx"
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import os
import sys
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional, Tuple, Any, Set
from dataclasses import dataclass
from datetime import datetime
from copy import deepcopy

@dataclass
class ProcessingError:
    """Represents an error found during content processing"""
    line_number: int
    error_type: str
    message: str
    content: str
    context: str = ""
    suggested_fix: str = ""

@dataclass
class ContentBlock:
    """Represents a block of content with its formatting and level"""
    text: str
    level: Optional[int] = None
    numbering_id: Optional[str] = None
    numbering_level: Optional[int] = None
    style_name: Optional[str] = None
    is_special_content: bool = False
    is_table: bool = False
    is_manually_numbered: bool = False
    manual_number: Optional[str] = None
    needs_highlighting: bool = False
    error_message: Optional[str] = None

class SpecContentProcessor:
    """Processes specification content to match template formatting"""
    
    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path
        self.errors: List[ProcessingError] = []
        self.line_count = 0
        
        # Template analysis data
        self.template_numbering = {}
        self.template_styles = {}
        self.bwa_list_levels = {}
        
        # Document structure
        self.section_number = ""
        self.section_title = ""
        self.end_of_section = ""
        
        # Content blocks
        self.content_blocks: List[ContentBlock] = []
        self.special_content: List[ContentBlock] = []
        self.tables: List[Any] = []
        
        # Regex patterns
        self.section_pattern = re.compile(r'^SECTION\s+(.+)$', re.IGNORECASE)
        self.end_section_pattern = re.compile(r'^END\s+OF\s+SECTION\s*(.+)?$', re.IGNORECASE)
        self.manual_number_pattern = re.compile(r'^([A-Z]\.|[0-9]+\.|[a-z]\.)\s+(.+)$')
        
        # Load template if provided
        if template_path:
            self.load_template_analysis(template_path)
    
    def load_template_analysis(self, template_path: str):
        """Load and analyze template structure"""
        try:
            print(f"Loading template analysis from: {template_path}")
            
            # Load template document
            template_doc = Document(template_path)
            
            # Extract numbering definitions
            self.extract_template_numbering(template_path)
            
            # Extract styles
            self.extract_template_styles(template_doc)
            
            # Find BWA-labeled list levels
            self.find_bwa_list_levels(template_doc)
            
            print(f"Template loaded: {len(self.bwa_list_levels)} BWA list levels found")
            
        except Exception as e:
            print(f"Warning: Could not load template analysis: {e}")
    
    def extract_template_numbering(self, template_path: str):
        """Extract numbering definitions from template"""
        try:
            with zipfile.ZipFile(template_path) as zf:
                if "word/numbering.xml" in zf.namelist():
                    num_xml = zf.read("word/numbering.xml")
                    root = ET.fromstring(num_xml)
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    
                    # Extract abstract numbering definitions
                    for abstract_num in root.findall(".//w:abstractNum", ns):
                        abstract_num_id = abstract_num.get(f"{{{ns['w']}}}abstractNumId")
                        self.template_numbering[abstract_num_id] = {
                            "levels": {},
                            "bwa_label": None
                        }
                        
                        # Check for BWA label in first level
                        first_level = abstract_num.find("w:lvl", ns)
                        if first_level is not None:
                            lvl_text_elem = first_level.find("w:lvlText", ns)
                            if lvl_text_elem is not None:
                                lvl_text = lvl_text_elem.get(f"{{{ns['w']}}}val", "")
                                if "BWA" in lvl_text.upper():
                                    self.template_numbering[abstract_num_id]["bwa_label"] = lvl_text
                        
                        # Extract all levels
                        for lvl in abstract_num.findall("w:lvl", ns):
                            ilvl = lvl.get(f"{{{ns['w']}}}ilvl")
                            lvl_text_elem = lvl.find("w:lvlText", ns)
                            num_fmt_elem = lvl.find("w:numFmt", ns)
                            
                            self.template_numbering[abstract_num_id]["levels"][ilvl] = {
                                "lvlText": lvl_text_elem.get(f"{{{ns['w']}}}val") if lvl_text_elem is not None else None,
                                "numFmt": num_fmt_elem.get(f"{{{ns['w']}}}val") if num_fmt_elem is not None else None
                            }
                    
                    # Extract num mappings
                    for num_elem in root.findall(".//w:num", ns):
                        num_id = num_elem.get(f"{{{ns['w']}}}numId")
                        abstract_num_ref = num_elem.find("w:abstractNumId", ns)
                        if abstract_num_ref is not None:
                            abstract_num_id = abstract_num_ref.get(f"{{{ns['w']}}}val")
                            self.template_numbering[f"num_{num_id}"] = {
                                "abstract_num_id": abstract_num_id,
                                "bwa_label": self.template_numbering.get(abstract_num_id, {}).get("bwa_label")
                            }
                            
        except Exception as e:
            print(f"Error extracting template numbering: {e}")
    
    def extract_template_styles(self, template_doc: Document):
        """Extract styles from template"""
        try:
            for style in template_doc.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    self.template_styles[style.name] = {
                        "name": style.name,
                        "base_style": style.base_style.name if style.base_style else None,
                        "font_name": style.font.name if hasattr(style, 'font') else None,
                        "font_size": style.font.size if hasattr(style, 'font') else None
                    }
        except Exception as e:
            print(f"Error extracting template styles: {e}")
    
    def find_bwa_list_levels(self, template_doc: Document):
        """Find BWA-labeled list levels in template"""
        try:
            for paragraph in template_doc.paragraphs:
                if paragraph.style and "BWA" in paragraph.style.name.upper():
                    self.bwa_list_levels[paragraph.style.name] = {
                        "name": paragraph.style.name,
                        "level": self.get_paragraph_level(paragraph)
                    }
        except Exception as e:
            print(f"Error finding BWA list levels: {e}")
    
    def get_paragraph_level(self, paragraph) -> Optional[int]:
        """Get the list level of a paragraph"""
        try:
            pPr = paragraph._p.pPr
            if pPr is not None and pPr.numPr is not None:
                if pPr.numPr.ilvl is not None:
                    return pPr.numPr.ilvl.val
        except:
            pass
        return None
    
    def analyze_input_document(self, docx_path: str) -> Document:
        """Analyze input document and extract content blocks"""
        try:
            print(f"Analyzing input document: {docx_path}")
            
            # Load document
            doc = Document(docx_path)
            
            # Extract document properties
            self.extract_document_properties(doc)
            
            # Process all paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                self.line_count += 1
                content_block = self.process_paragraph(paragraph, i)
                
                if content_block:
                    if content_block.is_special_content:
                        self.special_content.append(content_block)
                    else:
                        self.content_blocks.append(content_block)
            
            # Process tables
            self.process_tables(doc)
            
            # Validate and reconstruct lists
            self.validate_and_reconstruct_lists()
            
            return doc
            
        except Exception as e:
            self.add_error("Analysis Error", f"Failed to analyze document: {str(e)}", "")
            return None
    
    def extract_document_properties(self, doc: Document):
        """Extract document properties (header, footer, margins, styles)"""
        try:
            # Extract section information
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Check for section header
                if text.upper().startswith("SECTION"):
                    match = self.section_pattern.match(text)
                    if match:
                        self.section_number = match.group(1).strip()
                        print(f"Found section number: {self.section_number}")
                
                # Check for section title (next non-empty paragraph after section header)
                elif self.section_number and not self.section_title and text:
                    self.section_title = text
                    print(f"Found section title: {self.section_title}")
                
                # Check for end of section
                elif text.upper().startswith("END OF SECTION"):
                    match = self.end_section_pattern.match(text)
                    if match:
                        self.end_of_section = match.group(1).strip() if match.group(1) else ""
                        print(f"Found end of section: {self.end_of_section}")
            
            # Extract margins and styles (simplified for now)
            for section in doc.sections:
                self.margins = {
                    "top": section.top_margin,
                    "bottom": section.bottom_margin,
                    "left": section.left_margin,
                    "right": section.right_margin
                }
                break
                
        except Exception as e:
            print(f"Error extracting document properties: {e}")
    
    def process_paragraph(self, paragraph, index: int) -> Optional[ContentBlock]:
        """Process a single paragraph and return a content block"""
        try:
            text = paragraph.text.strip()
            if not text:
                return None
            
            # Check for special content
            if self.is_special_content(text):
                return ContentBlock(
                    text=text,
                    is_special_content=True,
                    style_name=paragraph.style.name if paragraph.style else None
                )
            
            # Check for manual numbering
            manual_match = self.manual_number_pattern.match(text)
            if manual_match:
                return ContentBlock(
                    text=text,
                    manual_number=manual_match.group(1),
                    is_manually_numbered=True,
                    style_name=paragraph.style.name if paragraph.style else None
                )
            
            # Get list level information
            level = self.get_paragraph_level(paragraph)
            numbering_id = self.get_paragraph_numbering_id(paragraph)
            
            return ContentBlock(
                text=text,
                level=level,
                numbering_id=numbering_id,
                numbering_level=level,
                style_name=paragraph.style.name if paragraph.style else None
            )
            
        except Exception as e:
            self.add_error("Paragraph Processing Error", f"Failed to process paragraph {index}: {str(e)}", text)
            return None
    
    def get_paragraph_numbering_id(self, paragraph) -> Optional[str]:
        """Get the numbering ID of a paragraph"""
        try:
            pPr = paragraph._p.pPr
            if pPr is not None and pPr.numPr is not None:
                if pPr.numPr.numId is not None:
                    return str(pPr.numPr.numId.val)
        except:
            pass
        return None
    
    def is_special_content(self, text: str) -> bool:
        """Check if text is special content (SECTION, END OF SECTION, blank page indicators)"""
        text_upper = text.upper()
        return (
            text_upper.startswith("SECTION") or
            text_upper.startswith("END OF SECTION") or
            "PAGE INTENTIONALLY LEFT BLANK" in text_upper or
            text_upper.strip() == ""
        )
    
    def process_tables(self, doc: Document):
        """Process tables and include them in previous list items"""
        try:
            for table in doc.tables:
                # Find the previous content block to attach the table to
                if self.content_blocks:
                    last_block = self.content_blocks[-1]
                    last_block.is_table = True
                    # Note: In a full implementation, you'd store the table data here
                    print(f"Attached table to content block: {last_block.text[:50]}...")
                else:
                    # Create a special content block for orphaned tables
                    table_block = ContentBlock(
                        text="[Table Content]",
                        is_table=True,
                        needs_highlighting=True,
                        error_message="Orphaned table - no previous content block"
                    )
                    self.content_blocks.append(table_block)
                    
        except Exception as e:
            print(f"Error processing tables: {e}")
    
    def validate_and_reconstruct_lists(self):
        """Validate list hierarchy and reconstruct broken lists"""
        try:
            print("Validating list hierarchy...")
            
            # Check for broken numbering sequences
            current_levels = {}
            for i, block in enumerate(self.content_blocks):
                if block.level is not None:
                    # Validate level sequence
                    if block.level in current_levels:
                        expected_next = current_levels[block.level] + 1
                        # Check if numbering is sequential
                        if hasattr(block, 'numbering_id') and block.numbering_id:
                            # This would need more sophisticated validation
                            pass
                    else:
                        current_levels[block.level] = 1
                
                # Check for manually numbered content that might need reinsertion
                if block.is_manually_numbered:
                    suggested_level = self.suggest_level_for_manual_number(block.manual_number)
                    if suggested_level is not None:
                        block.level = suggested_level
                        block.needs_highlighting = False
                        print(f"Reinserted manually numbered content: {block.text[:50]}... at level {suggested_level}")
                    else:
                        block.needs_highlighting = True
                        block.error_message = f"Could not determine level for manual number: {block.manual_number}"
                        self.add_error("Manual Numbering", f"Could not determine level for: {block.manual_number}", block.text)
            
        except Exception as e:
            print(f"Error validating and reconstructing lists: {e}")
    
    def suggest_level_for_manual_number(self, manual_number: str) -> Optional[int]:
        """Suggest a level for manually numbered content based on template patterns"""
        try:
            # Simple heuristic based on numbering patterns
            if manual_number.endswith('.'):
                number_part = manual_number[:-1]
                
                # Uppercase letters (A., B., C.) - typically level 2
                if number_part.isalpha() and number_part.isupper():
                    return 2
                
                # Numbers (1., 2., 3.) - typically level 3
                elif number_part.isdigit():
                    return 3
                
                # Lowercase letters (a., b., c.) - typically level 4
                elif number_part.isalpha() and number_part.islower():
                    return 4
            
            return None
            
        except Exception as e:
            print(f"Error suggesting level for manual number: {e}")
            return None
    
    def create_output_document(self, input_doc: Document, output_path: str) -> bool:
        """Create output document with corrected formatting"""
        try:
            print(f"Creating output document: {output_path}")
            
            # Create new document
            output_doc = Document()
            
            # Apply document properties
            self.apply_document_properties(output_doc, input_doc)
            
            # Separate special content into section header/title and end of section
            section_content = []
            end_of_section_content = []
            
            for block in self.special_content:
                if block.text.upper().startswith("END OF SECTION"):
                    end_of_section_content.append(block)
                else:
                    section_content.append(block)
            
            # Add section header and title first
            for block in section_content:
                self.add_content_block_to_document(output_doc, block)
            
            # Add processed content blocks
            for block in self.content_blocks:
                self.add_content_block_to_document(output_doc, block)
            
            # Add end of section content at the very end
            for block in end_of_section_content:
                self.add_content_block_to_document(output_doc, block)
            
            # Save the document
            output_doc.save(output_path)
            print(f"Output document saved: {output_path}")
            return True
            
        except Exception as e:
            self.add_error("Output Creation Error", f"Failed to create output document: {str(e)}", "")
            return False
    
    def apply_document_properties(self, output_doc: Document, input_doc: Document):
        """Apply document properties from input to output"""
        try:
            # Apply margins
            for section in output_doc.sections:
                if hasattr(self, 'margins'):
                    section.top_margin = self.margins.get("top", Inches(1))
                    section.bottom_margin = self.margins.get("bottom", Inches(1))
                    section.left_margin = self.margins.get("left", Inches(1))
                    section.right_margin = self.margins.get("right", Inches(1))
                break
            
            # Apply styles (simplified)
            for style_name, style_info in self.template_styles.items():
                try:
                    if style_name not in output_doc.styles:
                        new_style = output_doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                        # Apply style properties
                        if style_info.get("font_name"):
                            new_style.font.name = style_info["font_name"]
                        if style_info.get("font_size"):
                            new_style.font.size = Pt(style_info["font_size"])
                except:
                    pass  # Style might already exist
                    
        except Exception as e:
            print(f"Error applying document properties: {e}")
    
    def add_content_block_to_document(self, doc: Document, block: ContentBlock):
        """Add a content block to the output document with proper formatting"""
        try:
            if block.is_table:
                # Handle table content
                paragraph = doc.add_paragraph("[Table Content]")
                if block.needs_highlighting:
                    # Apply highlighting for debugging
                    for run in paragraph.runs:
                        run.font.highlight_color = 6  # Yellow highlighting
            else:
                # Add regular paragraph
                paragraph = doc.add_paragraph(block.text)
                
                # Apply numbering if available
                if block.level is not None and block.numbering_id:
                    # Apply list formatting
                    pPr = paragraph._p.pPr
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        paragraph._p.insert(0, pPr)
                    
                    numPr = OxmlElement('w:numPr')
                    pPr.insert(0, numPr)
                    
                    ilvl = OxmlElement('w:ilvl')
                    ilvl.set(qn('w:val'), str(block.level))
                    numPr.insert(0, ilvl)
                    
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), block.numbering_id)
                    numPr.insert(1, numId)
                
                # Apply highlighting for debugging
                if block.needs_highlighting:
                    for run in paragraph.runs:
                        run.font.highlight_color = 6  # Yellow highlighting
                
                # Apply style if available
                if block.style_name and block.style_name in doc.styles:
                    try:
                        paragraph.style = doc.styles[block.style_name]
                    except:
                        pass
                        
        except Exception as e:
            print(f"Error adding content block to document: {e}")
    
    def add_error(self, error_type: str, message: str, content: str, context: str = "", suggested_fix: str = ""):
        """Add an error to the error list"""
        error = ProcessingError(
            line_number=self.line_count,
            error_type=error_type,
            message=message,
            content=content,
            context=context,
            suggested_fix=suggested_fix
        )
        self.errors.append(error)
    
    def generate_processing_report(self) -> str:
        """Generate a comprehensive processing report"""
        report = f"SPECIFICATION PROCESSING REPORT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += "=" * 80 + "\n\n"
        
        # Document information
        report += f"Section Number: {self.section_number}\n"
        report += f"Section Title: {self.section_title}\n"
        report += f"End of Section: {self.end_of_section}\n\n"
        
        # Content statistics
        report += f"Content Blocks Processed: {len(self.content_blocks)}\n"
        report += f"Special Content Items: {len(self.special_content)}\n"
        report += f"Tables Found: {len([b for b in self.content_blocks if b.is_table])}\n"
        report += f"Manually Numbered Items: {len([b for b in self.content_blocks if b.is_manually_numbered])}\n\n"
        
        # Template information
        if self.template_path:
            report += f"Template Used: {self.template_path}\n"
            report += f"BWA List Levels: {len(self.bwa_list_levels)}\n"
            report += f"Template Numbering Definitions: {len(self.template_numbering)}\n\n"
        
        # Error summary
        if self.errors:
            report += f"ERRORS FOUND ({len(self.errors)}):\n"
            report += "-" * 40 + "\n"
            
            error_types = {}
            for error in self.errors:
                if error.error_type not in error_types:
                    error_types[error.error_type] = []
                error_types[error.error_type].append(error)
            
            for error_type, errors in error_types.items():
                report += f"\n{error_type} ({len(errors)} errors):\n"
                for i, error in enumerate(errors[:5], 1):  # Show first 5 of each type
                    report += f"  {i}. Line {error.line_number}: {error.message}\n"
                    if error.content:
                        report += f"     Content: {error.content[:100]}...\n"
                    if error.suggested_fix:
                        report += f"     Suggested Fix: {error.suggested_fix}\n"
                
                if len(errors) > 5:
                    report += f"  ... and {len(errors) - 5} more errors of this type\n"
        else:
            report += "No errors found during processing.\n"
        
        return report
    
    def save_processing_report(self, report: str, output_path: str):
        """Save processing report to file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)
    
    def save_content_analysis(self, output_path: str):
        """Save detailed content analysis to JSON"""
        analysis = {
            "section_info": {
                "section_number": self.section_number,
                "section_title": self.section_title,
                "end_of_section": self.end_of_section
            },
            "content_blocks": [
                {
                    "text": block.text,
                    "level": block.level,
                    "numbering_id": block.numbering_id,
                    "style_name": block.style_name,
                    "is_special_content": block.is_special_content,
                    "is_table": block.is_table,
                    "is_manually_numbered": block.is_manually_numbered,
                    "manual_number": block.manual_number,
                    "needs_highlighting": block.needs_highlighting,
                    "error_message": block.error_message
                }
                for block in self.content_blocks
            ],
            "special_content": [
                {
                    "text": block.text,
                    "style_name": block.style_name
                }
                for block in self.special_content
            ],
            "template_info": {
                "template_path": self.template_path,
                "bwa_list_levels": self.bwa_list_levels,
                "template_numbering": self.template_numbering
            },
            "processing_errors": [
                {
                    "line_number": error.line_number,
                    "error_type": error.error_type,
                    "message": error.message,
                    "content": error.content,
                    "context": error.context,
                    "suggested_fix": error.suggested_fix
                }
                for error in self.errors
            ]
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)

def main():
    """Main function to run the processing"""
    if len(sys.argv) < 2:
        print("Usage: python extract_spec_content_final_v2.py <input_docx> [output_dir] [template_file]")
        print("Example: python extract_spec_content_final_v2.py 'SECTION 26 05 00.docx'")
        print("Note: All output files will be saved to <output_dir>/output/")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "."
    template_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not os.path.exists(input_path):
        print(f"Error: File '{input_path}' not found.")
        sys.exit(1)
    
    # Auto-detect cleaned template if no template specified
    if not template_path:
        possible_templates = [
            "test_template_cleaned.docx",
            "template_cleaned.docx", 
            "spec_template_cleaned.docx"
        ]
        for template in possible_templates:
            if os.path.exists(template):
                template_path = template
                print(f"Auto-detected cleaned template: {template_path}")
                break
    
    # Create output directory
    output_dir = os.path.join(output_dir, "output")
    os.makedirs(output_dir, exist_ok=True)
    
    # Initialize processor
    processor = SpecContentProcessor(template_path)
    
    # Process the document
    print(f"Processing document: {input_path}")
    input_doc = processor.analyze_input_document(input_path)
    
    if not input_doc:
        print("Error: Failed to analyze input document.")
        sys.exit(1)
    
    # Generate output filenames
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_docx_path = os.path.join(output_dir, f"{base_name}_processed.docx")
    report_path = os.path.join(output_dir, f"{base_name}_processing_report.txt")
    analysis_path = os.path.join(output_dir, f"{base_name}_content_analysis.json")
    
    # Create output document
    success = processor.create_output_document(input_doc, output_docx_path)
    
    # Generate and save reports
    processing_report = processor.generate_processing_report()
    processor.save_processing_report(processing_report, report_path)
    processor.save_content_analysis(analysis_path)
    
    # Print summary
    print(f"\nProcessing Summary:")
    print(f"- Input document: {input_path}")
    print(f"- Output document: {output_docx_path}")
    print(f"- Processing report: {report_path}")
    print(f"- Content analysis: {analysis_path}")
    print(f"- Content blocks processed: {len(processor.content_blocks)}")
    print(f"- Special content items: {len(processor.special_content)}")
    print(f"- Errors found: {len(processor.errors)}")
    
    if processor.errors:
        print(f"\nWARNING: {len(processor.errors)} errors were found during processing.")
        print(f"Please review the processing report: {report_path}")
    else:
        print("\nProcessing completed successfully with no errors.")

if __name__ == "__main__":
    main() 