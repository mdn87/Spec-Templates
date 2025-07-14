from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json

# Configuration variables - change these to modify font and size for all text
TEMPLATE_PATH = 'test_template.docx'
OUTPUT_PATH   = 'generated_spec.docx'
CONTENT_PATH  = 'SECTION 00 00 00.json'
FONT_NAME = 'Arial'
FONT_SIZE = 10

def clear_document(doc):
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

def set_font_and_size(paragraph):
    """Set font and size for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        # For compatibility with some versions of Word
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def apply_basic_numbering(paragraph, level=0):
    """Apply basic numbering to a paragraph"""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numId = numPr.get_or_add_numId()
    numId.val = 1
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = level

try:
    doc = Document(TEMPLATE_PATH)
    print(f"Using existing template file")
except ValueError as e:
    print("Template file error, creating new document...")
    doc = Document()

clear_document(doc)

# ───── Stage 1: Read & parse your test file ─────
def parse_spec_docx(CONTENT_PATH):
    with open(CONTENT_PATH, 'r') as file:
        content = json.load(file)
        doc = Document(content)
        spec = {"section_number": None, "section_title": None, "parts": []}
        current_part = None
        current_sub = None

    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name.lower() if p.style and hasattr(p.style, 'name') and p.style.name else ''

        # Example rule: Heading 1 → section header
        if style == 'heading 1' and text.upper().startswith('SECTION'):
            # e.g. "SECTION 00 00 00"
            _, num = text.split(' ', 1)
            spec['section_number'] = num
            continue

        # Example rule: Heading 2 → part
        if style == 'heading 2':
            current_part = {"part_title": text, "subsections": []}
            spec['parts'].append(current_part)
            continue

        # Example rule: Heading 3 → subsection
        if style == 'heading 3' and current_part is not None:
            current_sub = {"title": text, "items": []}
            current_part['subsections'].append(current_sub)
            continue

        # Example rule: List Paragraph → items
        if style == 'list paragraph' and current_sub is not None:
            current_sub['items'].append({"text": text})
            continue

        # …you can test for p.paragraph_format.left_indent
        # or p.text prefixes like "Item1", "   •" etc…

    return spec

doc.save(OUTPUT_PATH)
print(f"Document saved as '{OUTPUT_PATH}' with {FONT_SIZE}pt {FONT_NAME} font")
print(f"Numbering for parts, subsections, and items is applied from JSON content '{CONTENT_PATH}'")
print("Note: You must apply multilevel list font and size manually in Word")
print("by using a template that already has the multilevel list defined.")






