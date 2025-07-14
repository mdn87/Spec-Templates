from docx import Document
import json
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def extract_header_and_margins(docx_path):
    """Extract header information and margin settings from Word document"""
    doc = Document(docx_path)
    sec = doc.sections[0]
    
    # Extract margin settings with null checks
    margins = {}
    try:
        if sec.top_margin:
            margins["top_margin"] = sec.top_margin.inches
        if sec.bottom_margin:
            margins["bottom_margin"] = sec.bottom_margin.inches
        if sec.left_margin:
            margins["left_margin"] = sec.left_margin.inches
        if sec.right_margin:
            margins["right_margin"] = sec.right_margin.inches
        if sec.header_distance:
            margins["header_distance"] = sec.header_distance.inches
        if sec.footer_distance:
            margins["footer_distance"] = sec.footer_distance.inches
    except Exception as e:
        print(f"Warning: Could not extract margin settings: {e}")

    header_data = {}
    
    # 1) Simple paragraph scan
    for p in sec.header.paragraphs:
        text = p.text.strip()
        if ':' in text:
            key, val = text.split(':', 1)
            header_data[key.strip().lower().replace(' ', '_')] = val.strip()

    # 2) Scan header tables
    for tbl in sec.header.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) >= 2:
                key = cells[0].text.strip()
                val = cells[1].text.strip()
                if ':' in key:
                    # if the key itself has a colon
                    k, _ = key.split(':', 1)
                    header_data[k.strip().lower().replace(' ', '_')] = val
                else:
                    header_data[key.lower().replace(' ', '_')] = val

    # 3) Scan any text boxes (runs within drawing objects)
    try:
        hdr_element = sec.header._element
        for drawing in hdr_element.findall('.//w:txbxContent', namespaces=hdr_element.nsmap):
            for p in drawing.findall('.//w:p', namespaces=hdr_element.nsmap):
                texts = [n.text for n in p.findall('.//w:t', namespaces=hdr_element.nsmap) if n.text]
                line = ''.join(texts).strip()
                if ':' in line:
                    key, val = line.split(':', 1)
                    header_data[key.strip().lower().replace(' ', '_')] = val.strip()
    except Exception as e:
        print(f"Warning: Could not scan text boxes in header: {e}")

    return {"header": header_data, "margins": margins}

def save_header_to_json(header_data, output_path):
    """Save header data to JSON file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(header_data, f, indent=2)
    print(f"Header data saved to JSON: {output_path}")

def save_header_to_txt(header_data, output_path):
    """Save header data to TXT file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("HEADER AND MARGIN DATA FROM DOCUMENT\n")
        f.write("=" * 50 + "\n\n")
        
        # Header information
        f.write("HEADER INFORMATION:\n")
        f.write("-" * 20 + "\n")
        if header_data.get('header'):
            for key, value in header_data['header'].items():
                f.write(f"{key.replace('_', ' ').title()}: {value}\n")
        else:
            f.write("No header information found\n")
        
        f.write("\nMARGIN SETTINGS:\n")
        f.write("-" * 20 + "\n")
        if header_data.get('margins'):
            for key, value in header_data['margins'].items():
                f.write(f"{key.replace('_', ' ').title()}: {value} inches\n")
        else:
            f.write("No margin information found\n")
    
    print(f"Header data saved to TXT: {output_path}")

if __name__ == "__main__":
    try:
        docx_file = "SECTION 26 05 00.docx"
        header_data = extract_header_and_margins(docx_file)
        
        if header_data:
            # Save to JSON
            json_file = "SECTION 26 05 00_header.json"
            save_header_to_json(header_data, json_file)
            
            # Save to TXT
            txt_file = "SECTION 26 05 00_header.txt"
            save_header_to_txt(header_data, txt_file)
            
            print(f"\nExtracted header and margin data from {docx_file}")
            print("Files created:")
            print(f"  - {json_file}")
            print(f"  - {txt_file}")
            
            # Show a preview of the extracted data
            print("\nHeader Data Preview:")
            if header_data.get('header'):
                print("  Header Information:")
                for key, value in list(header_data['header'].items())[:3]:  # Show first 3 items
                    print(f"    {key.replace('_', ' ').title()}: {value}")
            
            if header_data.get('margins'):
                print("  Margin Settings:")
                for key, value in list(header_data['margins'].items())[:3]:  # Show first 3 items
                    print(f"    {key.replace('_', ' ').title()}: {value} inches")
        else:
            print("No header data found to save")
            
    except FileNotFoundError:
        print("File 'SECTION 26 05 00.docx' not found")
    except Exception as e:
        print(f"Error processing document: {e}")
