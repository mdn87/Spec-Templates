from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT

# Open the template file - python-docx can handle .dotx files
try:
    doc = Document('test_template.docx')
except ValueError as e:
    # If the template file gives an error, try creating a new document
    print("Template file error, creating new document...")
    doc = Document()

# Level 1
doc.add_paragraph('PART 1', style='Heading 1')
# Level 2
doc.add_paragraph('Test',  style='Heading 2')
# Level 1 again (auto resets Level 2)
doc.add_paragraph('PART 2', style='Heading 1')
doc.add_paragraph('Jest',  style='Heading 2')

doc.save('generated_spec.docx')
print("Document saved as 'generated_spec.docx'")
