from docx import Document
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def extract_header_and_margins(docx_path):
    doc = Document(docx_path)
    sec = doc.sections[0]
    margins = {
        "top_margin": sec.top_margin.inches,
        "bottom_margin": sec.bottom_margin.inches,
        "left_margin": sec.left_margin.inches,
        "right_margin": sec.right_margin.inches,
        "header_distance": sec.header_distance.inches,
        "footer_distance": sec.footer_distance.inches
    }

    header_data = {}
    # 1) Simple paragraph scan
    for p in sec.header.paragraphs:
        text = p.text.strip()
        if ':' in text:
            key, val = text.split(':', 1)
            header_data[key.strip().lower().replace(' ', '_')] = val.strip()

    # 2) Scan header tables
    for tbl in sec.header.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) >= 2:
                key = cells[0].text.strip()
                val = cells[1].text.strip()
                if ':' in key:
                    # if the key itself has a colon
                    k, _ = key.split(':',1)
                    header_data[k.strip().lower().replace(' ', '_')] = val
                else:
                    header_data[key.lower().replace(' ', '_')] = val

    # 3) Scan any text boxes (runs within drawing objects)
    hdr_element = sec.header._element
    for drawing in hdr_element.findall('.//w:txbxContent', namespaces=hdr_element.nsmap):
        for p in drawing.findall('.//w:p', namespaces=hdr_element.nsmap):
            texts = [n.text for n in p.findall('.//w:t', namespaces=hdr_element.nsmap) if n.text]
            line = ''.join(texts).strip()
            if ':' in line:
                key, val = line.split(':',1)
                header_data[key.strip().lower().replace(' ', '_')] = val.strip()

    return {"header": header_data, "margins": margins}

if __name__ == "__main__":
    result = extract_header_and_margins('/mnt/data/SECTION 26 05 00.docx')
    print(json.dumps(result, indent=2))
