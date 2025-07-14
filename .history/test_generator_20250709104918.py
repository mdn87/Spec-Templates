from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def apply_multilevel_numbering(doc):
    """Apply multilevel numbering using a simpler approach"""
    
    # Create a simple multilevel list definition
    numbering_xml = '''
    <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNum w:abstractNumId="0">
            <w:nsid w:val="12345678"/>
            <w:multiLevelType w:val="hybridMultilevel"/>
            
            <!-- Level 1 -->
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.0"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
            </w:lvl>
            
            <!-- Level 2 -->
            <w:lvl w:ilvl="1">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimalZero"/>
                <w:lvlText w:val="%1.%2"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="360" w:hanging="0"/>
                </w:pPr>
            </w:lvl>
            
            <!-- Level 3 -->
            <w:lvl w:ilvl="2">
                <w:start w:val="1"/>
                <w:numFmt w:val="upperLetter"/>
                <w:lvlText w:val="%3."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="0"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        
        <w:num w:numId="1">
            <w:abstractNumId w:val="0"/>
        </w:num>
    </w:numbering>
    '''
    
    # Parse the XML
    numbering_element = parse_xml(numbering_xml)
    
    # Add to document's numbering part
    if not hasattr(doc.part, 'numbering_part'):
        # Create numbering part if it doesn't exist
        from docx.opc.part import Part
        from docx.opc.constants import CONTENT_TYPE as CT
        
        numbering_part = Part(
            partname=doc.part.partname.replace('document.xml', 'numbering.xml'),
            content_type=CT.WML_NUMBERING,
            element=numbering_element
        )
        doc.part.numbering_part = numbering_part
    else:
        doc.part.numbering_part._element = numbering_element
    
    return doc

def apply_numbering_to_paragraph(paragraph, level=0):
    """Apply multilevel numbering to a paragraph"""
    # Get or create paragraph properties
    pPr = paragraph._p.get_or_add_pPr()
    
    # Get or create numbering properties
    numPr = pPr.get_or_add_numPr()
    
    # Set the numbering ID
    numId = numPr.get_or_add_numId()
    numId.val = 1
    
    # Set the level
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = level

# Open the template file
try:
    doc = Document('test_template.docx')
except ValueError as e:
    print("Template file error, creating new document...")
    doc = Document()

# Apply multilevel numbering
doc = apply_multilevel_numbering(doc)

# Add content with automatic numbering
# Level 1 (Parts)
p1 = doc.add_paragraph('PART 1', style='Heading 1')
apply_numbering_to_paragraph(p1, 0)

# Level 2 (Subsections)
p2 = doc.add_paragraph('Test Subsection', style='Heading 2')
apply_numbering_to_paragraph(p2, 1)

# Level 3 (Items)
p3 = doc.add_paragraph('Test Item', style='Heading 3')
apply_numbering_to_paragraph(p3, 2)

# Level 1 again (auto resets Level 2)
p4 = doc.add_paragraph('PART 2', style='Heading 1')
apply_numbering_to_paragraph(p4, 0)

p5 = doc.add_paragraph('Another Subsection', style='Heading 2')
apply_numbering_to_paragraph(p5, 1)

doc.save('generated_spec.docx')
print("Document saved as 'generated_spec.docx' with automatic multilevel numbering!")
