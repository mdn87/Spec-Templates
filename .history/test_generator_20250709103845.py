from docx import Document
doc = Document('spec_template.dotx')

# Level 1
doc.add_paragraph('PART 1', style='Heading 1')
# Level 2
doc.add_paragraph('Test',  style='Heading 2')
# Level 1 again (auto resets Level 2)
doc.add_paragraph('PART 2', style='Heading 1')
doc.add_paragraph('Jest',  style='Heading 2')

doc.save('generated_spec.docx')
