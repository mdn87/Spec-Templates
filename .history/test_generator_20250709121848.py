from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import os

# Configuration variables - change these to modify font and size for all text
TEMPLATE_PATH = 'test_template.docx'
OUTPUT_PATH   = 'generated_spec.docx'
CONTENT_PATH  = 'SECTION 00 00 00.json'
FONT_NAME = 'Arial'
FONT_SIZE = 10

def list_available_styles(doc):
    """List all available styles in the document"""
    print("DEBUG: Available styles in template:")
    for style in doc.styles:
        print(f"  - {style.name}")

def clear_document(doc):
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

def set_font_and_size(paragraph):
    """Set font and size for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        # For compatibility with some versions of Word
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def parse_spec_json(json_path):
    """Parse JSON file and return structured data"""
    print(f"DEBUG: Attempting to read JSON file: {json_path}")
    
    if not os.path.exists(json_path):
        print(f"ERROR: JSON file not found: {json_path}")
        return None
    
    try:
        with open(json_path, 'r', encoding='utf-8') as file:
            content = json.load(file)
            print(f"DEBUG: Successfully loaded JSON with {len(content) if isinstance(content, (list, dict)) else 'unknown'} items")
            print(f"DEBUG: JSON content type: {type(content)}")
            print(f"DEBUG: JSON keys (if dict): {list(content.keys()) if isinstance(content, dict) else 'Not a dict'}")
            return content
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON format: {e}")
        return None
    except Exception as e:
        print(f"ERROR: Failed to read JSON file: {e}")
        return None

def generate_content_from_json(doc, json_data):
    """Generate document content from JSON data using template styles"""
    print(f"DEBUG: Starting content generation from JSON data")
    
    if json_data is None:
        print("ERROR: No JSON data to process")
        return
    
    # Add section header (outside multilevel list numbering)
    if isinstance(json_data, dict) and 'section_number' in json_data:
        section_text = f"SECTION {json_data['section_number']}"
        p1 = doc.add_paragraph(section_text, style='Heading 1')
        set_font_and_size(p1)
        print(f"DEBUG: Added section header: {section_text}")
    
    # Add section title (outside multilevel list numbering)
    if isinstance(json_data, dict) and 'section_title' in json_data:
        title_text = json_data['section_title']
        p2 = doc.add_paragraph(title_text, style='Heading 2')
        set_font_and_size(p2)
        print(f"DEBUG: Added section title: {title_text}")
    
    # Add parts (use template's multilevel list style)
    if isinstance(json_data, dict) and 'parts' in json_data:
        for i, part in enumerate(json_data['parts']):
            print(f"DEBUG: Processing part {i+1}: {part.get('part_title', 'Unknown')}")
            
            # Add part title using template's multilevel list style
            # Try different style names that might be in your template
            part_style = 'List Paragraph'  # Default fallback
            try:
                # Try to use a style that might be defined in your template
                p3 = doc.add_paragraph(part.get('part_title', f'Part {i+1}'), style=part_style)
                set_font_and_size(p3)
            except:
                # Fallback to basic paragraph
                p3 = doc.add_paragraph(part.get('part_title', f'Part {i+1}'))
                set_font_and_size(p3)
            
            # Add subsections
            if 'subsections' in part:
                for j, sub in enumerate(part['subsections']):
                    print(f"DEBUG: Processing subsection {j+1}: {sub.get('title', 'Unknown')}")
                    
                    # Add subsection title
                    p4 = doc.add_paragraph(sub.get('title', f'Subsection {j+1}'), style='List Paragraph')
                    set_font_and_size(p4)
                    
                    # Add items
                    if 'items' in sub:
                        for k, item in enumerate(sub['items']):
                            print(f"DEBUG: Processing item {k+1}: {item.get('text', 'Unknown')[:50]}...")
                            
                            # Add main item text
                            p5 = doc.add_paragraph(item.get('text', f'Item {k+1}'), style='List Paragraph')
                            set_font_and_size(p5)
                            
                            # Add lists
                            if 'lists' in item:
                                for l, list_item in enumerate(item['lists']):
                                    print(f"DEBUG: Processing list item {l+1}: {list_item.get('text', 'Unknown')[:50]}...")
                                    p6 = doc.add_paragraph(list_item.get('text', f'List Item {l+1}'), style='List Paragraph')
                                    set_font_and_size(p6)
                            
                            # Add sublists
                            if 'sublists' in item:
                                for m, sublist_item in enumerate(item['sublists']):
                                    print(f"DEBUG: Processing sublist item {m+1}: {sublist_item.get('text', 'Unknown')[:50]}...")
                                    p7 = doc.add_paragraph(sublist_item.get('text', f'SubList Item {m+1}'), style='List Paragraph')
                                    set_font_and_size(p7)
    else:
        print("DEBUG: No 'parts' found in JSON data, adding fallback content")
        # Fallback: add some test content
        p1 = doc.add_paragraph('PART 1', style='List Paragraph')
        set_font_and_size(p1)
        
        p2 = doc.add_paragraph('Test Subsection', style='List Paragraph')
        set_font_and_size(p2)

# Main execution
print("DEBUG: Starting document generation process")

# Load template
try:
    doc = Document(TEMPLATE_PATH)
    print(f"DEBUG: Successfully loaded template: {TEMPLATE_PATH}")
    
    # List available styles to see what's in the template
    list_available_styles(doc)
    
except ValueError as e:
    print(f"DEBUG: Template file error, creating new document: {e}")
    doc = Document()

# Clear existing content
clear_document(doc)
print("DEBUG: Cleared existing document content")

# Parse JSON content
json_data = parse_spec_json(CONTENT_PATH)

# Generate content from JSON
generate_content_from_json(doc, json_data)

# Save document
doc.save(OUTPUT_PATH)
print(f"DEBUG: Document saved as '{OUTPUT_PATH}' with {FONT_SIZE}pt {FONT_NAME} font")
print(f"DEBUG: Content source: {CONTENT_PATH}")
print("Note: The template's multilevel list style should be applied automatically")
print("if the paragraphs use the correct style names from the template.")






