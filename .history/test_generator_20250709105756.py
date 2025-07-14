from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Configuration variables - change these to modify font and size for all text
FONT_NAME = 'Arial'
FONT_SIZE = 10

def clear_document(doc):
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

def set_font_and_size(paragraph):
    """Set font and size for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        # For compatibility with some versions of Word
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def apply_basic_numbering(paragraph, level=0):
    """Apply basic numbering to a paragraph and set numbering font/size"""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numId = numPr.get_or_add_numId()
    numId.val = 1
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = level
    
    # Set numbering font and size
    # This affects the numbering itself, not just the text
    numPr.get_or_add_numPr().get_or_add_numId().val = 1
    # Add numbering level properties for font
    lvl = numPr.get_or_add_lvl()
    lvl.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'), FONT_NAME)
    lvl.get_or_add_rPr().get_or_add_sz().val = FONT_SIZE * 2  # Word uses half-points

try:
    doc = Document('test_template.docx')
    print(f"Using existing template file")
except ValueError as e:
    print("Template file error, creating new document...")
    doc = Document()

clear_document(doc)

p1 = doc.add_paragraph('PART 1', style='Heading 1')
apply_basic_numbering(p1, 0)
set_font_and_size(p1)

p2 = doc.add_paragraph('Test Subsection', style='Heading 2')
apply_basic_numbering(p2, 1)
set_font_and_size(p2)

p3 = doc.add_paragraph('Test Item', style='Heading 3')
apply_basic_numbering(p3, 2)
set_font_and_size(p3)

p4 = doc.add_paragraph('PART 2', style='Heading 1')
apply_basic_numbering(p4, 0)
set_font_and_size(p4)

p5 = doc.add_paragraph('Another Subsection', style='Heading 2')
apply_basic_numbering(p5, 1)
set_font_and_size(p5)

p6 = doc.add_paragraph('This is regular paragraph text without numbering.')
set_font_and_size(p6)

doc.save('generated_spec.docx')
print(f"Document saved as 'generated_spec.docx' with {FONT_SIZE}pt {FONT_NAME} font and numbering!")
print("Note: You may need to apply multilevel list formatting manually in Word")
print("or use a template that already has the multilevel list defined.")
