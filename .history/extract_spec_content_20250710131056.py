#!/usr/bin/env python3
"""
Specification Content Extractor

This script extracts multi-level list content from Word documents (.docx) and converts it to JSON format.
It handles the hierarchical structure of specification documents with parts, subsections, and items.
The script includes error detection and reporting for broken or inconsistent content.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os
import re
import sys
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime

@dataclass
class ExtractionError:
    """Represents an error found during content extraction"""
    line_number: int
    error_type: str
    message: str
    context: str
    expected: Optional[str] = None
    found: Optional[str] = None

class SpecContentExtractor:
    """Extracts specification content from Word documents"""
    
    def __init__(self):
        self.errors: List[ExtractionError] = []
        self.current_part_number: Optional[str] = None
        self.current_subsection_number: Optional[str] = None
        self.current_item_number: Optional[str] = None
        self.line_count = 0
        
        # Regex patterns for different levels
        self.part_pattern = re.compile(r'^(\d+\.0)\s+(.+)$')
        self.subsection_pattern = re.compile(r'^(\d+\.\d{2})\s+(.+)$')
        self.item_pattern = re.compile(r'^([A-Z])\.\s+(.+)$')
        self.list_pattern = re.compile(r'^(\d+)\.\s+(.+)$')
        self.sub_list_pattern = re.compile(r'^([a-z])\.\s+(.+)$')
        
        # Additional patterns for different formats
        self.section_pattern = re.compile(r'^SECTION\s+(\d+\s+\d+\s+\d+)$')
        self.title_pattern = re.compile(r'^([A-Z][A-Z\s]+)$')
        
        # Patterns for actual document structure
        self.part_title_pattern = re.compile(r'^([A-Z][A-Z\s]+)$')  # GENERAL, PRODUCTS, EXECUTION
        self.subsection_title_pattern = re.compile(r'^([A-Z][A-Z\s&]+)$')  # SCOPE, EXISTING CONDITIONS, etc.
        
    def add_error(self, error_type: str, message: str, context: str = "", 
                  expected: Optional[str] = None, found: Optional[str] = None):
        """Add an error to the error list"""
        error = ExtractionError(
            line_number=self.line_count,
            error_type=error_type,
            message=message,
            context=context,
            expected=expected,
            found=found
        )
        self.errors.append(error)
        print(f"ERROR (line {self.line_count}): {error_type} - {message}")
        if context:
            print(f"  Context: {context}")
        if expected and found:
            print(f"  Expected: {expected}, Found: {found}")
    
    def extract_text_from_element(self, element, nsmap) -> str:
        """Extract all text from an element and its children"""
        texts = []
        for text_elem in element.findall('.//w:t', namespaces=nsmap):
            if text_elem.text:
                texts.append(text_elem.text)
        return ''.join(texts).strip()
    
    def get_paragraph_style(self, paragraph) -> str:
        """Get the style name of a paragraph"""
        try:
            return paragraph.style.name if paragraph.style else "Normal"
        except:
            return "Normal"
    
    def get_paragraph_numbering(self, paragraph) -> Optional[str]:
        """Get the numbering level and format of a paragraph"""
        try:
            if paragraph._element.pPr and paragraph._element.pPr.numPr:
                numPr = paragraph._element.pPr.numPr
                if numPr.numId and numPr.ilvl:
                    return f"Level {numPr.ilvl.val}"
        except:
            pass
        return None
    
    def parse_paragraph_content(self, text: str) -> Tuple[Optional[str], Optional[str], str]:
        """
        Parse paragraph content to determine level and extract content
        Returns: (level_type, number, content)
        """
        text = text.strip()
        if not text:
            return None, None, ""
        
        # Check for section header
        match = self.section_pattern.match(text)
        if match:
            return "section", match.group(1), ""
        
        # Check for section title (all caps, longer text)
        if text == "Common Work Results for Electrical":
            return "title", None, text
        
        # Check for part level with numbering (1.0, 2.0, etc.)
        match = self.part_pattern.match(text)
        if match:
            return "part", match.group(1), match.group(2)
        
        # Check for part titles without numbering (GENERAL, PRODUCTS, EXECUTION)
        if text in ["GENERAL", "PRODUCTS", "EXECUTION"]:
            return "part_title", None, text
        
        # Check for subsection level with numbering (1.01, 1.02, etc.)
        match = self.subsection_pattern.match(text)
        if match:
            return "subsection", match.group(1), match.group(2)
        
        # Check for subsection titles without numbering (SCOPE, EXISTING CONDITIONS, etc.)
        subsection_titles = [
            "SCOPE", "EXISTING CONDITIONS", "CODES AND REGULATIONS", "DEFINITIONS",
            "DRAWINGS AND SPECIFICATIONS", "SITE VISIT", "DEVIATIONS",
            "STANDARDS FOR MATERIALS AND WORKMANSHIP", "SHOP DRAWINGS AND SUBMITTAL",
            "RECORD (AS-BUILT) DRAWINGS AND MAINTENANCE MANUALS",
            "COORDINATION", "PROTECTION OF MATERIALS", "TESTS, DEMONSTRATION AND INSTRUCTIONS",
            "GUARANTEE"
        ]
        if text in subsection_titles:
            return "subsection_title", None, text
        
        # Check for item level (A., B., C., etc.)
        match = self.item_pattern.match(text)
        if match:
            return "item", match.group(1), match.group(2)
        
        # Check for list level (1., 2., etc.)
        match = self.list_pattern.match(text)
        if match:
            return "list", match.group(1), match.group(2)
        
        # Check for sub-list level (a., b., etc.)
        match = self.sub_list_pattern.match(text)
        if match:
            return "sub_list", match.group(1), match.group(2)
        
        # If no pattern matches, it's regular content
        return "content", None, text
    
    def validate_numbering_sequence(self, level_type: str, number: str, content: str):
        """Validate that numbering follows expected sequence"""
        if level_type == "part":
            expected_part = f"{len(self.extracted_data.get('parts', [])) + 1}.0"
            if number != expected_part:
                self.add_error(
                    "Numbering Sequence Error",
                    f"Unexpected part number",
                    content,
                    expected_part,
                    number
                )
            self.current_part_number = number
            
        elif level_type == "subsection":
            if not self.current_part_number:
                self.add_error(
                    "Structure Error",
                    "Subsection found without preceding part",
                    content
                )
            else:
                part_num = self.current_part_number.split('.')[0]
                # Find the current part to get its subsections
                current_part = None
                for part in self.extracted_data.get('parts', []):
                    if part.get('part_number') == self.current_part_number:
                        current_part = part
                        break
                
                if current_part:
                    expected_subsection = f"{part_num}.{len(current_part.get('subsections', [])) + 1:02d}"
                    if number != expected_subsection:
                        self.add_error(
                            "Numbering Sequence Error",
                            f"Unexpected subsection number",
                            content,
                            expected_subsection,
                            number
                        )
            self.current_subsection_number = number
            
        elif level_type == "item":
            if not self.current_subsection_number:
                self.add_error(
                    "Structure Error",
                    "Item found without preceding subsection",
                    content
                )
            else:
                # Find the current subsection to get its items
                current_subsection = None
                for part in self.extracted_data.get('parts', []):
                    for sub in part.get('subsections', []):
                        if sub.get('subsection_number') == self.current_subsection_number:
                            current_subsection = sub
                            break
                    if current_subsection:
                        break
                
                if current_subsection:
                    expected_item = chr(ord('A') + len(current_subsection.get('items', [])))
                    if number != expected_item:
                        self.add_error(
                            "Numbering Sequence Error",
                            f"Unexpected item number",
                            content,
                            expected_item,
                            number
                        )
            self.current_item_number = number
    
    def extract_header_info(self, doc) -> Dict[str, Any]:
        """Extract header information from document"""
        header_info = {
            "bwa_number": "2025-XXXX",
            "client_number": "ZZZ# 00000", 
            "project_name": "PROJECT NAME",
            "company_name": "CLIENT NAME",
            "section_number": "",
            "section_title": ""
        }
        
        # Try to extract section number and title from first few paragraphs
        for i, para in enumerate(doc.paragraphs[:10]):
            text = para.text.strip()
            if text.startswith("SECTION"):
                # Extract section number
                section_match = re.search(r'SECTION\s+(\d+\s+\d+\s+\d+)', text)
                if section_match:
                    header_info["section_number"] = section_match.group(1).replace(" ", "")
                
                # Look for title in next paragraph
                if i + 1 < len(doc.paragraphs):
                    title_text = doc.paragraphs[i + 1].text.strip()
                    if title_text and not title_text.startswith("SECTION"):
                        header_info["section_title"] = title_text
                break
        
        return header_info
    
    def extract_content(self, docx_path: str) -> Dict[str, Any]:
        """Extract all content from the Word document"""
        print(f"Extracting content from: {docx_path}")
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            self.add_error("File Error", f"Could not open document: {e}")
            return {}
        
        # Initialize extraction data
        self.extracted_data = {
            "header": self.extract_header_info(doc),
            "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
            "margins": {},
            "comments": [],
            "section_number": "",
            "section_title": "",
            "parts": []
        }
        
        # Extract margins
        try:
            sec = doc.sections[0]
            self.extracted_data["margins"] = {
                "top_margin": sec.top_margin.inches if sec.top_margin else 1.0,
                "bottom_margin": sec.bottom_margin.inches if sec.bottom_margin else 1.0,
                "left_margin": sec.left_margin.inches if sec.left_margin else 1.0,
                "right_margin": sec.right_margin.inches if sec.right_margin else 1.0,
                "header_distance": sec.header_distance.inches if sec.header_distance else 1.0,
                "footer_distance": sec.footer_distance.inches if sec.footer_distance else 1.0
            }
        except Exception as e:
            self.add_error("Margin Error", f"Could not extract margins: {e}")
        
        # Process paragraphs
        current_part = None
        current_subsection = None
        current_item = None
        
        for para in doc.paragraphs:
            self.line_count += 1
            text = para.text.strip()
            
            if not text:
                continue
            
            # Parse the paragraph content
            level_type, number, content = self.parse_paragraph_content(text)
            
            if level_type == "section":
                # Extract section number
                if number:
                    self.extracted_data["section_number"] = number.replace(" ", "")
                
            elif level_type == "title":
                # Set section title
                self.extracted_data["section_title"] = content
                
            elif level_type == "part":
                # Validate numbering
                if number:
                    self.validate_numbering_sequence(level_type, number, content)
                
                # Create new part
                current_part = {
                    "part_number": number,
                    "title": content,
                    "subsections": []
                }
                self.extracted_data["parts"].append(current_part)
                current_subsection = None
                current_item = None
                
            elif level_type == "subsection":
                # Validate numbering
                if number:
                    self.validate_numbering_sequence(level_type, number, content)
                
                if not current_part:
                    self.add_error("Structure Error", "Subsection found without part", content)
                    # Create a default part
                    current_part = {
                        "part_number": f"{len(self.extracted_data['parts']) + 1}.0",
                        "title": "GENERAL",
                        "subsections": []
                    }
                    self.extracted_data["parts"].append(current_part)
                
                # Create new subsection
                current_subsection = {
                    "subsection_number": number,
                    "title": content,
                    "items": []
                }
                current_part["subsections"].append(current_subsection)
                current_item = None
                
            elif level_type == "item":
                # Validate numbering
                if number:
                    self.validate_numbering_sequence(level_type, number, content)
                
                if not current_subsection:
                    self.add_error("Structure Error", "Item found without subsection", content)
                    # Create a default subsection
                    if not current_part:
                        current_part = {
                            "part_number": f"{len(self.extracted_data['parts']) + 1}.0",
                            "title": "GENERAL",
                            "subsections": []
                        }
                        self.extracted_data["parts"].append(current_part)
                    
                    current_subsection = {
                        "subsection_number": f"{current_part['part_number'].split('.')[0]}.01",
                        "title": "GENERAL",
                        "items": []
                    }
                    current_part["subsections"].append(current_subsection)
                
                # Create new item
                current_item = {
                    "item_number": number,
                    "text": content,
                    "lists": []
                }
                current_subsection["items"].append(current_item)
                
            elif level_type == "list":
                if not current_item:
                    self.add_error("Structure Error", "List item found without parent item", content)
                    continue
                
                # Add to current item's lists
                current_item["lists"].append({
                    "list_number": number,
                    "text": content
                })
                
            elif level_type == "sub_list":
                if not current_item:
                    self.add_error("Structure Error", "Sub-list item found without parent item", content)
                    continue
                
                # Add to current item's sublists (if not already present)
                if "sublists" not in current_item:
                    current_item["sublists"] = []
                
                current_item["sublists"].append({
                    "list_number": number,
                    "text": content
                })
                
            elif level_type == "content":
                # This might be continuation text for the current item
                if current_item:
                    # Append to the current item's text
                    current_item["text"] += " " + content
                else:
                    # Only report content warnings for non-empty content that's not just whitespace
                    if content.strip() and len(content.strip()) > 10:
                        self.add_error("Content Warning", "Unstructured content found", content[:100] + "..." if len(content) > 100 else content)
        
        # Set section number and title from header
        self.extracted_data["section_number"] = self.extracted_data["header"]["section_number"]
        self.extracted_data["section_title"] = self.extracted_data["header"]["section_title"]
        
        return self.extracted_data
    
    def generate_error_report(self) -> str:
        """Generate a detailed error report"""
        if not self.errors:
            return "No errors detected during extraction."
        
        report = f"ERROR REPORT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += "=" * 60 + "\n\n"
        
        # Group errors by type
        error_types = {}
        for error in self.errors:
            if error.error_type not in error_types:
                error_types[error.error_type] = []
            error_types[error.error_type].append(error)
        
        for error_type, errors in error_types.items():
            report += f"{error_type} ({len(errors)} errors):\n"
            report += "-" * (len(error_type) + 10) + "\n"
            
            for error in errors:
                report += f"Line {error.line_number}: {error.message}\n"
                if error.context:
                    report += f"  Context: {error.context}\n"
                if error.expected and error.found:
                    report += f"  Expected: {error.expected}, Found: {error.found}\n"
                report += "\n"
        
        return report
    
    def save_to_json(self, data: Dict[str, Any], output_path: str):
        """Save extracted data to JSON file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            print(f"Content saved to: {output_path}")
        except Exception as e:
            self.add_error("Save Error", f"Could not save JSON file: {e}")
    
    def save_error_report(self, report: str, output_path: str):
        """Save error report to text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report)
            print(f"Error report saved to: {output_path}")
        except Exception as e:
            print(f"Could not save error report: {e}")

def main():
    """Main function to process Word documents"""
    if len(sys.argv) < 2:
        print("Usage: python extract_spec_content.py <docx_file> [output_dir]")
        print("Example: python extract_spec_content.py 'SECTION 26 05 00.docx'")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "."
    
    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)
    
    # Create extractor and process document
    extractor = SpecContentExtractor()
    data = extractor.extract_content(docx_path)
    
    if not data:
        print("Error: No data extracted from document")
        sys.exit(1)
    
    # Generate output filenames
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    json_path = os.path.join(output_dir, f"{base_name}_content.json")
    error_path = os.path.join(output_dir, f"{base_name}_errors.txt")
    
    # Save results
    extractor.save_to_json(data, json_path)
    
    # Generate and save error report
    error_report = extractor.generate_error_report()
    extractor.save_error_report(error_report, error_path)
    
    # Print summary
    print(f"\nExtraction Summary:")
    print(f"  Document: {docx_path}")
    print(f"  Parts found: {len(data.get('parts', []))}")
    print(f"  Total subsections: {sum(len(part.get('subsections', [])) for part in data.get('parts', []))}")
    print(f"  Total items: {sum(len(sub.get('items', [])) for part in data.get('parts', []) for sub in part.get('subsections', []))}")
    print(f"  Errors detected: {len(extractor.errors)}")
    print(f"  Output files:")
    print(f"    - {json_path}")
    print(f"    - {error_path}")

if __name__ == "__main__":
    main() 