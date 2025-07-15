#!/usr/bin/env python3
"""
Specification Content Extractor - Final Version

This script extracts multi-level list content from Word documents (.docx) and converts it to JSON format.
It handles the hierarchical structure of specification documents with parts, subsections, and items.
The script includes error detection and reporting for broken or inconsistent content.

Features:
- Extracts section headers, titles, parts, subsections, items, and lists
- Handles both numbered and unnumbered structures
- Validates numbering sequences and reports errors
- Generates comprehensive error reports
- Outputs JSON in the expected format for specification documents
- Enhanced subsection detection with format normalization
- Recovery logic for broken lists and missing items
- Hierarchical content processing to avoid confusion between levels

Usage:
    python extract_spec_content_final.py <docx_file> [output_dir]

Example:
    python extract_spec_content_final.py "SECTION 26 05 00.docx"
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os
import re
import sys
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime

@dataclass
class ExtractionError:
    """Represents an error found during content extraction"""
    line_number: int
    error_type: str
    message: str
    context: str
    expected: Optional[str] = None
    found: Optional[str] = None

class SpecContentExtractor:
    """Extracts specification content from Word documents using hierarchical processing"""
    
    def __init__(self, template_path: Optional[str] = None):
        self.errors: List[ExtractionError] = []
        self.current_part_number: Optional[str] = None
        self.current_subsection_number: Optional[str] = None
        self.current_item_number: Optional[str] = None
        self.line_count = 0
        self.section_header_found = False
        self.section_title_found = False
        
        # Store all paragraphs for hierarchical analysis
        self.all_paragraphs = []
        self.paragraph_analysis = []
        
        # Template structure for validation
        self.template_structure = {}
        
        # Regex patterns for different levels
        self.part_pattern = re.compile(r'^(\d+\.0)\s+(.+)$')
        self.subsection_pattern = re.compile(r'^(\d+\.\d{2})\s+(.+)$')
        self.subsection_alt_pattern = re.compile(r'^(\d+\.\d)\s+(.+)$')  # For 1.1, 1.2 format
        self.item_pattern = re.compile(r'^([A-Z])\.\s+(.+)$')
        self.list_pattern = re.compile(r'^(\d+)\.\s+(.+)$')
        self.sub_list_pattern = re.compile(r'^([a-z])\.\s+(.+)$')
        
        # Additional patterns for different formats
        self.section_pattern = re.compile(r'^SECTION\s+(.+)$')
        self.title_pattern = re.compile(r'^([A-Z][A-Z\s]+)$')
        
        # Patterns for actual document structure
        self.part_title_pattern = re.compile(r'^([A-Z][A-Z\s]+)$')  # GENERAL, PRODUCTS, EXECUTION
        self.subsection_title_pattern = re.compile(r'^([A-Z][A-Z\s&]+)$')  # SCOPE, EXISTING CONDITIONS, etc.
        
        # Load template structure after patterns are initialized
        if template_path:
            self.load_template_structure(template_path)
    
    def load_template_structure(self, template_path: str):
        """Load and analyze template structure for validation"""
        try:
            print(f"Loading template structure from: {template_path}")
            doc = Document(template_path)
            
            # Extract all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Debug printout: show interpretation of each paragraph
            print("\n--- Template Paragraph Interpretation ---")
            for text in paragraphs:
                level_type, number, content = self.classify_paragraph_level(text)
                print(f"[{level_type}] number={number!r} content={content!r} | raw: {text!r}")
            print("--- End Template Paragraph Interpretation ---\n")
            
            # Analyze template structure
            self.template_structure = self.analyze_template_structure(paragraphs)
            print(f"Template structure loaded with {len(self.template_structure)} subsections")
            
            # Perform comprehensive template analysis
            self.comprehensive_template_analysis(template_path)
            
        except Exception as e:
            print(f"Warning: Could not load template structure: {e}")
            self.template_structure = {}
    
    def comprehensive_template_analysis(self, template_path: str):
        """Perform comprehensive analysis of template numbering structure"""
        try:
            print("\n--- Comprehensive Template Analysis ---")
            
            doc = Document(template_path)
            analysis_data = {
                "template_path": template_path,
                "analysis_timestamp": datetime.now().isoformat(),
                "paragraphs": [],
                "numbering_definitions": {},
                "level_patterns": {},
                "summary": {}
            }
            
            # 1. Analyze each paragraph for numbering information
            print("Analyzing paragraph numbering...")
            for i, para in enumerate(doc.paragraphs):
                para_data = {
                    "index": i,
                    "text": para.text.strip(),
                    "level_type": None,
                    "number": None,
                    "content": None,
                    "word_numbering": {}
                }
                
                # Get Word document numbering info
                pPr = para._p.pPr
                if pPr is not None and pPr.numPr is not None:
                    if pPr.numPr.ilvl is not None:
                        para_data["word_numbering"]["ilvl"] = pPr.numPr.ilvl.val
                    if pPr.numPr.numId is not None:
                        para_data["word_numbering"]["numId"] = pPr.numPr.numId.val
                
                # Classify the paragraph
                if para_data["text"]:
                    level_type, number, content = self.classify_paragraph_level(para_data["text"])
                    para_data["level_type"] = level_type
                    para_data["number"] = number
                    para_data["content"] = content
                    
                    if para_data["word_numbering"]:
                        print(f"PARA {i}: {para_data['text']!r} → level={level_type}, number={number}, ilvl={para_data['word_numbering'].get('ilvl')}, numId={para_data['word_numbering'].get('numId')}")
                
                analysis_data["paragraphs"].append(para_data)
            
            # 2. Extract numbering.xml information
            print("\nExtracting numbering.xml patterns...")
            try:
                import zipfile
                import lxml.etree as ET
                
                with zipfile.ZipFile(template_path) as zf:
                    if "word/numbering.xml" in zf.namelist():
                        num_xml = zf.read("word/numbering.xml")
                        root = ET.fromstring(num_xml)
                        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                        
                        # Find all abstract numbering definitions
                        for abstract_num in root.findall(".//w:abstractNum", ns):
                            abstract_num_id = abstract_num.get(qn("w:abstractNumId"))
                            analysis_data["numbering_definitions"][abstract_num_id] = {
                                "levels": {},
                                "nsid": abstract_num.get(qn("w:nsid")),
                                "multiLevelType": abstract_num.get(qn("w:multiLevelType")),
                                "tmpl": abstract_num.get(qn("w:tmpl"))
                            }
                            
                            for lvl in abstract_num.findall("w:lvl", ns):
                                ilvl = lvl.get(qn("w:ilvl"))
                                level_data = {
                                    "ilvl": ilvl,
                                    "lvlText": None,
                                    "numFmt": None,
                                    "lvlPicBulletId": None,
                                    "lvlJc": None,
                                    "pPr": {},
                                    "rPr": {},
                                    "start": None,
                                    "numFmt": None,
                                    "suff": None,
                                    "lvlRestart": None,
                                    "isLgl": None,
                                    "pStyle": None,
                                    "legacy": {},
                                    "legacyIndent": None,
                                    "legacySpace": None,
                                    "legacyNum": None
                                }
                                
                                # Extract lvlText (the pattern like "%1.0", "%1.%2")
                                lvl_text_elem = lvl.find("w:lvlText", ns)
                                if lvl_text_elem is not None:
                                    level_data["lvlText"] = lvl_text_elem.get(qn("w:val"))
                                
                                # Extract numFmt (decimal, lowerLetter, upperLetter, etc.)
                                num_fmt_elem = lvl.find("w:numFmt", ns)
                                if num_fmt_elem is not None:
                                    level_data["numFmt"] = num_fmt_elem.get(qn("w:val"))
                                
                                # Extract lvlJc (justification: left, center, right)
                                lvl_jc_elem = lvl.find("w:lvlJc", ns)
                                if lvl_jc_elem is not None:
                                    level_data["lvlJc"] = lvl_jc_elem.get(qn("w:val"))
                                
                                # Extract start value
                                start_elem = lvl.find("w:start", ns)
                                if start_elem is not None:
                                    level_data["start"] = start_elem.get(qn("w:val"))
                                
                                # Extract suffix (tab, space, nothing)
                                suff_elem = lvl.find("w:suff", ns)
                                if suff_elem is not None:
                                    level_data["suff"] = suff_elem.get(qn("w:val"))
                                
                                # Extract lvlRestart
                                lvl_restart_elem = lvl.find("w:lvlRestart", ns)
                                if lvl_restart_elem is not None:
                                    level_data["lvlRestart"] = lvl_restart_elem.get(qn("w:val"))
                                
                                # Extract isLgl (legal numbering)
                                is_lgl_elem = lvl.find("w:isLgl", ns)
                                if is_lgl_elem is not None:
                                    level_data["isLgl"] = is_lgl_elem.get(qn("w:val"))
                                
                                # Extract pStyle (paragraph style)
                                p_style_elem = lvl.find("w:pStyle", ns)
                                if p_style_elem is not None:
                                    level_data["pStyle"] = p_style_elem.get(qn("w:val"))
                                
                                # Extract paragraph properties (pPr)
                                p_pr_elem = lvl.find("w:pPr", ns)
                                if p_pr_elem is not None:
                                    # Indentation
                                    indent_elem = p_pr_elem.find("w:ind", ns)
                                    if indent_elem is not None:
                                        level_data["pPr"]["indent"] = {
                                            "left": indent_elem.get(qn("w:left")),
                                            "right": indent_elem.get(qn("w:right")),
                                            "hanging": indent_elem.get(qn("w:hanging")),
                                            "firstLine": indent_elem.get(qn("w:firstLine"))
                                        }
                                    
                                    # Spacing
                                    spacing_elem = p_pr_elem.find("w:spacing", ns)
                                    if spacing_elem is not None:
                                        level_data["pPr"]["spacing"] = {
                                            "before": spacing_elem.get(qn("w:before")),
                                            "after": spacing_elem.get(qn("w:after")),
                                            "line": spacing_elem.get(qn("w:line")),
                                            "lineRule": spacing_elem.get(qn("w:lineRule"))
                                        }
                                
                                # Extract run properties (rPr) - font info
                                r_pr_elem = lvl.find("w:rPr", ns)
                                if r_pr_elem is not None:
                                    # Font family
                                    r_fonts_elem = r_pr_elem.find("w:rFonts", ns)
                                    if r_fonts_elem is not None:
                                        level_data["rPr"]["rFonts"] = {
                                            "ascii": r_fonts_elem.get(qn("w:ascii")),
                                            "hAnsi": r_fonts_elem.get(qn("w:hAnsi")),
                                            "eastAsia": r_fonts_elem.get(qn("w:eastAsia")),
                                            "cs": r_fonts_elem.get(qn("w:cs"))
                                        }
                                    
                                    # Font size
                                    sz_elem = r_pr_elem.find("w:sz", ns)
                                    if sz_elem is not None:
                                        level_data["rPr"]["sz"] = sz_elem.get(qn("w:val"))
                                    
                                    # Bold
                                    b_elem = r_pr_elem.find("w:b", ns)
                                    if b_elem is not None:
                                        level_data["rPr"]["bold"] = b_elem.get(qn("w:val"))
                                    
                                    # Italic
                                    i_elem = r_pr_elem.find("w:i", ns)
                                    if i_elem is not None:
                                        level_data["rPr"]["italic"] = i_elem.get(qn("w:val"))
                                
                                # Extract legacy numbering info
                                legacy_elem = lvl.find("w:legacy", ns)
                                if legacy_elem is not None:
                                    level_data["legacy"] = {
                                        "legacyIndent": legacy_elem.get(qn("w:legacyIndent")),
                                        "legacySpace": legacy_elem.get(qn("w:legacySpace")),
                                        "legacyNum": legacy_elem.get(qn("w:legacyNum"))
                                    }
                                
                                analysis_data["numbering_definitions"][abstract_num_id]["levels"][ilvl] = level_data
                                
                                # Print detailed level info
                                print(f"Level {ilvl}:")
                                print(f"  Pattern: {level_data['lvlText']}")
                                print(f"  Format: {level_data['numFmt']}")
                                print(f"  Start: {level_data['start']}")
                                print(f"  Suffix: {level_data['suff']}")
                                print(f"  Justification: {level_data['lvlJc']}")
                                if level_data['pPr'].get('indent'):
                                    print(f"  Indent: {level_data['pPr']['indent']}")
                                
                                # Store in level patterns for easy access
                                analysis_data["level_patterns"][f"level_{ilvl}"] = {
                                    "ilvl": ilvl,
                                    "pattern": level_data["lvlText"],
                                    "format": level_data["numFmt"],
                                    "start": level_data["start"],
                                    "suffix": level_data["suff"],
                                    "abstract_num_id": abstract_num_id,
                                    "full_data": level_data
                                }
                        
                        # Find all num elements that map abstractNum to numId
                        print("\nExtracting num mappings...")
                        for num_elem in root.findall(".//w:num", ns):
                            num_id = num_elem.get(qn("w:numId"))
                            abstract_num_ref = num_elem.find("w:abstractNumId", ns)
                            if abstract_num_ref is not None:
                                abstract_num_id_ref = abstract_num_ref.get(qn("w:val"))
                                analysis_data["num_mappings"] = analysis_data.get("num_mappings", {})
                                analysis_data["num_mappings"][num_id] = {
                                    "abstract_num_id": abstract_num_id_ref,
                                    "abstract_num_data": analysis_data["numbering_definitions"].get(abstract_num_id_ref, {})
                                }
                                print(f"numId {num_id} → abstractNumId {abstract_num_id_ref}")
                    else:
                        print("No numbering.xml found in template")
                        
            except ImportError:
                print("lxml not available, skipping numbering.xml analysis")
            except Exception as e:
                print(f"Error reading numbering.xml: {e}")
            
            # 3. Generate summary statistics
            level_counts = {}
            numbering_counts = {}
            
            for para in analysis_data["paragraphs"]:
                if para["level_type"]:
                    level_counts[para["level_type"]] = level_counts.get(para["level_type"], 0) + 1
                
                if para["word_numbering"]:
                    ilvl = para["word_numbering"].get("ilvl")
                    if ilvl is not None:
                        numbering_counts[f"ilvl_{ilvl}"] = numbering_counts.get(f"ilvl_{ilvl}", 0) + 1
            
            analysis_data["summary"] = {
                "total_paragraphs": len(analysis_data["paragraphs"]),
                "level_counts": level_counts,
                "numbering_counts": numbering_counts,
                "numbering_definitions_count": len(analysis_data["numbering_definitions"]),
                "level_patterns_count": len(analysis_data["level_patterns"])
            }
            
            print(f"\nSummary: {analysis_data['summary']}")
            print("--- End Comprehensive Template Analysis ---\n")
            
            # Save analysis to JSON file
            analysis_file = "template_analysis.json"
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_data, f, indent=2, ensure_ascii=False)
            print(f"Template analysis saved to: {analysis_file}")
            
        except Exception as e:
            print(f"Error in comprehensive template analysis: {e}")
            import traceback
            traceback.print_exc()
    
    def extract_numbering_patterns(self, text: str) -> List[Tuple[str, str]]:
        """Extract numbering patterns from text, returning (level_type, number) pairs"""
        patterns = []
        
        # Check for various numbering patterns
        if re.match(r'^\d+\.0\s', text):
            match = re.match(r'^(\d+\.0)\s', text)
            if match:
                patterns.append(('part', match.group(1)))
                
        elif re.match(r'^\d+\.\d{2}\s', text):
            match = re.match(r'^(\d+\.\d{2})\s', text)
            if match:
                patterns.append(('subsection', match.group(1)))
                
        elif re.match(r'^\d+\.\d\s', text):
            match = re.match(r'^(\d+\.\d)\s', text)
            if match:
                patterns.append(('subsection', match.group(1)))
                
        elif re.match(r'^[A-Z]\.\s', text):
            match = re.match(r'^([A-Z])\.\s', text)
            if match:
                patterns.append(('item', match.group(1)))
                
        elif re.match(r'^\d+\.\s', text):
            match = re.match(r'^(\d+)\.\s', text)
            if match:
                patterns.append(('list', match.group(1)))
                
        elif re.match(r'^[a-z]\.\s', text):
            match = re.match(r'^([a-z])\.\s', text)
            if match:
                patterns.append(('sub_list', match.group(1)))
        
        return patterns

    def analyze_template_structure(self, paragraphs: List[str]) -> Dict[str, Dict]:
        """Analyze template to extract hierarchical levels and structure"""
        template_structure = {
            "levels": [],  # Array of level types (part, subsection, item, list, sub_list)
            "labels": [],  # Array of corresponding labels
            "hierarchy": {}  # Detailed structure for validation
        }
        
        print("\n--- Template Level Analysis ---")
        
        for text in paragraphs:
            text = text.strip()
            if not text:
                continue
                
            # Determine the level based on the label
            level_type = None
            if text.upper().startswith('PART'):
                level_type = 'part'
            elif text.upper().startswith('SUBSECTION'):
                level_type = 'subsection'
            elif text.upper().startswith('ITEM'):
                level_type = 'item'
            elif text.upper().startswith('LIST'):
                level_type = 'list'
            elif text.upper().startswith('SUBLIST'):
                level_type = 'sub_list'
            elif text.upper().startswith('SUBITEM'):
                level_type = 'sub_item'
            
            if level_type:
                template_structure["levels"].append(level_type)
                template_structure["labels"].append(text)
                print(f"Level {len(template_structure['levels'])}: {level_type} - {text}")
        
        print(f"Template has {len(template_structure['levels'])} levels")
        print("Level sequence:", template_structure["levels"])
        print("Label sequence:", template_structure["labels"])
        print("--- End Template Level Analysis ---\n")
        
        return template_structure
    
    def get_expected_structure(self, part_title: str, subsection_title: str) -> Optional[Dict]:
        """Get expected structure for a specific part/subsection combination"""
        # For now, return the overall template structure
        return self.template_structure if self.template_structure else None
    
    def validate_against_template(self, part_title: str, subsection_title: str, 
                                found_items: List[Dict]) -> List[str]:
        """Validate found items against template structure"""
        validation_errors = []
        expected_structure = self.get_expected_structure(part_title, subsection_title)
        
        if not expected_structure or not expected_structure.get("levels"):
            return validation_errors  # No template to validate against
        
        # Check if the found structure follows the expected levels
        expected_levels = expected_structure["levels"]
        print(f"Validating {part_title}:{subsection_title} against template levels: {expected_levels}")
        
        # Analyze the found structure
        found_levels = []
        for item in found_items:
            found_levels.append("item")
            if item.get("lists"):
                found_levels.append("list")
                for list_item in item["lists"]:
                    if list_item.get("sub_lists"):
                        found_levels.append("sub_list")
        
        print(f"Found levels: {found_levels}")
        
        # Check if the found levels match the expected pattern
        if found_levels and expected_levels:
            # Look for the expected pattern in the found levels
            for i, expected_level in enumerate(expected_levels):
                if i < len(found_levels):
                    if found_levels[i] != expected_level:
                        validation_errors.append(
                            f"Level mismatch at position {i+1}. "
                            f"Expected: {expected_level}, Found: {found_levels[i]}"
                        )
                else:
                    validation_errors.append(
                        f"Missing expected level: {expected_level} at position {i+1}"
                    )
        
        return validation_errors
        
    def normalize_subsection_number(self, number: str) -> str:
        """Convert subsection numbers like '1.1' to '1.01' format"""
        if re.match(r'^\d+\.\d$', number):
            parts = number.split('.')
            if len(parts) == 2:
                return f"{parts[0]}.{parts[1]:0>2}"
        return number
    
    def classify_paragraph_level(self, text: str) -> Tuple[str, Optional[str], str]:
        """
        Classify a paragraph into its hierarchical level
        Returns: (level_type, number, content)
        """
        text = text.strip()
        if not text:
            return "empty", None, ""
        
        # Check for section header (must be the very first line)
        if text.upper().startswith("SECTION") and not self.section_header_found:
            match = self.section_pattern.match(text)
            if match:
                self.section_header_found = True
                return "section", match.group(1), ""
        elif text.upper().startswith("SECTION") and self.section_header_found:
            self.add_error("Structure Error", "Multiple section headers found", text)
            return "content", None, text
        
        # Check for section title (must be the second line after section header)
        if (self.section_header_found and 
            not self.section_title_found and
            len(text.strip()) > 0):
            self.section_title_found = True
            return "title", None, text
        
        # Check for part level with numbering (1.0, 2.0, etc.)
        match = self.part_pattern.match(text)
        if match:
            return "part", match.group(1), match.group(2)
        
        # Check for part titles with various formats
        part_names = ["DESCRIPTION", "PRODUCTS", "EXECUTION", "GENERAL"]
        for part_name in part_names:
            match = re.match(rf'(?:PART\s*)?(\d+)\.0?\s*[-]?\s*{part_name}$', text.upper())
            if match:
                part_number = f"{match.group(1)}.0"
                return "part_title", part_number, part_name
            elif text.strip().upper() == part_name:
                part_number = f"{len(self.extracted_data.get('parts', [])) + 1}.0"
                return "part_title", part_number, part_name
        
        # Check for subsection level with numbering (1.01, 1.02, etc.)
        match = self.subsection_pattern.match(text)
        if match:
            normalized_number = self.normalize_subsection_number(match.group(1))
            return "subsection", normalized_number, match.group(2)
        
        # Check for subsection level with alternative numbering (1.1, 1.2, etc.)
        match = self.subsection_alt_pattern.match(text)
        if match:
            normalized_number = self.normalize_subsection_number(match.group(1))
            return "subsection", normalized_number, match.group(2)
        
        # Check for subsection titles without numbering
        subsection_titles = [
            "SCOPE", "EXISTING CONDITIONS", "CODES AND REGULATIONS", "DEFINITIONS",
            "DRAWINGS AND SPECIFICATIONS", "SITE VISIT", "DEVIATIONS",
            "STANDARDS FOR MATERIALS AND WORKMANSHIP", "SHOP DRAWINGS AND SUBMITTAL",
            "RECORD (AS-BUILT) DRAWINGS AND MAINTENANCE MANUALS",
            "COORDINATION", "PROTECTION OF MATERIALS", "TESTS, DEMONSTRATION AND INSTRUCTIONS",
            "GUARANTEE"
        ]
        
        # Check for exact match
        if text.upper() in [title.upper() for title in subsection_titles]:
            return "subsection_title", None, text
        
        # Check for partial matches
        text_upper = text.upper().strip()
        for title in subsection_titles:
            if title.upper() in text_upper or text_upper in title.upper():
                return "subsection_title", None, title
        
        # Check for item level (A., B., C., etc.)
        match = self.item_pattern.match(text)
        if match:
            return "item", match.group(1), match.group(2)
        
        # Check for list level (1., 2., etc.)
        match = self.list_pattern.match(text)
        if match:
            return "list", match.group(1), match.group(2)
        
        # Check for sub-list level (a., b., etc.)
        match = self.sub_list_pattern.match(text)
        if match:
            return "sub_list", match.group(1), match.group(2)
        
        # If no pattern matches, it's regular content
        return "content", None, text
    
    def extract_section_header_and_title(self, paragraphs: List[str]) -> Tuple[str, str]:
        """Extract section header and title from the first few paragraphs"""
        section_number = ""
        section_title = ""
        
        if len(paragraphs) >= 2:
            # First paragraph should be section header
            section_text = paragraphs[0].strip()
            if section_text.upper().startswith("SECTION"):
                section_match = re.search(r'^SECTION\s+(.+)$', section_text, re.IGNORECASE)
                if section_match:
                    section_content = section_match.group(1).strip()
                    
                    # Try to extract section number from various formats
                    number_match = re.search(r'(\d+)\s+(\d+)\s+(\d+)', section_content)
                    if number_match:
                        section_number = f"{number_match.group(1)}{number_match.group(2)}{number_match.group(3)}"
                    else:
                        number_match = re.search(r'(\d+)-(\d+)-(\d+)', section_content)
                        if number_match:
                            section_number = f"{number_match.group(1)}{number_match.group(2)}{number_match.group(3)}"
                        else:
                            number_match = re.search(r'(\d{6})', section_content)
                            if number_match:
                                section_number = number_match.group(1)
                            else:
                                section_number = section_content.replace(" ", "").replace("-", "")
            
            # Second paragraph should be section title
            title_text = paragraphs[1].strip()
            if title_text and not title_text.upper().startswith("SECTION"):
                section_title = title_text
        
        return section_number, section_title
    
    def find_part_boundaries(self, paragraphs: List[str]) -> List[Dict]:
        """Find the boundaries of parts in the document"""
        parts = []
        current_part = None
        
        for i, text in enumerate(paragraphs):
            if not text.strip():
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type in ["part", "part_title"]:
                if current_part:
                    parts.append(current_part)
                
                current_part = {
                    "start_index": i,
                    "end_index": len(paragraphs) - 1,  # Will be updated when next part is found
                    "part_number": number,
                    "title": content,
                    "level_type": level_type
                }
        
        # Add the last part
        if current_part:
            parts.append(current_part)
        
        # Update end indices
        for i in range(len(parts) - 1):
            parts[i]["end_index"] = parts[i + 1]["start_index"] - 1
        
        return parts
    
    def find_subsection_boundaries(self, paragraphs: List[str], start_index: int, end_index: int) -> List[Dict]:
        """Find the boundaries of subsections within a part"""
        subsections = []
        current_subsection = None
        
        for i in range(start_index, end_index + 1):
            text = paragraphs[i].strip()
            if not text:
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type in ["subsection", "subsection_title"]:
                if current_subsection:
                    subsections.append(current_subsection)
                
                current_subsection = {
                    "start_index": i,
                    "end_index": end_index,  # Will be updated when next subsection is found
                    "subsection_number": number,
                    "title": content,
                    "level_type": level_type
                }
        
        # Add the last subsection
        if current_subsection:
            subsections.append(current_subsection)
        
        # Update end indices
        for i in range(len(subsections) - 1):
            subsections[i]["end_index"] = subsections[i + 1]["start_index"] - 1
        
        return subsections
    
    def extract_items_and_lists(self, paragraphs: List[str], start_index: int, end_index: int) -> List[Dict]:
        """Extract items and lists from a subsection"""
        items = []
        current_item = None
        current_list_sequence = 1
        current_item_sequence = 'A'
        
        # Track all content that doesn't fit expected patterns
        unexpected_content = []
        
        for i in range(start_index, end_index + 1):
            text = paragraphs[i].strip()
            if not text:
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type == "item":
                # Check if item numbering is logical
                if number and number != current_item_sequence:
                    self.add_error(
                        "Numbering Error",
                        f"Item numbering sequence broken - expected {current_item_sequence}, found {number}",
                        text,
                        current_item_sequence,
                        number
                    )
                
                # Save previous item
                if current_item:
                    items.append(current_item)
                
                # Start new item
                current_item = {
                    "item_number": number,
                    "text": content,
                    "lists": []
                }
                current_list_sequence = 1  # Reset list sequence for new item
                
                # Update expected next item
                if current_item_sequence == 'Z':
                    current_item_sequence = 'AA'
                else:
                    current_item_sequence = chr(ord(current_item_sequence) + 1)
                    
            elif level_type == "list":
                if current_item:
                    # Check if list numbering is logical
                    if number and int(number) != current_list_sequence:
                        self.add_error(
                            "Numbering Error",
                            f"List numbering sequence broken - expected {current_list_sequence}, found {number}",
                            text,
                            str(current_list_sequence),
                            number
                        )
                    
                    current_item["lists"].append({
                        "list_number": number,
                        "text": content
                    })
                    current_list_sequence += 1
                else:
                    # List without parent item - log as error
                    self.add_error(
                        "Structure Error",
                        f"List item found without parent item",
                        text,
                        "Item A, B, C, etc.",
                        f"List {number}"
                    )
                    unexpected_content.append(text)
                    
            elif level_type == "sub_list":
                if current_item and current_item["lists"]:
                    # Add sub-list to the last list item
                    current_item["lists"][-1]["sub_lists"] = current_item["lists"][-1].get("sub_lists", [])
                    current_item["lists"][-1]["sub_lists"].append({
                        "sub_list_number": number,
                        "text": content
                    })
                else:
                    # Sub-list without parent - log as error
                    self.add_error(
                        "Structure Error",
                        f"Sub-list item found without parent list",
                        text,
                        "List 1, 2, 3, etc.",
                        f"Sub-list {number}"
                    )
                    unexpected_content.append(text)
                    
            elif level_type == "content":
                # Regular content - could be continuation of previous item or standalone text
                if current_item:
                    # Append to current item's text
                    if current_item["text"]:
                        current_item["text"] += " " + content
                    else:
                        current_item["text"] = content
                else:
                    # Standalone content - log as unexpected
                    self.add_error(
                        "Content Error",
                        f"Standalone content found without item structure",
                        text,
                        "Item A, B, C, etc.",
                        "Plain text"
                    )
                    unexpected_content.append(text)
                    
            elif level_type in ["part", "part_title", "subsection", "subsection_title"]:
                # Found structural element inside subsection - this shouldn't happen
                self.add_error(
                    "Structure Error",
                    f"Found {level_type} inside subsection - possible document structure issue",
                    text,
                    "Item or list content",
                    level_type
                )
                unexpected_content.append(text)
                
            else:
                # Any other unexpected content
                self.add_error(
                    "Content Error",
                    f"Unexpected content type: {level_type}",
                    text,
                    "Item, list, or content",
                    level_type
                )
                unexpected_content.append(text)
        
        # Add the last item
        if current_item:
            items.append(current_item)
        
        # Log any unexpected content as a block
        if unexpected_content:
            self.add_error(
                "Content Block Error",
                f"Found {len(unexpected_content)} unexpected content blocks in subsection",
                "\n".join(unexpected_content),
                "Properly structured items and lists",
                "Unexpected content"
            )
        
        return items
    
    def process_part_hierarchically(self, paragraphs: List[str], part_info: Dict) -> Dict:
        """Process a part hierarchically to extract its structure"""
        part_data = {
            "part_number": part_info["part_number"],
            "title": part_info["title"],
            "subsections": []
        }
        
        # Find subsections within this part
        subsections = self.find_subsection_boundaries(
            paragraphs, 
            part_info["start_index"], 
            part_info["end_index"]
        )
        
        for subsection_info in subsections:
            items = self.extract_items_and_lists(
                paragraphs,
                subsection_info["start_index"],
                subsection_info["end_index"]
            )
            
            # Validate against template if available
            if self.template_structure:
                validation_errors = self.validate_against_template(
                    part_data["title"], 
                    subsection_info["title"], 
                    items
                )
                for error in validation_errors:
                    self.add_error("Template Validation Error", error, subsection_info["title"])
            
            subsection_data = {
                "subsection_number": subsection_info["subsection_number"],
                "title": subsection_info["title"],
                "items": items
            }
            part_data["subsections"].append(subsection_data)
        
        return part_data
    
    def add_error(self, error_type: str, message: str, context: str = "", 
                  expected: Optional[str] = None, found: Optional[str] = None):
        """Add an error to the error list"""
        error = ExtractionError(
            line_number=self.line_count,
            error_type=error_type,
            message=message,
            context=context,
            expected=expected,
            found=found
        )
        self.errors.append(error)
    
    def extract_text_from_element(self, element, nsmap) -> str:
        """Extract text from a Word document element"""
        text = ""
        for child in element.iter():
            if child.tag == qn('w:t'):
                text += child.text or ""
        return text
    
    def get_paragraph_style(self, paragraph) -> str:
        """Get the style name of a paragraph"""
        try:
            return paragraph.style.name
        except:
            return "Normal"
    
    def get_paragraph_numbering(self, paragraph) -> Optional[str]:
        """Get the numbering level of a paragraph"""
        try:
            numbering = paragraph._p.pPr.numPr
            if numbering is not None:
                ilvl = numbering.ilvl
                if ilvl is not None:
                    return ilvl.val
        except:
            pass
        return None
    
    def parse_paragraph_content(self, text: str) -> Tuple[Optional[str], Optional[str], str]:
        """
        Parse paragraph content to extract numbering and text
        Returns: (numbering, level_type, content)
        """
        text = text.strip()
        if not text:
            return None, None, ""
        
        # Check for various numbering patterns
        patterns = [
            # Part level: 1.0, 2.0, etc.
            (r'^(\d+\.0)\s+(.+)$', 'part'),
            # Subsection level: 1.01, 1.02, etc.
            (r'^(\d+\.\d{2})\s+(.+)$', 'subsection'),
            # Alternative subsection: 1.1, 1.2, etc.
            (r'^(\d+\.\d)\s+(.+)$', 'subsection'),
            # Item level: A., B., C., etc.
            (r'^([A-Z])\.\s+(.+)$', 'item'),
            # List level: 1., 2., 3., etc.
            (r'^(\d+)\.\s+(.+)$', 'list'),
            # Sub-list level: a., b., c., etc.
            (r'^([a-z])\.\s+(.+)$', 'sub_list'),
        ]
        
        for pattern, level_type in patterns:
            match = re.match(pattern, text)
            if match:
                numbering = match.group(1)
                content = match.group(2)
                return numbering, level_type, content
        
        # No numbering found
        return None, None, text
    
    def validate_numbering_sequence(self, level_type: str, number: str, content: str):
        """Validate numbering sequences and report errors"""
        if level_type == "part":
            try:
                part_num = int(number.split('.')[0])
                if part_num != len(self.extracted_data.get('parts', [])) + 1:
                    self.add_error(
                        "Numbering Error",
                        f"Part numbering sequence broken",
                        content,
                        str(len(self.extracted_data.get('parts', [])) + 1),
                        number
                    )
            except ValueError:
                self.add_error("Numbering Error", f"Invalid part number format: {number}", content)
        
        elif level_type == "subsection":
            try:
                parts = number.split('.')
                if len(parts) == 2:
                    part_num = int(parts[0])
                    subsection_num = int(parts[1])
                    
                    # Check if this subsection belongs to the current part
                    if self.current_part_number:
                        expected_part = int(self.current_part_number.split('.')[0])
                        if part_num != expected_part:
                            self.add_error(
                                "Structure Error",
                                f"Subsection {number} does not belong to part {self.current_part_number}",
                                content
                            )
                    
                    # Check subsection numbering sequence
                    current_subsections = []
                    for part in self.extracted_data.get('parts', []):
                        if part.get('part_number') == f"{part_num}.0":
                            current_subsections = part.get('subsections', [])
                            break
                    
                    expected_subsection = len(current_subsections) + 1
                    if subsection_num != expected_subsection:
                        self.add_error(
                            "Numbering Error",
                            f"Subsection numbering sequence broken",
                            content,
                            f"{part_num}.{expected_subsection:02d}",
                            number
                        )
            except ValueError:
                self.add_error("Numbering Error", f"Invalid subsection number format: {number}", content)
        
        elif level_type == "item":
            if self.current_subsection_number:
                # Check item numbering sequence (A, B, C, ...)
                # Note: This validation would need to track current items in the subsection
                pass
        
        elif level_type == "list":
            if self.current_item_number:
                # Check list numbering sequence (1, 2, 3, ...)
                # Note: This validation would need to track current lists in the item
                pass
    
    def extract_header_info(self, doc) -> Dict[str, Any]:
        """Extract header information from the document"""
        header_info = {
            "bwa_number": "2025-XXXX",
            "client_number": "ZZZ# 00000",
            "project_name": "PROJECT NAME",
            "company_name": "CLIENT NAME",
            "section_number": "",
            "section_title": ""
        }
        
        # Try to extract from headers
        for section in doc.sections:
            if section.header:
                header_text = ""
                for paragraph in section.header.paragraphs:
                    header_text += paragraph.text + "\n"
                
                # Look for patterns in header text
                # This is a simplified extraction - you may need to customize based on your header format
                lines = header_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if "BWA" in line.upper():
                        match = re.search(r'(\d{4}-\w+)', line)
                        if match:
                            header_info["bwa_number"] = match.group(1)
                    elif "CLIENT" in line.upper() or "ZZZ" in line.upper():
                        match = re.search(r'(ZZZ#\s*\d+)', line)
                        if match:
                            header_info["client_number"] = match.group(1)
                    elif "PROJECT" in line.upper():
                        # Extract project name
                        pass
        
        return header_info
    
    def extract_content(self, docx_path: str) -> Dict[str, Any]:
        """Extract content from a Word document"""
        try:
            doc = Document(docx_path)
            
            # Extract header information
            header_info = self.extract_header_info(doc)
            
            # Extract all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
                    self.line_count += 1
            
            # Extract section header and title
            section_number, section_title = self.extract_section_header_and_title(paragraphs)
            header_info["section_number"] = section_number
            header_info["section_title"] = section_title
            
            # Initialize extracted data structure
            self.extracted_data = {
                "header": header_info,
                "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
                "margins": {
                    "top_margin": 1.0,
                    "bottom_margin": 1.0,
                    "left_margin": 1.0833333333333333,
                    "right_margin": 1.0833333333333333,
                    "header_distance": 1.0,
                    "footer_distance": 1.0
                },
                "comments": [],
                "section_number": section_number,
                "section_title": section_title,
                "parts": []
            }
            
            # Find part boundaries
            parts = self.find_part_boundaries(paragraphs)
            
            # Process each part hierarchically
            for part_info in parts:
                part_data = self.process_part_hierarchically(paragraphs, part_info)
                self.extracted_data["parts"].append(part_data)
            
            return self.extracted_data
            
        except Exception as e:
            self.add_error("Extraction Error", f"Failed to extract content: {str(e)}", "")
            return {}
    
    def generate_error_report(self) -> str:
        """Generate a comprehensive error report"""
        if not self.errors:
            return "No errors found during extraction.\n"
        
        report = f"ERROR REPORT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += "=" * 60 + "\n\n"
        
        # Group errors by type
        error_types = {}
        for error in self.errors:
            if error.error_type not in error_types:
                error_types[error.error_type] = []
            error_types[error.error_type].append(error)
        
        for error_type, errors in error_types.items():
            report += f"{error_type} ERRORS ({len(errors)} found):\n"
            report += "-" * 40 + "\n"
            
            for i, error in enumerate(errors, 1):
                report += f"{i}. Line {error.line_number}: {error.message}\n"
                if error.context:
                    # For content block errors, format the context nicely
                    if error.error_type == "Content Block Error" and "\n" in error.context:
                        report += f"   Content Block:\n"
                        lines = error.context.split('\n')
                        for line in lines:
                            if line.strip():
                                report += f"   > {line.strip()}\n"
                    else:
                        report += f"   Context: {error.context}\n"
                if error.expected and error.found:
                    report += f"   Expected: {error.expected}, Found: {error.found}\n"
                report += "\n"
        
        # Add summary statistics
        report += "SUMMARY:\n"
        report += "-" * 20 + "\n"
        total_errors = len(self.errors)
        report += f"Total errors: {total_errors}\n"
        for error_type, errors in error_types.items():
            report += f"{error_type}: {len(errors)} errors\n"
        
        return report
    
    def save_to_json(self, data: Dict[str, Any], output_path: str):
        """Save extracted data to JSON file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def save_error_report(self, report: str, output_path: str):
        """Save error report to text file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)

def main():
    """Main function to run the extraction"""
    if len(sys.argv) < 2:
        print("Usage: python extract_spec_content_final.py <docx_file> [output_dir] [template_file]")
        print("Example: python extract_spec_content_final.py 'SECTION 26 05 00.docx' . 'test_template.docx'")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "."
    template_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not os.path.exists(docx_path):
        print(f"Error: File '{docx_path}' not found.")
        sys.exit(1)
    
    if template_path and not os.path.exists(template_path):
        print(f"Warning: Template file '{template_path}' not found. Proceeding without template validation.")
        template_path = None
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Initialize extractor with template if provided
    extractor = SpecContentExtractor(template_path)
    
    # Extract content
    print(f"Extracting content from '{docx_path}'...")
    data = extractor.extract_content(docx_path)
    
    if not data:
        print("Error: Failed to extract content.")
        sys.exit(1)
    
    # Generate output filenames
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    json_path = os.path.join(output_dir, f"{base_name}.json")
    error_path = os.path.join(output_dir, f"{base_name}_errors.txt")
    
    # Save results
    extractor.save_to_json(data, json_path)
    print(f"Content saved to: {json_path}")
    
    # Generate and save error report
    error_report = extractor.generate_error_report()
    extractor.save_error_report(error_report, error_path)
    print(f"Error report saved to: {error_path}")
    
    # Print summary
    print(f"\nExtraction Summary:")
    print(f"- Parts found: {len(data.get('parts', []))}")
    total_subsections = sum(len(part.get('subsections', [])) for part in data.get('parts', []))
    print(f"- Subsections found: {total_subsections}")
    total_items = sum(
        len(subsection.get('items', [])) 
        for part in data.get('parts', []) 
        for subsection in part.get('subsections', [])
    )
    print(f"- Items found: {total_items}")
    print(f"- Errors found: {len(extractor.errors)}")
    
    if extractor.errors:
        print(f"\nWARNING: {len(extractor.errors)} errors were found during extraction.")
        print(f"Please review the error report: {error_path}")
    else:
        print("\nExtraction completed successfully with no errors.")

if __name__ == "__main__":
    main() 