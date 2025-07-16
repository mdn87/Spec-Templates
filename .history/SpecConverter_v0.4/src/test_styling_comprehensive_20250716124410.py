#!/usr/bin/env python3
"""
Comprehensive styling test to verify enhanced formatting comes through
"""

from docx import Document
import json
import os
from extract_spec_content_v3 import SpecContentExtractorV3

def test_comprehensive_styling():
    """Test comprehensive styling extraction and regeneration"""
    print("=== COMPREHENSIVE STYLING TEST ===")
    
    # Test extraction
    print("\n1. Testing extraction with enhanced template...")
    extractor = SpecContentExtractorV3("../templates/test_template_cleaned.docx")
    
    # Extract content
    data = extractor.extract_content("../examples/SECTION 26 05 00.docx")
    
    if not data:
        print("ERROR: Extraction failed")
        return
    
    content_blocks = data.get('content_blocks', [])
    print(f"✓ Extracted {len(content_blocks)} content blocks")
    
    # Analyze styling extraction
    print("\n2. Analyzing extracted styling...")
    
    styling_stats = {
        'font_size': 0,
        'font_bold': 0,
        'font_italic': 0,
        'font_underline': 0,
        'font_color': 0,
        'paragraph_indent_left': 0,
        'paragraph_indent_right': 0,
        'paragraph_indent_first_line': 0,
        'paragraph_spacing_before': 0,
        'paragraph_spacing_after': 0,
        'paragraph_line_spacing': 0,
        'paragraph_alignment': 0,
        'number_alignment': 0,
        'follow_number_with': 0,
        'aligned_at': 0,
        'text_indent_at': 0,
        'add_tab_stop_at': 0,
        'link_level_to_style': 0
    }
    
    # Count styling properties
    for block in content_blocks:
        for prop in styling_stats.keys():
            if block.get(prop) is not None:
                styling_stats[prop] += 1
    
    print("Extracted styling properties:")
    for prop, count in styling_stats.items():
        if count > 0:
            print(f"  ✓ {prop}: {count} instances")
        else:
            print(f"  ✗ {prop}: {count} instances")
    
    # Check for specific styling examples
    print("\n3. Looking for specific styling examples...")
    
    # Find paragraphs with font color
    colored_paragraphs = [b for b in content_blocks if b.get('font_color')]
    if colored_paragraphs:
        print(f"✓ Found {len(colored_paragraphs)} paragraphs with font color")
        for i, block in enumerate(colored_paragraphs[:3]):  # Show first 3
            print(f"  - Color: {block['font_color']}, Text: {block['text'][:50]}...")
    
    # Find paragraphs with indentation
    indented_paragraphs = [b for b in content_blocks if b.get('paragraph_indent_left')]
    if indented_paragraphs:
        print(f"✓ Found {len(indented_paragraphs)} paragraphs with left indentation")
        indent_levels = set(b['paragraph_indent_left'] for b in indented_paragraphs)
        print(f"  - Indentation levels: {sorted(indent_levels)}")
    
    # Find paragraphs with alignment
    aligned_paragraphs = [b for b in content_blocks if b.get('paragraph_alignment')]
    if aligned_paragraphs:
        print(f"✓ Found {len(aligned_paragraphs)} paragraphs with alignment")
        alignments = set(b['paragraph_alignment'] for b in aligned_paragraphs)
        print(f"  - Alignment types: {alignments}")
    
    # Find numbered paragraphs with enhanced properties
    numbered_paragraphs = [b for b in content_blocks if b.get('number_alignment')]
    if numbered_paragraphs:
        print(f"✓ Found {len(numbered_paragraphs)} numbered paragraphs with enhanced properties")
        for i, block in enumerate(numbered_paragraphs[:3]):  # Show first 3
            print(f"  - Number alignment: {block['number_alignment']}, Follow with: {block['follow_number_with']}")
    
    # Test regeneration
    print("\n4. Testing regeneration...")
    
    # Save test JSON
    test_json_path = "../output/test_comprehensive_styling.json"
    with open(test_json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    # Import and run test generator
    from test_generator import Document, parse_spec_json, generate_content_from_v3_json, apply_document_settings_from_json, apply_margins_from_json
    
    # Create new document
    doc = Document("../templates/test_template_cleaned.docx")
    
    # Clear content
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    # Apply settings and generate content
    apply_document_settings_from_json(doc, data)
    apply_margins_from_json(doc, data)
    generate_content_from_v3_json(doc, data)
    
    # Save regenerated document
    output_path = "../output/test_comprehensive_styling_output.docx"
    doc.save(output_path)
    
    print(f"✓ Regenerated document saved to: {output_path}")
    
    # Verify the regenerated document has content
    regenerated_doc = Document(output_path)
    regenerated_paragraphs = len(regenerated_doc.paragraphs)
    print(f"✓ Regenerated document has {regenerated_paragraphs} paragraphs")
    
    # Check if BWA styles were applied
    bwa_styles_found = 0
    for paragraph in regenerated_doc.paragraphs:
        if paragraph.style and hasattr(paragraph.style, 'name') and paragraph.style.name and paragraph.style.name.startswith('BWA-'):
            bwa_styles_found += 1
    
    print(f"✓ Found {bwa_styles_found} paragraphs with BWA styles")
    
    print("\n=== STYLING TEST COMPLETE ===")
    print("Check the regenerated document to verify styling is applied correctly.")
    print("The document should show:")
    print("- Proper BWA styles applied to different level types")
    print("- Correct indentation levels")
    print("- Font colors and formatting where specified")
    print("- Proper paragraph alignment")
    print("- Enhanced numbering properties")

if __name__ == "__main__":
    test_comprehensive_styling() 