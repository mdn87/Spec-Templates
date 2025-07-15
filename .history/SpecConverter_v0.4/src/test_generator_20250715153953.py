from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import os

# Configuration variables - change these to modify font and size for all text
TEMPLATE_PATH = '../templates/test_template_cleaned.docx'
OUTPUT_PATH   = '../output/generated_spec_v3.docx'
CONTENT_PATH  = '../output/SECTION 26 05 00_v3.json'
FONT_NAME = 'Arial'
FONT_SIZE = 10

def list_available_styles(doc):
    """List all available styles in the document"""
    print("DEBUG: Available styles in template:")
    for style in doc.styles:
        print(f"  - {style.name}")

def clear_document(doc):
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

def set_font_and_size(paragraph):
    """Set font and size for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        # For compatibility with some versions of Word
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def parse_spec_json(json_path):
    """Parse JSON file and return structured data"""
    print(f"DEBUG: Attempting to read JSON file: {json_path}")
    
    if not os.path.exists(json_path):
        print(f"ERROR: JSON file not found: {json_path}")
        return None
    
    try:
        with open(json_path, 'r', encoding='utf-8') as file:
            content = json.load(file)
            print(f"DEBUG: Successfully loaded JSON with content_blocks: {len(content.get('content_blocks', []))}")
            print(f"DEBUG: JSON content type: {type(content)}")
            print(f"DEBUG: JSON keys: {list(content.keys())}")
            return content
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON format: {e}")
        return None
    except Exception as e:
        print(f"ERROR: Failed to read JSON file: {e}")
        return None

def get_style_for_bwa_level(bwa_level_name):
    """Map BWA level names to template style names"""
    style_mapping = {
        "BWA-PART": "BWA-PART",
        "BWA-SUBSECTION": "BWA-SUBSECTION", 
        "BWA-Item": "BWA-Item",
        "BWA-List": "BWA-List",
        "BWA-SubList": "BWA-SubList",
        "BWA-SubItem": "BWA-SubItem",
        "BWA-SubSubItem": "BWA-SubSubItem",
        "BWA-SubSubList": "BWA-SubSubList"
    }
    return style_mapping.get(bwa_level_name, "Normal")

def generate_content_from_v3_json(doc, json_data):
    """Generate document content from v3 JSON data using template styles"""
    print(f"DEBUG: Starting content generation from v3 JSON data")
    
    if json_data is None:
        print("ERROR: No JSON data to process")
        return
    
    content_blocks = json_data.get('content_blocks', [])
    print(f"DEBUG: Processing {len(content_blocks)} content blocks")
    
    for i, block in enumerate(content_blocks):
        text = block.get('text', '')
        level_type = block.get('level_type', 'content')
        bwa_level_name = block.get('bwa_level_name')
        number = block.get('number')
        level_number = block.get('level_number')
        
        print(f"DEBUG: Block {i+1}: {level_type} -> {bwa_level_name} (level {level_number})")
        
        # Skip empty content
        if not text.strip():
            continue
        
        # Determine the style to use
        if bwa_level_name:
            style_name = get_style_for_bwa_level(bwa_level_name)
            print(f"DEBUG: Using BWA style: {style_name}")
        else:
            # Fallback styles based on level_type
            if level_type == "section":
                style_name = "Heading 1"
            elif level_type == "title":
                style_name = "Heading 2"
            elif level_type == "part_title":
                style_name = "Heading 3"
            elif level_type == "subsection_title":
                style_name = "Heading 4"
            else:
                style_name = "Normal"
            print(f"DEBUG: Using fallback style: {style_name}")
        
        # Add paragraph with appropriate style
        try:
            paragraph = doc.add_paragraph(text, style=style_name)
            set_font_and_size(paragraph)
            print(f"DEBUG: Added paragraph with style '{style_name}': {text[:50]}...")
        except Exception as e:
            print(f"DEBUG: Style '{style_name}' not found, using Normal: {e}")
            paragraph = doc.add_paragraph(text, style="Normal")
            set_font_and_size(paragraph)
        
        # Add some spacing for better readability
        if level_type in ["section", "title", "part_title"]:
            doc.add_paragraph()  # Add blank line after major sections

# Main execution
print("DEBUG: Starting document generation process")

# Load template
try:
    doc = Document(TEMPLATE_PATH)
    print(f"DEBUG: Successfully loaded template: {TEMPLATE_PATH}")
    
    # List available styles to see what's in the template
    list_available_styles(doc)
    
except ValueError as e:
    print(f"DEBUG: Template file error, creating new document: {e}")
    doc = Document()

# Clear existing content
clear_document(doc)
print("DEBUG: Cleared existing document content")

# Parse JSON content
json_data = parse_spec_json(CONTENT_PATH)

# Generate content from JSON
generate_content_from_v3_json(doc, json_data)

# Save document
doc.save(OUTPUT_PATH)
print(f"DEBUG: Document saved as '{OUTPUT_PATH}' with {FONT_SIZE}pt {FONT_NAME} font")
print(f"DEBUG: Content source: {CONTENT_PATH}")
print("Note: The template's multilevel list style should be applied automatically")
print("if the paragraphs use the correct style names from the template.")






