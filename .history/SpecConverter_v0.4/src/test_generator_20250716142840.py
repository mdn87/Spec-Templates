from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
import json
import os
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE

# Configuration variables - change these to modify font and size for all text
TEMPLATE_PATH = '../templates/test_template_cleaned.docx'
OUTPUT_PATH   = '../output/generated_spec_v3_fixed.docx'
CONTENT_PATH  = '../output/SECTION 26 05 00_v3.json'
FONT_NAME = 'Arial'
FONT_SIZE = 10

def list_available_styles(doc):
    """List all available styles in the document"""
    print("DEBUG: Available styles in template:")
    for style in doc.styles:
        print(f"  - {style.name}")

def clear_document(doc):
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

def set_font_and_size(paragraph):
    """Set font and size for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        # For compatibility with some versions of Word
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def apply_styling_from_json(paragraph, block):
    """Apply styling information from JSON block to paragraph"""
    try:
        # Font properties
        if block.get('font_name'):
            for run in paragraph.runs:
                run.font.name = block['font_name']
                # For compatibility with some versions of Word
                r = run._element
                if r.rPr is None:
                    r.rPr = OxmlElement('w:rPr')
                r_fonts = r.rPr.find(qn('w:rFonts'))
                if r_fonts is None:
                    r_fonts = OxmlElement('w:rFonts')
                    r.rPr.append(r_fonts)
                r_fonts.set(qn('w:eastAsia'), block['font_name'])
        
        if block.get('font_size'):
            for run in paragraph.runs:
                run.font.size = Pt(block['font_size'])
        
        if block.get('font_bold') is not None:
            for run in paragraph.runs:
                run.font.bold = block['font_bold']
        
        if block.get('font_italic') is not None:
            for run in paragraph.runs:
                run.font.italic = block['font_italic']
        
        if block.get('font_underline'):
            for run in paragraph.runs:
                # Map underline values to valid WD_UNDERLINE constants
                underline_map = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wavy': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE
                }
                underline_value = underline_map.get(block['font_underline'], WD_UNDERLINE.SINGLE)
                run.font.underline = underline_value
        
        if block.get('font_color'):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(block['font_color'])
        
        # Additional font properties
        if block.get('font_strike') is not None:
            for run in paragraph.runs:
                run.font.strike = block['font_strike']
        
        if block.get('font_small_caps') is not None:
            for run in paragraph.runs:
                run.font.small_caps = block['font_small_caps']
        
        if block.get('font_all_caps') is not None:
            for run in paragraph.runs:
                run.font.all_caps = block['font_all_caps']
        
        # Paragraph properties
        if block.get('paragraph_alignment'):
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'both': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if block['paragraph_alignment'] in alignment_map:
                paragraph.alignment = alignment_map[block['paragraph_alignment']]
        
        # Indentation (convert points to inches for Word)
        if block.get('paragraph_indent_left'):
            paragraph.paragraph_format.left_indent = Inches(block['paragraph_indent_left'] / 72.0)
        
        if block.get('paragraph_indent_right'):
            paragraph.paragraph_format.right_indent = Inches(block['paragraph_indent_right'] / 72.0)
        
        if block.get('paragraph_indent_first_line'):
            paragraph.paragraph_format.first_line_indent = Inches(block['paragraph_indent_first_line'] / 72.0)
        
        # Spacing
        if block.get('paragraph_spacing_before'):
            paragraph.paragraph_format.space_before = Pt(block['paragraph_spacing_before'])
        
        if block.get('paragraph_spacing_after'):
            paragraph.paragraph_format.space_after = Pt(block['paragraph_spacing_after'])
        
        if block.get('paragraph_line_spacing'):
            paragraph.paragraph_format.line_spacing = block['paragraph_line_spacing']
        
        # Additional paragraph properties
        if block.get('paragraph_keep_with_next') is not None:
            paragraph.paragraph_format.keep_with_next = block['paragraph_keep_with_next']
        
        if block.get('paragraph_keep_lines_together') is not None:
            paragraph.paragraph_format.keep_lines_together = block['paragraph_keep_lines_together']
        
        if block.get('paragraph_page_break_before') is not None:
            paragraph.paragraph_format.page_break_before = block['paragraph_page_break_before']
        
        if block.get('paragraph_widow_control') is not None:
            paragraph.paragraph_format.widow_control = block['paragraph_widow_control']
        
        # Don't add space between paragraphs of the same style
        if block.get('paragraph_dont_add_space_between_same_style') is not None:
            paragraph.paragraph_format.dont_add_space_between_same_style = block['paragraph_dont_add_space_between_same_style']
        
    except Exception as e:
        print(f"Warning: Could not apply styling from JSON: {e}")
        # Fallback to default styling
        set_font_and_size(paragraph)

def apply_style_definitions_from_json(doc, json_data):
    """Apply style definitions from JSON to ensure proper styling in regenerated document"""
    try:
        content_blocks = json_data.get('content_blocks', [])
        if not content_blocks:
            return
        
        # Collect all unique BWA level names used in content blocks
        bwa_style_names = set()
        for block in content_blocks:
            bwa_level_name = block.get('bwa_level_name')
            if bwa_level_name and bwa_level_name != 'Normal':
                bwa_style_names.add(bwa_level_name)
        
        # Apply styling for each BWA style
        for bwa_style_name in bwa_style_names:
            apply_style_definition(doc, bwa_style_name, content_blocks)
        
        print(f"Applied style definitions for {len(bwa_style_names)} BWA styles")
        
    except Exception as e:
        print(f"Warning: Could not apply style definitions: {e}")

def apply_style_definition(doc, style_name, content_blocks):
    """Apply definition for a specific style based on content blocks using that style"""
    try:
        # Get or create the style
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            # Create new style if it doesn't exist
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        # Find content blocks using this BWA level name to determine its properties
        style_blocks = [block for block in content_blocks if block.get('bwa_level_name') == style_name]
        if not style_blocks:
            return
        
        # Use the first block to determine style properties (they should be consistent)
        sample_block = style_blocks[0]
        
        # Apply paragraph format properties
        if hasattr(style, 'paragraph_format') and style.paragraph_format:
            pf = style.paragraph_format
            
            # Alignment
            if sample_block.get('paragraph_alignment'):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'both': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if sample_block['paragraph_alignment'] in alignment_map:
                    pf.alignment = alignment_map[sample_block['paragraph_alignment']]
            
            # Indentation
            if sample_block.get('paragraph_indent_left'):
                from docx.shared import Inches
                pf.left_indent = Inches(sample_block['paragraph_indent_left'] / 72.0)
            
            if sample_block.get('paragraph_indent_right'):
                from docx.shared import Inches
                pf.right_indent = Inches(sample_block['paragraph_indent_right'] / 72.0)
            
            if sample_block.get('paragraph_indent_first_line'):
                from docx.shared import Inches
                pf.first_line_indent = Inches(sample_block['paragraph_indent_first_line'] / 72.0)
            
            # Spacing
            if sample_block.get('paragraph_spacing_before'):
                from docx.shared import Pt
                pf.space_before = Pt(sample_block['paragraph_spacing_before'])
            
            if sample_block.get('paragraph_spacing_after'):
                from docx.shared import Pt
                pf.space_after = Pt(sample_block['paragraph_spacing_after'])
            
            if sample_block.get('paragraph_line_spacing'):
                pf.line_spacing = sample_block['paragraph_line_spacing']
            
            # Other paragraph properties
            if sample_block.get('paragraph_keep_with_next') is not None:
                pf.keep_with_next = sample_block['paragraph_keep_with_next']
            
            if sample_block.get('paragraph_keep_lines_together') is not None:
                pf.keep_lines_together = sample_block['paragraph_keep_lines_together']
            
            if sample_block.get('paragraph_page_break_before') is not None:
                pf.page_break_before = sample_block['paragraph_page_break_before']
            
            if sample_block.get('paragraph_widow_control') is not None:
                pf.widow_control = sample_block['paragraph_widow_control']
        
        # Apply font properties
        if hasattr(style, 'font') and style.font:
            font = style.font
            
            if sample_block.get('font_name'):
                font.name = sample_block['font_name']
            
            if sample_block.get('font_size'):
                from docx.shared import Pt
                font.size = Pt(sample_block['font_size'])
            
            if sample_block.get('font_bold') is not None:
                font.bold = sample_block['font_bold']
            
            if sample_block.get('font_italic') is not None:
                font.italic = sample_block['font_italic']
            
            if sample_block.get('font_underline'):
                from docx.enum.text import WD_UNDERLINE
                underline_map = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wavy': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE
                }
                if sample_block['font_underline'] in underline_map:
                    font.underline = underline_map[sample_block['font_underline']]
            
            if sample_block.get('font_color'):
                from docx.shared import RGBColor
                font.color.rgb = RGBColor.from_string(sample_block['font_color'])
            
            if sample_block.get('font_strike') is not None:
                font.strike = sample_block['font_strike']
            
            if sample_block.get('font_small_caps') is not None:
                font.small_caps = sample_block['font_small_caps']
            
            if sample_block.get('font_all_caps') is not None:
                font.all_caps = sample_block['font_all_caps']
        
    except Exception as e:
        print(f"Warning: Could not apply style definition for {style_name}: {e}")

def apply_document_settings_from_json(doc, json_data):
    """Apply document-level settings from JSON to the document"""
    try:
        document_settings = json_data.get('document_settings', {})
        if not document_settings:
            return
        
        # Apply settings to the first section
        section = doc.sections[0]
        
        # Page size and orientation
        if document_settings.get('page_width') and document_settings.get('page_height'):
            from docx.shared import Inches
            section.page_width = Inches(document_settings['page_width'])
            section.page_height = Inches(document_settings['page_height'])
        
        # Margins (if not already set by template)
        if document_settings.get('top_margin'):
            section.top_margin = Inches(document_settings['top_margin'])
        if document_settings.get('bottom_margin'):
            section.bottom_margin = Inches(document_settings['bottom_margin'])
        if document_settings.get('left_margin'):
            section.left_margin = Inches(document_settings['left_margin'])
        if document_settings.get('right_margin'):
            section.right_margin = Inches(document_settings['right_margin'])
        
        # Header and footer distances
        if document_settings.get('header_distance'):
            section.header_distance = Inches(document_settings['header_distance'])
        if document_settings.get('footer_distance'):
            section.footer_distance = Inches(document_settings['footer_distance'])
        
        # Gutter settings
        if document_settings.get('gutter'):
            section.gutter = Inches(document_settings['gutter'])
        
        # Different first page header/footer
        if document_settings.get('different_first_page_header_footer') is not None:
            section.different_first_page_header_footer = document_settings['different_first_page_header_footer']
        
        # Different odd and even pages
        if document_settings.get('different_odd_and_even_pages') is not None:
            section.different_odd_and_even_pages = document_settings['different_odd_and_even_pages']
        
        # Page numbering
        if document_settings.get('page_numbering'):
            page_num = document_settings['page_numbering']
            if hasattr(section, 'page_numbering') and section.page_numbering:
                if page_num.get('start') is not None:
                    section.page_numbering.start = page_num['start']
                if page_num.get('restart') is not None:
                    section.page_numbering.restart = page_num['restart']
                if page_num.get('format') is not None:
                    section.page_numbering.format = page_num['format']
        
        # Line numbering
        if document_settings.get('line_numbering'):
            line_num = document_settings['line_numbering']
            if hasattr(section, 'line_numbering') and section.line_numbering:
                if line_num.get('start') is not None:
                    section.line_numbering.start = line_num['start']
                if line_num.get('increment') is not None:
                    section.line_numbering.increment = line_num['increment']
                if line_num.get('restart') is not None:
                    section.line_numbering.restart = line_num['restart']
                if line_num.get('distance') is not None:
                    section.line_numbering.distance = Inches(line_num['distance'])
        
        # Document properties
        if document_settings.get('document_properties') and doc.core_properties:
            props = document_settings['document_properties']
            if props.get('title'):
                doc.core_properties.title = props['title']
            if props.get('subject'):
                doc.core_properties.subject = props['subject']
            if props.get('author'):
                doc.core_properties.author = props['author']
            if props.get('keywords'):
                doc.core_properties.keywords = props['keywords']
            if props.get('category'):
                doc.core_properties.category = props['category']
            if props.get('comments'):
                doc.core_properties.comments = props['comments']
            if props.get('last_modified_by'):
                doc.core_properties.last_modified_by = props['last_modified_by']
            if props.get('revision'):
                doc.core_properties.revision = props['revision']
        
        # Apply default formatting settings
        if document_settings.get('default_formatting'):
            apply_default_formatting_from_json(doc, document_settings['default_formatting'])
        
        # Apply document-wide settings
        if document_settings.get('document_wide_settings'):
            apply_document_wide_settings_from_json(doc, document_settings['document_wide_settings'])
        
    except Exception as e:
        print(f"Warning: Could not apply document settings from JSON: {e}")

def apply_default_formatting_from_json(doc, default_formatting):
    """Apply default formatting settings from JSON to document styles"""
    try:
        # Get the Normal style
        normal_style = doc.styles['Normal'] if 'Normal' in doc.styles else None
        if not normal_style:
            return
        
        # Apply default paragraph format
        if default_formatting.get('default_paragraph_format'):
            pf_data = default_formatting['default_paragraph_format']
            pf = normal_style.paragraph_format
            
            if pf_data.get('alignment'):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if pf_data['alignment'] in alignment_map:
                    pf.alignment = alignment_map[pf_data['alignment']]
            
            if pf_data.get('left_indent'):
                from docx.shared import Inches
                pf.left_indent = Inches(pf_data['left_indent'])
            
            if pf_data.get('right_indent'):
                from docx.shared import Inches
                pf.right_indent = Inches(pf_data['right_indent'])
            
            if pf_data.get('first_line_indent'):
                from docx.shared import Inches
                pf.first_line_indent = Inches(pf_data['first_line_indent'])
            
            if pf_data.get('space_before'):
                from docx.shared import Pt
                pf.space_before = Pt(pf_data['space_before'])
            
            if pf_data.get('space_after'):
                from docx.shared import Pt
                pf.space_after = Pt(pf_data['space_after'])
            
            if pf_data.get('line_spacing'):
                pf.line_spacing = pf_data['line_spacing']
            
            if pf_data.get('keep_with_next') is not None:
                pf.keep_with_next = pf_data['keep_with_next']
            
            if pf_data.get('keep_lines_together') is not None:
                pf.keep_lines_together = pf_data['keep_lines_together']
            
            if pf_data.get('page_break_before') is not None:
                pf.page_break_before = pf_data['page_break_before']
            
            if pf_data.get('widow_control') is not None:
                pf.widow_control = pf_data['widow_control']
        
        # Apply default run format
        if default_formatting.get('default_run_format'):
            font_data = default_formatting['default_run_format']
            font = normal_style.font
            
            if font_data.get('name'):
                font.name = font_data['name']
            
            if font_data.get('size'):
                from docx.shared import Pt
                font.size = Pt(font_data['size'])
            
            if font_data.get('bold') is not None:
                font.bold = font_data['bold']
            
            if font_data.get('italic') is not None:
                font.italic = font_data['italic']
            
            if font_data.get('underline'):
                from docx.enum.text import WD_UNDERLINE
                underline_map = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wavy': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE
                }
                if font_data['underline'] in underline_map:
                    font.underline = underline_map[font_data['underline']]
            
            if font_data.get('color'):
                from docx.shared import RGBColor
                font.color.rgb = RGBColor.from_string(font_data['color'])
            
            if font_data.get('strike') is not None:
                font.strike = font_data['strike']
            
            if font_data.get('small_caps') is not None:
                font.small_caps = font_data['small_caps']
            
            if font_data.get('all_caps') is not None:
                font.all_caps = font_data['all_caps']
        
    except Exception as e:
        print(f"Warning: Could not apply default formatting from JSON: {e}")

def apply_document_wide_settings_from_json(doc, doc_wide_settings):
    """Apply document-wide settings from JSON to document"""
    try:
        # Note: Most document-wide settings are read-only or require direct XML manipulation
        # We'll focus on settings that can be applied through the python-docx API
        
        # Default tab stop can be applied through styles
        if doc_wide_settings.get('default_tab_stop'):
            # Apply to Normal style paragraph format
            normal_style = doc.styles.get('Normal')
            if normal_style and normal_style.paragraph_format:
                # Convert twips to inches (1 inch = 1440 twips)
                tab_stop_inches = float(doc_wide_settings['default_tab_stop']) / 1440.0
                from docx.shared import Inches
                normal_style.paragraph_format.tab_stops.add_tab_stop(Inches(tab_stop_inches))
        
        # Track changes setting
        if doc_wide_settings.get('track_changes') is not None:
            # This would require direct XML manipulation
            print(f"DEBUG: Track changes setting found: {doc_wide_settings['track_changes']} (requires XML manipulation)")
        
        # Document protection
        if doc_wide_settings.get('document_protection'):
            protection = doc_wide_settings['document_protection']
            print(f"DEBUG: Document protection settings found: {protection} (requires XML manipulation)")
        
    except Exception as e:
        print(f"Warning: Could not apply document-wide settings from JSON: {e}")

def apply_margins_from_json(doc, json_data):
    """Apply margin settings from JSON to the document"""
    try:
        margins = json_data.get('margins', {})
        if not margins:
            return
        
        # Apply margins to the first section
        section = doc.sections[0]
        
        if margins.get('top_margin'):
            section.top_margin = Inches(margins['top_margin'])
        if margins.get('bottom_margin'):
            section.bottom_margin = Inches(margins['bottom_margin'])
        if margins.get('left_margin'):
            section.left_margin = Inches(margins['left_margin'])
        if margins.get('right_margin'):
            section.right_margin = Inches(margins['right_margin'])
        if margins.get('header_distance'):
            section.header_distance = Inches(margins['header_distance'])
        if margins.get('footer_distance'):
            section.footer_distance = Inches(margins['footer_distance'])
        
    except Exception as e:
        print(f"Warning: Could not apply margins from JSON: {e}")

def parse_spec_json(json_path):
    """Parse JSON file and return structured data"""
    
    if not os.path.exists(json_path):
        print(f"ERROR: JSON file not found: {json_path}")
        return None
    
    try:
        with open(json_path, 'r', encoding='utf-8') as file:
            content = json.load(file)
            return content
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON format: {e}")
        return None
    except Exception as e:
        print(f"ERROR: Failed to read JSON file: {e}")
        return None

def get_style_for_bwa_level(bwa_level_name):
    """Map BWA level names to template style names"""
    style_mapping = {
        "BWA-SectionNumber": "BWA-SectionNumber",
        "BWA-SectionTitle": "BWA-SectionTitle",
        "BWA-PART": "BWA-PART",
        "BWA-SUBSECTION": "BWA-SUBSECTION", 
        "BWA-Item": "BWA-Item",
        "BWA-List": "BWA-List",
        "BWA-SubList": "BWA-SubList",
        "BWA-SubItem": "BWA-SubItem",
        "BWA-SubSubItem": "BWA-SubSubItem",
        "BWA-SubSubList": "BWA-SubSubList"
    }
    return style_mapping.get(bwa_level_name, "Normal")

def clean_text_for_display(text, level_type, number):
    """Clean text by removing numbering prefixes while preserving content"""
    if not text:
        return text
    
    # Remove common numbering patterns from the beginning of text
    import re
    
    # Patterns to remove (in order of specificity)
    patterns = [
        # Item patterns: "A.\t", "B.\t", "C.\t", etc.
        r'^[A-Z]\.\s*\t\s*',
        # List patterns: "1.\t", "2.\t", "3.\t", etc.
        r'^\d+\.\s*\t\s*',
        # Sub-list patterns: "a.\t", "b.\t", "c.\t", etc.
        r'^[a-z]\.\s*\t\s*',
        # Part patterns: "1.0\t", "2.0\t", etc.
        r'^\d+\.0\s*\t\s*',
        # Subsection patterns: "1.01\t", "1.02\t", etc.
        r'^\d+\.\d{2}\s*\t\s*',
        # Alternative subsection patterns: "1.1\t", "1.2\t", etc.
        r'^\d+\.\d\s*\t\s*',
        # Section patterns: "SECTION 26 05 00\t"
        r'^SECTION\s+[^\t]*\s*\t\s*',
        # Generic tab removal at start
        r'^\s*\t\s*'
    ]
    
    cleaned_text = text
    for pattern in patterns:
        cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE)
    
    # If we removed everything, return the original text
    if not cleaned_text.strip():
        return text
    
    return cleaned_text.strip()

def generate_content_from_v3_json(doc, json_data):
    """Generate document content from v3 JSON data using template styles"""
    
    if json_data is None:
        print("ERROR: No JSON data to process")
        return
    
    content_blocks = json_data.get('content_blocks', [])
    
    for i, block in enumerate(content_blocks):
        text = block.get('text', '')
        level_type = block.get('level_type', 'content')
        bwa_level_name = block.get('bwa_level_name')
        number = block.get('number')
        level_number = block.get('level_number')
        
        # Skip empty content
        if not text.strip():
            continue
        
        # Clean the text for display (remove numbering prefixes)
        display_text = clean_text_for_display(text, level_type, number)
        
        # Determine the style to use
        if bwa_level_name:
            style_name = get_style_for_bwa_level(bwa_level_name)
        else:
            # Map level types to BWA styles first, then fallback to Word styles
            if level_type == "section":
                style_name = "BWA-SectionNumber"  # Try BWA style first
            elif level_type == "title":
                style_name = "BWA-SectionTitle"   # Try BWA style first
            elif level_type == "part_title":
                style_name = "BWA-PART"           # Use BWA style
            elif level_type == "subsection_title":
                style_name = "BWA-SUBSECTION"     # Use BWA style
            else:
                style_name = "Normal"
        
        # Add paragraph with appropriate style using cleaned text
        try:
            paragraph = doc.add_paragraph(display_text, style=style_name)
            apply_styling_from_json(paragraph, block)
            
        except Exception as e:
            # Fallback to Normal style if the specified style is not found
            paragraph = doc.add_paragraph(display_text, style="Normal")
            apply_styling_from_json(paragraph, block)
        
        # Add some spacing for better readability
        if level_type in ["section", "title", "part_title"]:
            doc.add_paragraph()  # Add blank line after major sections

def check_template_styles(template_path):
    """Check template styles and provide feedback about style definitions"""
    try:
        doc = Document(template_path)
        print(f"\nTemplate Style Analysis for: {template_path}")
        print("-" * 50)
        
        # Check for BWA styles specifically
        bwa_styles = []
        other_styles = []
        
        for style in doc.styles:
            if style.name and style.name.startswith('BWA-'):
                bwa_styles.append(style.name)
            elif style.name and style.name not in ['Normal', 'Default Paragraph Font', 'Default Paragraph Font (Asian)', 'Default Paragraph Font (Complex Script)']:
                other_styles.append(style.name)
        
        print(f"BWA Styles found: {len(bwa_styles)}")
        for style_name in bwa_styles:
            print(f"  - {style_name}")
        
        if other_styles:
            print(f"\nOther custom styles found: {len(other_styles)}")
            for style_name in other_styles[:10]:  # Show first 10
                print(f"  - {style_name}")
            if len(other_styles) > 10:
                print(f"  ... and {len(other_styles) - 10} more")
        
        print(f"\nTotal styles in template: {len(doc.styles)}")
        print("\nNote: If styles appear with default formatting in regenerated documents,")
        print("consider setting them to 'New documents based on this template' in Word.")
        
    except Exception as e:
        print(f"Warning: Could not analyze template styles: {e}")

def clone_template_styles(template_doc, target_doc):
    """Clone all styles from template document to target document, removing default inheritance"""
    try:
        print("Cloning template styles to generated document...")
        
        # Clear existing styles from target document (except built-in ones we can't remove)
        styles_to_remove = []
        for style in target_doc.styles:
            # Don't remove built-in styles that can't be deleted
            if style.name not in ['Normal', 'Default Paragraph Font', 'Default Paragraph Font (Asian)', 'Default Paragraph Font (Complex Script)']:
                styles_to_remove.append(style.name)
        
        # Remove custom styles
        for style_name in styles_to_remove:
            try:
                del target_doc.styles[style_name]
            except:
                pass  # Some styles can't be deleted
        
        # Clone styles from template
        cloned_styles = 0
        for template_style in template_doc.styles:
            try:
                # Skip built-in styles that already exist
                if template_style.name in ['Normal', 'Default Paragraph Font', 'Default Paragraph Font (Asian)', 'Default Paragraph Font (Complex Script)']:
                    continue
                
                # Check if style already exists in target
                if template_style.name in target_doc.styles:
                    # Update existing style
                    target_style = target_doc.styles[template_style.name]
                else:
                    # Create new style
                    target_style = target_doc.styles.add_style(template_style.name, WD_STYLE_TYPE.PARAGRAPH)
                
                # Clone paragraph format
                if hasattr(template_style, 'paragraph_format') and template_style.paragraph_format:
                    if hasattr(target_style, 'paragraph_format') and target_style.paragraph_format:
                        pf_template = template_style.paragraph_format
                        pf_target = target_style.paragraph_format
                        
                        # Copy all paragraph format properties
                        if pf_template.alignment:
                            pf_target.alignment = pf_template.alignment
                        if pf_template.left_indent:
                            pf_target.left_indent = pf_template.left_indent
                        if pf_template.right_indent:
                            pf_target.right_indent = pf_template.right_indent
                        if pf_template.first_line_indent:
                            pf_target.first_line_indent = pf_template.first_line_indent
                        if pf_template.space_before:
                            pf_target.space_before = pf_template.space_before
                        if pf_template.space_after:
                            pf_target.space_after = pf_template.space_after
                        if pf_template.line_spacing:
                            pf_target.line_spacing = pf_template.line_spacing
                        if pf_template.keep_with_next is not None:
                            pf_target.keep_with_next = pf_template.keep_with_next
                        if hasattr(pf_template, 'keep_lines_together') and pf_template.keep_lines_together is not None:
                            pf_target.keep_lines_together = pf_template.keep_lines_together
                        if pf_template.page_break_before is not None:
                            pf_target.page_break_before = pf_template.page_break_before
                        if pf_template.widow_control is not None:
                            pf_target.widow_control = pf_template.widow_control
                
                # Clone font properties
                if hasattr(template_style, 'font') and template_style.font:
                    if hasattr(target_style, 'font') and target_style.font:
                        font_template = template_style.font
                        font_target = target_style.font
                        
                        # Copy all font properties
                        if font_template.name:
                            font_target.name = font_template.name
                        if font_template.size:
                            font_target.size = font_template.size
                        if font_template.bold is not None:
                            font_target.bold = font_template.bold
                        if font_template.italic is not None:
                            font_target.italic = font_template.italic
                        if font_template.underline:
                            font_target.underline = font_template.underline
                        if font_template.color.rgb:
                            font_target.color.rgb = font_template.color.rgb
                        if font_template.strike is not None:
                            font_target.strike = font_template.strike
                        if font_template.small_caps is not None:
                            font_target.small_caps = font_template.small_caps
                        if font_template.all_caps is not None:
                            font_target.all_caps = font_template.all_caps
                
                cloned_styles += 1
                
            except Exception as e:
                print(f"Warning: Could not clone style {template_style.name}: {e}")
        
        print(f"Successfully cloned {cloned_styles} styles from template")
        
        # Also clone the Normal style to ensure it doesn't inherit unwanted properties
        if 'Normal' in template_doc.styles and 'Normal' in target_doc.styles:
            try:
                normal_template = template_doc.styles['Normal']
                normal_target = target_doc.styles['Normal']
                
                # Clone Normal style properties
                if hasattr(normal_template, 'paragraph_format') and normal_template.paragraph_format:
                    if hasattr(normal_target, 'paragraph_format') and normal_target.paragraph_format:
                        pf_template = normal_template.paragraph_format
                        pf_target = normal_target.paragraph_format
                        
                        if pf_template.alignment:
                            pf_target.alignment = pf_template.alignment
                        if pf_template.left_indent:
                            pf_target.left_indent = pf_template.left_indent
                        if pf_template.right_indent:
                            pf_target.right_indent = pf_template.right_indent
                        if pf_template.first_line_indent:
                            pf_target.first_line_indent = pf_template.first_line_indent
                        if pf_template.space_before:
                            pf_target.space_before = pf_template.space_before
                        if pf_template.space_after:
                            pf_target.space_after = pf_template.space_after
                        if pf_template.line_spacing:
                            pf_target.line_spacing = pf_template.line_spacing
                
                if hasattr(normal_template, 'font') and normal_template.font:
                    if hasattr(normal_target, 'font') and normal_target.font:
                        font_template = normal_template.font
                        font_target = normal_target.font
                        
                        if font_template.name:
                            font_target.name = font_template.name
                        if font_template.size:
                            font_target.size = font_template.size
                        if font_template.bold is not None:
                            font_target.bold = font_template.bold
                        if font_template.italic is not None:
                            font_target.italic = font_template.italic
                        if font_template.underline:
                            font_target.underline = font_template.underline
                        if font_template.color.rgb:
                            font_target.color.rgb = font_template.color.rgb
                
                print("Cloned Normal style from template")
                
            except Exception as e:
                print(f"Warning: Could not clone Normal style: {e}")
        
    except Exception as e:
        print(f"Warning: Could not clone template styles: {e}")

def create_document_from_template(template_path):
    """Create a new document by cloning the template completely"""
    try:
        # Load the template
        template_doc = Document(template_path)
        
        # Create a new document
        new_doc = Document()
        
        # Clone all styles from template
        clone_template_styles(template_doc, new_doc)
        
        # Clone document settings
        if template_doc.sections and new_doc.sections:
            template_section = template_doc.sections[0]
            new_section = new_doc.sections[0]
            
            # Copy page size and margins
            if template_section.page_width:
                new_section.page_width = template_section.page_width
            if template_section.page_height:
                new_section.page_height = template_section.page_height
            if template_section.top_margin:
                new_section.top_margin = template_section.top_margin
            if template_section.bottom_margin:
                new_section.bottom_margin = template_section.bottom_margin
            if template_section.left_margin:
                new_section.left_margin = template_section.left_margin
            if template_section.right_margin:
                new_section.right_margin = template_section.right_margin
            if template_section.header_distance:
                new_section.header_distance = template_section.header_distance
            if template_section.footer_distance:
                new_section.footer_distance = template_section.footer_distance
        
        print("Created new document with cloned template styles and settings")
        return new_doc
        
    except Exception as e:
        print(f"Error creating document from template: {e}")
        # Fallback to regular document creation
        return Document()

# Main execution
print("Starting document generation process...")

# Check template styles first
check_template_styles(TEMPLATE_PATH)

# Create document by cloning template instead of loading it directly
doc = create_document_from_template(TEMPLATE_PATH)

# Clear existing content
clear_document(doc)

# Parse JSON content
json_data = parse_spec_json(CONTENT_PATH)

# Apply document settings from JSON
apply_document_settings_from_json(doc, json_data)

# Apply margins from JSON
apply_margins_from_json(doc, json_data)

# Apply style definitions from JSON
apply_style_definitions_from_json(doc, json_data)

# Generate content from JSON
generate_content_from_v3_json(doc, json_data)

# Save document
doc.save(OUTPUT_PATH)
print(f"Document saved as '{OUTPUT_PATH}' with {FONT_SIZE}pt {FONT_NAME} font")
print(f"Content source: {CONTENT_PATH}")
print("Note: The template's multilevel list style should be applied automatically")
print("if the paragraphs use the correct style names from the template.")






