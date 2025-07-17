#!/usr/bin/env python3
"""
Batch Processing Script for Specification Documents

This script processes all specification documents in the examples/Specs folder,
extracting content and regenerating documents with proper formatting.

Usage:
    python batch_process_specs.py
"""

import os
import sys
import subprocess
import time
import shutil
from pathlib import Path

def run_command(command, description):
    """Run a command and handle errors"""
    print(f"\n{'='*60}")
    print(f"Running: {description}")
    print(f"Command: {command}")
    print(f"{'='*60}")
    
    start_time = time.time()
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True)
        end_time = time.time()
        
        if result.returncode == 0:
            print(f"✓ SUCCESS: {description}")
            print(f"  Time: {end_time - start_time:.2f} seconds")
            if result.stdout.strip():
                print(f"  Output: {result.stdout.strip()}")
            return True
        else:
            print(f"✗ FAILED: {description}")
            print(f"  Return code: {result.returncode}")
            print(f"  Error: {result.stderr.strip()}")
            return False
            
    except Exception as e:
        print(f"✗ ERROR: {description}")
        print(f"  Exception: {e}")
        return False

def create_generator_script(json_path, output_path):
    """Create a temporary generator script for a specific document"""
    generator_template = '''#!/usr/bin/env python3
"""
Temporary Generator Script for Batch Processing
Generated automatically by batch_process_specs.py
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration
CONTENT_PATH = '{json_path}'
OUTPUT_PATH = '{output_path}'
TEMPLATE_PATH = '../templates/test_template_cleaned.docx'
FONT_SIZE = 10
FONT_NAME = 'Arial'

def load_template_analysis():
    """Load template analysis from the template file"""
    try:
        # Import the template list detector module
        from template_list_detector import TemplateListDetector
        
        detector = TemplateListDetector()
        analysis = detector.analyze_template(TEMPLATE_PATH)
        return analysis
    except Exception as e:
        print(f"Warning: Could not load template analysis: {{e}}")
        return None

def apply_styling_from_json(paragraph, block):
    """Apply styling from JSON block to paragraph"""
    try:
        # Apply font properties
        if block.get('font_name'):
            for run in paragraph.runs:
                run.font.name = block['font_name']
        
        if block.get('font_size'):
            for run in paragraph.runs:
                run.font.size = Inches(block['font_size'] / 72.0)  # Convert points to inches
        
        if block.get('font_bold') is not None:
            for run in paragraph.runs:
                run.font.bold = block['font_bold']
        
        if block.get('font_italic') is not None:
            for run in paragraph.runs:
                run.font.italic = block['font_italic']
        
        # Apply paragraph properties
        if block.get('paragraph_alignment'):
            paragraph.alignment = getattr(paragraph.alignment, block['paragraph_alignment'].upper(), None)
        
        # Apply highlighting for corrected blocks
        if block.get('used_fallback_styling'):
            # Add yellow highlighting
            for run in paragraph.runs:
                run.font.highlight_color = 'yellow'
                
    except Exception as e:
        print(f"Warning: Could not apply styling: {{e}}")

def apply_explicit_indentation(paragraph, level_type, level_number, template_analysis):
    """Apply explicit indentation based on template analysis"""
    try:
        if not template_analysis or not template_analysis.bwa_list_levels:
            return
        
        # Find the BWA level info for this level type
        bwa_style_name = None
        if level_type == "part":
            bwa_style_name = "BWA-PART"
        elif level_type == "subsection":
            bwa_style_name = "BWA-SUBSECTION"
        elif level_type == "item":
            bwa_style_name = "BWA-Item"
        elif level_type == "list":
            bwa_style_name = "BWA-List"
        elif level_type == "sub_list":
            bwa_style_name = "BWA-SubList"
        
        if bwa_style_name and bwa_style_name in template_analysis.bwa_list_levels:
            level_info = template_analysis.bwa_list_levels[bwa_style_name]
            
            # Get indentation values from template analysis
            if level_info.abstract_num_id and level_info.abstract_num_id in template_analysis.numbering_definitions:
                abstract_info = template_analysis.numbering_definitions[level_info.abstract_num_id]
                level_str = str(level_number) if level_number is not None else "0"
                
                if level_str in abstract_info.get("levels", {}):
                    level_data = abstract_info["levels"][level_str]
                    p_pr = level_data.get("pPr", {})
                    
                    if "indent" in p_pr:
                        indent = p_pr["indent"]
                        
                        # Convert twips to inches (1 inch = 1440 twips)
                        if indent.get("left"):
                            left_indent = float(indent["left"]) / 1440.0
                            paragraph.paragraph_format.left_indent = Inches(left_indent)
                            print(f"DEBUG: Applied left indent {{left_indent:.3f}} inches for {{bwa_style_name}}")
                        
                        if indent.get("hanging"):
                            hanging_indent = float(indent["hanging"]) / 1440.0
                            paragraph.paragraph_format.hanging_indent = Inches(hanging_indent)
                            print(f"DEBUG: Applied hanging indent {{hanging_indent:.3f}} inches for {{bwa_style_name}}")
                        
                        if indent.get("firstLine"):
                            first_line_indent = float(indent["firstLine"]) / 1440.0
                            paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)
                            print(f"DEBUG: Applied first line indent {{first_line_indent:.3f}} inches for {{bwa_style_name}}")
                            
    except Exception as e:
        print(f"Warning: Could not apply explicit indentation: {{e}}")

def clean_text_for_display(text, level_type, original_number):
    """Clean text for display by removing numbering prefixes"""
    if not text:
        return ""
    
    # Remove common numbering patterns
    text = text.strip()
    
    # Remove part numbering (1.0, 2.0, etc.)
    if level_type == "part":
        text = re.sub(r'^\\d+\\.0\\s*[-]?\\s*', '', text)
    
    # Remove subsection numbering (1.01, 1.02, etc.)
    elif level_type == "subsection":
        text = re.sub(r'^\\d+\\.\\d{{2}}\\s*[-]?\\s*', '', text)
    
    # Remove item numbering (A., B., C., etc.)
    elif level_type == "item":
        text = re.sub(r'^[A-Z]\\.\\s*', '', text)
    
    # Remove list numbering (1., 2., 3., etc.)
    elif level_type == "list":
        text = re.sub(r'^\\d+\\.\\s*', '', text)
    
    # Remove sub-list numbering (a., b., c., etc.)
    elif level_type == "sub_list":
        text = re.sub(r'^[a-z]\\.\\s*', '', text)
    
    return text.strip()

def update_numbering_context(numbering_context, level_type):
    """Update numbering context and return correct number"""
    if level_type == "part":
        numbering_context["part_counter"] = numbering_context.get("part_counter", 0) + 1
        numbering_context["subsection_counter"] = 0
        numbering_context["item_counter"] = 0
        numbering_context["list_counter"] = 0
        return str(numbering_context["part_counter"]) + ".0"
    
    elif level_type == "subsection":
        numbering_context["subsection_counter"] = numbering_context.get("subsection_counter", 0) + 1
        numbering_context["item_counter"] = 0
        numbering_context["list_counter"] = 0
        part_num = numbering_context.get("part_counter", 1)
        subsection_num = numbering_context["subsection_counter"]
        return f"{{part_num}}.{{subsection_num:02d}}"
    
    elif level_type == "item":
        numbering_context["item_counter"] = numbering_context.get("item_counter", 0) + 1
        numbering_context["list_counter"] = 0
        # Convert to letter (A, B, C, etc.)
        item_num = numbering_context["item_counter"]
        if item_num <= 26:
            return chr(64 + item_num)  # A=65, B=66, etc.
        else:
            return f"{{chr(64 + (item_num-1)//26)}}{{chr(64 + ((item_num-1)%26) + 1)}}"
    
    elif level_type == "list":
        numbering_context["list_counter"] = numbering_context.get("list_counter", 0) + 1
        return str(numbering_context["list_counter"])
    
    elif level_type == "sub_list":
        numbering_context["sub_list_counter"] = numbering_context.get("sub_list_counter", 0) + 1
        # Convert to lowercase letter (a, b, c, etc.)
        sub_list_num = numbering_context["sub_list_counter"]
        if sub_list_num <= 26:
            return chr(96 + sub_list_num)  # a=97, b=98, etc.
        else:
            return f"{{chr(96 + (sub_list_num-1)//26)}}{{chr(96 + ((sub_list_num-1)%26) + 1)}}"
    
    return None

def generate_content_from_v3_json(json_path, template_analysis):
    """Generate Word document content from v3 JSON data"""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        content_blocks = data.get('content_blocks', [])
        if not content_blocks:
            print("No content blocks found in JSON")
            return None
        
        # Initialize numbering context
        numbering_context = {
            "part_counter": 0,
            "subsection_counter": 0,
            "item_counter": 0,
            "list_counter": 0,
            "sub_list_counter": 0
        }
        
        doc = Document(TEMPLATE_PATH)
        
        for i, block in enumerate(content_blocks):
            text = block.get('text', '')
            level_type = block.get('level_type', 'content')
            bwa_level_name = block.get('bwa_level_name')
            original_number = block.get('number')
            level_number = block.get('level_number')
            
            # Skip empty content
            if not text.strip():
                continue
            
            # Generate correct sequential numbering based on level type and context
            correct_number = update_numbering_context(numbering_context, level_type)
            
            # Debug output for numbering changes
            if level_type in ["part", "subsection", "item"] and correct_number:
                print(f"DEBUG: {{level_type.upper()}} numbering: {{original_number}} -> {{correct_number}}")
            
            # Use content field if available, otherwise clean the text
            display_text = block.get('content', '')
            if not display_text.strip():
                # Fallback to cleaning the text field
                display_text = clean_text_for_display(text, level_type, original_number)
            
            # Determine style name
            style_name = bwa_level_name if bwa_level_name else 'Normal'
            
            # Add paragraph with appropriate style using cleaned text
            try:
                paragraph = doc.add_paragraph(display_text, style=style_name)
                apply_styling_from_json(paragraph, block)
                
                # Apply explicit indentation
                apply_explicit_indentation(paragraph, level_type, level_number, template_analysis)
                
                # Apply list numbering if level_number is specified
                if level_number is not None and level_type in ["list", "sub_list", "item"]:
                    # Apply list numbering based on level_number
                    # This will use the template's multilevel list definitions
                    try:
                        # Set the list level (0-based in python-docx)
                        paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level_number
                        
                        # Use the correct numbering definition from the template
                        # numId 10-25 map to abstract_num_id "1" which has the correct BWA level configurations
                        num_id = 10 + level_number  # Map level 0->10, 1->11, 2->12, etc.
                        paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num_id
                        
                        print(f"DEBUG: Applied list numbering - Level: {{level_number}}, NumId: {{num_id}}")
                        
                    except Exception as e:
                        print(f"Warning: Could not apply list numbering: {{e}}")
                
            except Exception as e:
                print(f"Warning: Could not add paragraph: {{e}}")
                # Fallback: add as plain text
                doc.add_paragraph(display_text)
        
        return doc
        
    except Exception as e:
        print(f"Error generating content: {{e}}")
        return None

def main():
    """Main function"""
    print("Starting document generation...")
    
    # Load template analysis
    template_analysis = load_template_analysis()
    
    # Generate content
    doc = generate_content_from_v3_json(CONTENT_PATH, template_analysis)
    
    if doc:
        # Save document
        doc.save(OUTPUT_PATH)
        print(f"Document saved as '{{OUTPUT_PATH}}'")
    else:
        print("Failed to generate document")

if __name__ == "__main__":
    import re
    main()
'''
    
    # Create the temporary script
    temp_script = f"temp_generator_{int(time.time())}.py"
    with open(temp_script, 'w', encoding='utf-8') as f:
        f.write(generator_template.format(
            json_path=json_path.replace('\\', '/'),
            output_path=output_path.replace('\\', '/')
        ))
    
    return temp_script

def main():
    """Main batch processing function"""
    print("SpecConverter v0.4 - Batch Processing")
    print("=" * 50)
    
    # Define paths
    examples_dir = Path("../examples/Specs")
    output_dir = Path("../output/Specs")
    template_path = Path("../templates/test_template_cleaned.docx")
    
    # Check if directories exist
    if not examples_dir.exists():
        print(f"Error: Examples directory not found: {examples_dir}")
        return
    
    if not template_path.exists():
        print(f"Error: Template file not found: {template_path}")
        return
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Get all DOCX files in the examples directory
    docx_files = list(examples_dir.glob("*.docx"))
    
    if not docx_files:
        print(f"No DOCX files found in {examples_dir}")
        return
    
    print(f"Found {len(docx_files)} documents to process:")
    for docx_file in docx_files:
        print(f"  - {docx_file.name}")
    
    # Process each document
    successful_extractions = 0
    successful_regenerations = 0
    
    for i, docx_file in enumerate(docx_files, 1):
        print(f"\n{'='*60}")
        print(f"Processing document {i}/{len(docx_files)}: {docx_file.name}")
        print(f"{'='*60}")
        
        # Step 1: Extract content
        extraction_cmd = f'python extract_spec_content_v3.py "{docx_file}" . "{template_path}"'
        if run_command(extraction_cmd, f"Extracting content from {docx_file.name}"):
            successful_extractions += 1
            
            # Step 2: Generate regenerated document
            base_name = docx_file.stem
            json_path = f"../output/{base_name}_v3.json"
            output_path = output_dir / f"{base_name}_regenerated.docx"
            
            # Create temporary generator script
            temp_script = create_generator_script(str(json_path), str(output_path))
            
            try:
                # Run the temporary generator
                if run_command(f"python {temp_script}", f"Regenerating {docx_file.name}"):
                    successful_regenerations += 1
            finally:
                # Clean up temporary script
                if os.path.exists(temp_script):
                    os.remove(temp_script)
        else:
            print(f"Skipping regeneration for {docx_file.name} due to extraction failure")
    
    # Summary
    print(f"\n{'='*60}")
    print("BATCH PROCESSING COMPLETE")
    print(f"{'='*60}")
    print(f"Total documents: {len(docx_files)}")
    print(f"Successful extractions: {successful_extractions}")
    print(f"Successful regenerations: {successful_regenerations}")
    print(f"Output location: {output_dir}")
    
    if successful_regenerations == len(docx_files):
        print("\n✓ All documents processed successfully!")
    else:
        print(f"\n⚠ {len(docx_files) - successful_regenerations} documents had issues")

if __name__ == "__main__":
    main() 