#!/usr/bin/env python3
"""
Specification Content Extractor - Final Version

This script extracts multi-level list content from Word documents (.docx) and converts it to JSON format.
It handles the hierarchical structure of specification documents with parts, subsections, and items.
The script includes error detection and reporting for broken or inconsistent content.

Features:
- Extracts section headers, titles, parts, subsections, items, and lists
- Handles both numbered and unnumbered structures
- Validates numbering sequences and reports errors
- Generates comprehensive error reports
- Outputs JSON in the expected format for specification documents
- Enhanced subsection detection with format normalization
- Recovery logic for broken lists and missing items
- Hierarchical content processing to avoid confusion between levels

Usage:
    python extract_spec_content_final.py <docx_file> [output_dir]

Example:
    python extract_spec_content_final.py "SECTION 26 05 00.docx"
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os
import re
import sys
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime

@dataclass
class ExtractionError:
    """Represents an error found during content extraction"""
    line_number: int
    error_type: str
    message: str
    context: str
    expected: Optional[str] = None
    found: Optional[str] = None

class SpecContentExtractor:
    """Extracts specification content from Word documents using hierarchical processing"""
    
    def __init__(self):
        self.errors: List[ExtractionError] = []
        self.current_part_number: Optional[str] = None
        self.current_subsection_number: Optional[str] = None
        self.current_item_number: Optional[str] = None
        self.line_count = 0
        self.section_header_found = False
        self.section_title_found = False
        
        # Store all paragraphs for hierarchical analysis
        self.all_paragraphs = []
        self.paragraph_analysis = []
        
        # Regex patterns for different levels
        self.part_pattern = re.compile(r'^(\d+\.0)\s+(.+)$')
        self.subsection_pattern = re.compile(r'^(\d+\.\d{2})\s+(.+)$')
        self.subsection_alt_pattern = re.compile(r'^(\d+\.\d)\s+(.+)$')  # For 1.1, 1.2 format
        self.item_pattern = re.compile(r'^([A-Z])\.\s+(.+)$')
        self.list_pattern = re.compile(r'^(\d+)\.\s+(.+)$')
        self.sub_list_pattern = re.compile(r'^([a-z])\.\s+(.+)$')
        
        # Additional patterns for different formats
        self.section_pattern = re.compile(r'^SECTION\s+(.+)$')
        self.title_pattern = re.compile(r'^([A-Z][A-Z\s]+)$')
        
        # Patterns for actual document structure
        self.part_title_pattern = re.compile(r'^([A-Z][A-Z\s]+)$')  # GENERAL, PRODUCTS, EXECUTION
        self.subsection_title_pattern = re.compile(r'^([A-Z][A-Z\s&]+)$')  # SCOPE, EXISTING CONDITIONS, etc.
        
    def normalize_subsection_number(self, number: str) -> str:
        """Convert subsection numbers like '1.1' to '1.01' format"""
        if re.match(r'^\d+\.\d$', number):
            parts = number.split('.')
            if len(parts) == 2:
                return f"{parts[0]}.{parts[1]:0>2}"
        return number
    
    def classify_paragraph_level(self, text: str) -> Tuple[str, Optional[str], str]:
        """
        Classify a paragraph into its hierarchical level
        Returns: (level_type, number, content)
        """
        text = text.strip()
        if not text:
            return "empty", None, ""
        
        # Check for section header (must be the very first line)
        if text.upper().startswith("SECTION") and not self.section_header_found:
            match = self.section_pattern.match(text)
            if match:
                self.section_header_found = True
                return "section", match.group(1), ""
        elif text.upper().startswith("SECTION") and self.section_header_found:
            self.add_error("Structure Error", "Multiple section headers found", text)
            return "content", None, text
        
        # Check for section title (must be the second line after section header)
        if (self.section_header_found and 
            not self.section_title_found and
            len(text.strip()) > 0):
            self.section_title_found = True
            return "title", None, text
        
        # Check for part level with numbering (1.0, 2.0, etc.)
        match = self.part_pattern.match(text)
        if match:
            return "part", match.group(1), match.group(2)
        
        # Check for part titles with various formats
        part_names = ["DESCRIPTION", "PRODUCTS", "EXECUTION", "GENERAL"]
        for part_name in part_names:
            match = re.match(rf'(?:PART\s*)?(\d+)\.0?\s*[-]?\s*{part_name}$', text.upper())
            if match:
                part_number = f"{match.group(1)}.0"
                return "part_title", part_number, part_name
            elif text.strip().upper() == part_name:
                part_number = f"{len(self.extracted_data.get('parts', [])) + 1}.0"
                return "part_title", part_number, part_name
        
        # Check for subsection level with numbering (1.01, 1.02, etc.)
        match = self.subsection_pattern.match(text)
        if match:
            normalized_number = self.normalize_subsection_number(match.group(1))
            return "subsection", normalized_number, match.group(2)
        
        # Check for subsection level with alternative numbering (1.1, 1.2, etc.)
        match = self.subsection_alt_pattern.match(text)
        if match:
            normalized_number = self.normalize_subsection_number(match.group(1))
            return "subsection", normalized_number, match.group(2)
        
        # Check for subsection titles without numbering
        subsection_titles = [
            "SCOPE", "EXISTING CONDITIONS", "CODES AND REGULATIONS", "DEFINITIONS",
            "DRAWINGS AND SPECIFICATIONS", "SITE VISIT", "DEVIATIONS",
            "STANDARDS FOR MATERIALS AND WORKMANSHIP", "SHOP DRAWINGS AND SUBMITTAL",
            "RECORD (AS-BUILT) DRAWINGS AND MAINTENANCE MANUALS",
            "COORDINATION", "PROTECTION OF MATERIALS", "TESTS, DEMONSTRATION AND INSTRUCTIONS",
            "GUARANTEE"
        ]
        
        # Check for exact match
        if text.upper() in [title.upper() for title in subsection_titles]:
            return "subsection_title", None, text
        
        # Check for partial matches
        text_upper = text.upper().strip()
        for title in subsection_titles:
            if title.upper() in text_upper or text_upper in title.upper():
                return "subsection_title", None, title
        
        # Check for item level (A., B., C., etc.)
        match = self.item_pattern.match(text)
        if match:
            return "item", match.group(1), match.group(2)
        
        # Check for list level (1., 2., etc.)
        match = self.list_pattern.match(text)
        if match:
            return "list", match.group(1), match.group(2)
        
        # Check for sub-list level (a., b., etc.)
        match = self.sub_list_pattern.match(text)
        if match:
            return "sub_list", match.group(1), match.group(2)
        
        # If no pattern matches, it's regular content
        return "content", None, text
    
    def extract_section_header_and_title(self, paragraphs: List[str]) -> Tuple[str, str]:
        """Extract section header and title from the first few paragraphs"""
        section_number = ""
        section_title = ""
        
        if len(paragraphs) >= 2:
            # First paragraph should be section header
            section_text = paragraphs[0].strip()
            if section_text.upper().startswith("SECTION"):
                section_match = re.search(r'^SECTION\s+(.+)$', section_text, re.IGNORECASE)
                if section_match:
                    section_content = section_match.group(1).strip()
                    
                    # Try to extract section number from various formats
                    number_match = re.search(r'(\d+)\s+(\d+)\s+(\d+)', section_content)
                    if number_match:
                        section_number = f"{number_match.group(1)}{number_match.group(2)}{number_match.group(3)}"
                    else:
                        number_match = re.search(r'(\d+)-(\d+)-(\d+)', section_content)
                        if number_match:
                            section_number = f"{number_match.group(1)}{number_match.group(2)}{number_match.group(3)}"
                        else:
                            number_match = re.search(r'(\d{6})', section_content)
                            if number_match:
                                section_number = number_match.group(1)
                            else:
                                section_number = section_content.replace(" ", "").replace("-", "")
            
            # Second paragraph should be section title
            title_text = paragraphs[1].strip()
            if title_text and not title_text.upper().startswith("SECTION"):
                section_title = title_text
        
        return section_number, section_title
    
    def find_part_boundaries(self, paragraphs: List[str]) -> List[Dict]:
        """Find the boundaries of parts in the document"""
        parts = []
        current_part = None
        
        for i, text in enumerate(paragraphs):
            if not text.strip():
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type in ["part", "part_title"]:
                if current_part:
                    parts.append(current_part)
                
                current_part = {
                    "start_index": i,
                    "end_index": len(paragraphs) - 1,  # Will be updated when next part is found
                    "part_number": number,
                    "title": content,
                    "level_type": level_type
                }
        
        # Add the last part
        if current_part:
            parts.append(current_part)
        
        # Update end indices
        for i in range(len(parts) - 1):
            parts[i]["end_index"] = parts[i + 1]["start_index"] - 1
        
        return parts
    
    def find_subsection_boundaries(self, paragraphs: List[str], start_index: int, end_index: int) -> List[Dict]:
        """Find the boundaries of subsections within a part"""
        subsections = []
        current_subsection = None
        
        for i in range(start_index, end_index + 1):
            text = paragraphs[i].strip()
            if not text:
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type in ["subsection", "subsection_title"]:
                if current_subsection:
                    subsections.append(current_subsection)
                
                current_subsection = {
                    "start_index": i,
                    "end_index": end_index,  # Will be updated when next subsection is found
                    "subsection_number": number,
                    "title": content,
                    "level_type": level_type
                }
        
        # Add the last subsection
        if current_subsection:
            subsections.append(current_subsection)
        
        # Update end indices
        for i in range(len(subsections) - 1):
            subsections[i]["end_index"] = subsections[i + 1]["start_index"] - 1
        
        return subsections
    
    def extract_items_and_lists(self, paragraphs: List[str], start_index: int, end_index: int) -> List[Dict]:
        """Extract items and their associated lists from a subsection"""
        items = []
        current_item = None
        current_lists = []
        
        for i in range(start_index, end_index + 1):
            text = paragraphs[i].strip()
            if not text:
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type == "item":
                # Save previous item if exists
                if current_item:
                    current_item["lists"] = current_lists
                    items.append(current_item)
                
                # Start new item
                current_item = {
                    "item_number": number,
                    "text": content,
                    "lists": []
                }
                current_lists = []
                
            elif level_type == "list":
                if current_item:
                    current_lists.append({
                        "list_number": number,
                        "text": content
                    })
                else:
                    # List without parent item - create a default item
                    current_item = {
                        "item_number": "A",
                        "text": "",
                        "lists": [{
                            "list_number": number,
                            "text": content
                        }]
                    }
                    current_lists = []
                    
            elif level_type == "sub_list":
                if current_lists:
                    # Add to the last list item
                    current_lists[-1]["sub_lists"] = current_lists[-1].get("sub_lists", [])
                    current_lists[-1]["sub_lists"].append({
                        "list_number": number,
                        "text": content
                    })
                    
            elif level_type == "content":
                if current_item:
                    # Append to current item's text
                    if current_item["text"]:
                        current_item["text"] += " " + content
                    else:
                        current_item["text"] = content
        
        # Save the last item
        if current_item:
            current_item["lists"] = current_lists
            items.append(current_item)
        
        return items
    
    def process_part_hierarchically(self, paragraphs: List[str], part_info: Dict) -> Dict:
        """Process a part and all its subsections hierarchically"""
        part_data = {
            "part_number": part_info["part_number"],
            "title": part_info["title"],
            "subsections": []
        }
        
        # Find subsections within this part
        subsections = self.find_subsection_boundaries(
            paragraphs, 
            part_info["start_index"], 
            part_info["end_index"]
        )
        
        for subsection_info in subsections:
            # Generate subsection number if not provided
            if not subsection_info["subsection_number"]:
                subsection_number = f"{part_data['part_number'].split('.')[0]}.{len(part_data['subsections']) + 1:02d}"
            else:
                subsection_number = subsection_info["subsection_number"]
            
            # Extract items and lists for this subsection
            items = self.extract_items_and_lists(
                paragraphs,
                subsection_info["start_index"],
                subsection_info["end_index"]
            )
            
            subsection_data = {
                "subsection_number": subsection_number,
                "title": subsection_info["title"],
                "items": items
            }
            
            part_data["subsections"].append(subsection_data)
        
        return part_data
    
    def add_error(self, error_type: str, message: str, context: str = "", 
                  expected: Optional[str] = None, found: Optional[str] = None):
        """Add an error to the error list"""
        error = ExtractionError(
            line_number=self.line_count,
            error_type=error_type,
            message=message,
            context=context,
            expected=expected,
            found=found
        )
        self.errors.append(error)
        print(f"ERROR (line {self.line_count}): {error_type} - {message}")
        if context:
            print(f"  Context: {context}")
        if expected and found:
            print(f"  Expected: {expected}, Found: {found}")
    
    def extract_text_from_element(self, element, nsmap) -> str:
        """Extract all text from an element and its children"""
        texts = []
        for text_elem in element.findall('.//w:t', namespaces=nsmap):
            if text_elem.text:
                texts.append(text_elem.text)
        return ''.join(texts).strip()
    
    def get_paragraph_style(self, paragraph) -> str:
        """Get the style name of a paragraph"""
        try:
            return paragraph.style.name if paragraph.style else "Normal"
        except:
            return "Normal"
    
    def get_paragraph_numbering(self, paragraph) -> Optional[str]:
        """Get the numbering level and format of a paragraph"""
        try:
            if paragraph._element.pPr and paragraph._element.pPr.numPr:
                numPr = paragraph._element.pPr.numPr
                if numPr.numId and numPr.ilvl:
                    return f"Level {numPr.ilvl.val}"
        except:
            pass
        return None
    
    def parse_paragraph_content(self, text: str) -> Tuple[Optional[str], Optional[str], str]:
        """
        Parse paragraph content to determine level and extract content
        Returns: (level_type, number, content)
        """
        text = text.strip()
        if not text:
            return None, None, ""
        
        # Check for section header (must be the very first line)
        if text.upper().startswith("SECTION") and not self.section_header_found:
            match = self.section_pattern.match(text)
            if match:
                self.section_header_found = True
                return "section", match.group(1), ""
        elif text.upper().startswith("SECTION") and self.section_header_found:
            # Found another section header after the first one - this is an error
            self.add_error("Structure Error", "Multiple section headers found", text)
            return "content", None, text
        
        # Check for section title (must be the second line after section header)
        if (self.section_header_found and 
            not self.section_title_found and
            len(text.strip()) > 0):
            # Any non-empty text on the line after section header is the title
            self.section_title_found = True
            return "title", None, text
        
        # Check for part level with numbering (1.0, 2.0, etc.)
        match = self.part_pattern.match(text)
        if match:
            return "part", match.group(1), match.group(2)
        
        # Check for part titles with various formats
        # Look for the specific part names: DESCRIPTION, PRODUCTS, EXECUTION, GENERAL
        # Handle formats like: "PART 1 - GENERAL", "1.0 GENERAL", "GENERAL", etc.
        part_names = ["DESCRIPTION", "PRODUCTS", "EXECUTION", "GENERAL"]
        for part_name in part_names:
            # Accepts: "PART 1 - GENERAL", "1.0 GENERAL", "GENERAL", "2.0 PRODUCTS", etc.
            match = re.match(rf'(?:PART\s*)?(\d+)\.0?\s*[-]?\s*{part_name}$', text.upper())
            if match:
                part_number = f"{match.group(1)}.0"
                return "part_title", part_number, part_name
            elif text.strip().upper() == part_name:
                # No number, assign next in sequence
                part_number = f"{len(self.extracted_data.get('parts', [])) + 1}.0"
                return "part_title", part_number, part_name
        
        # Check for subsection level with numbering (1.01, 1.02, etc.)
        match = self.subsection_pattern.match(text)
        if match:
            normalized_number = self.normalize_subsection_number(match.group(1))
            return "subsection", normalized_number, match.group(2)
        
        # Check for subsection level with alternative numbering (1.1, 1.2, etc.)
        match = self.subsection_alt_pattern.match(text)
        if match:
            normalized_number = self.normalize_subsection_number(match.group(1))
            return "subsection", normalized_number, match.group(2)
        
        # Check for subsection titles without numbering
        subsection_titles = [
            "SCOPE", "EXISTING CONDITIONS", "CODES AND REGULATIONS", "DEFINITIONS",
            "DRAWINGS AND SPECIFICATIONS", "SITE VISIT", "DEVIATIONS",
            "STANDARDS FOR MATERIALS AND WORKMANSHIP", "SHOP DRAWINGS AND SUBMITTAL",
            "RECORD (AS-BUILT) DRAWINGS AND MAINTENANCE MANUALS",
            "COORDINATION", "PROTECTION OF MATERIALS", "TESTS, DEMONSTRATION AND INSTRUCTIONS",
            "GUARANTEE"
        ]
        
        # Check for exact match
        if text.upper() in [title.upper() for title in subsection_titles]:
            return "subsection_title", None, text
        
        # Check for partial matches (some titles might have extra spaces or formatting)
        text_upper = text.upper().strip()
        for title in subsection_titles:
            if title.upper() in text_upper or text_upper in title.upper():
                return "subsection_title", None, title
        
        # Check for item level (A., B., C., etc.)
        match = self.item_pattern.match(text)
        if match:
            return "item", match.group(1), match.group(2)
        
        # Check for list level (1., 2., etc.)
        match = self.list_pattern.match(text)
        if match:
            return "list", match.group(1), match.group(2)
        
        # Check for sub-list level (a., b., etc.)
        match = self.sub_list_pattern.match(text)
        if match:
            return "sub_list", match.group(1), match.group(2)
        
        # If no pattern matches, it's regular content
        return "content", None, text
    
    def validate_numbering_sequence(self, level_type: str, number: str, content: str):
        """Validate that numbering follows expected sequence"""
        if level_type == "part":
            expected_part = f"{len(self.extracted_data.get('parts', [])) + 1}.0"
            if number != expected_part:
                self.add_error(
                    "Numbering Sequence Error",
                    f"Unexpected part number",
                    content,
                    expected_part,
                    number
                )
            self.current_part_number = number
            
        elif level_type == "subsection":
            if not self.current_part_number:
                self.add_error(
                    "Structure Error",
                    "Subsection found without preceding part",
                    content
                )
            else:
                part_num = self.current_part_number.split('.')[0]
                # Find the current part to get its subsections
                current_part = None
                for part in self.extracted_data.get('parts', []):
                    if part.get('part_number') == self.current_part_number:
                        current_part = part
                        break
                
                if current_part:
                    expected_subsection = f"{part_num}.{len(current_part.get('subsections', [])) + 1:02d}"
                    if number != expected_subsection:
                        self.add_error(
                            "Numbering Sequence Error",
                            f"Unexpected subsection number",
                            content,
                            expected_subsection,
                            number
                        )
            self.current_subsection_number = number
            
        elif level_type == "item":
            if not self.current_subsection_number:
                self.add_error(
                    "Structure Error",
                    "Item found without preceding subsection",
                    content
                )
            else:
                # Find the current subsection to get its items
                current_subsection = None
                for part in self.extracted_data.get('parts', []):
                    for sub in part.get('subsections', []):
                        if sub.get('subsection_number') == self.current_subsection_number:
                            current_subsection = sub
                            break
                    if current_subsection:
                        break
                
                if current_subsection:
                    expected_item = chr(ord('A') + len(current_subsection.get('items', [])))
                    if number != expected_item:
                        self.add_error(
                            "Numbering Sequence Error",
                            f"Unexpected item number",
                            content,
                            expected_item,
                            number
                        )
            self.current_item_number = number
    
    def extract_header_info(self, doc) -> Dict[str, Any]:
        """Extract header information from document"""
        header_info = {
            "bwa_number": "2025-XXXX",
            "client_number": "ZZZ# 00000", 
            "project_name": "PROJECT NAME",
            "company_name": "CLIENT NAME",
            "section_number": "",
            "section_title": ""
        }
        
        # Extract section number and title from first two paragraphs
        if len(doc.paragraphs) >= 2:
            # First paragraph should be section header
            section_text = doc.paragraphs[0].text.strip()
            if section_text.upper().startswith("SECTION"):
                # Extract section content (everything after "SECTION")
                section_match = re.search(r'^SECTION\s+(.+)$', section_text, re.IGNORECASE)
                if section_match:
                    section_content = section_match.group(1).strip()
                    
                    # Try to extract section number from various formats
                    # Format 1: "26 05 00" -> "260500"
                    number_match = re.search(r'(\d+)\s+(\d+)\s+(\d+)', section_content)
                    if number_match:
                        header_info["section_number"] = f"{number_match.group(1)}{number_match.group(2)}{number_match.group(3)}"
                    else:
                        # Format 2: "26-05-00" -> "260500"
                        number_match = re.search(r'(\d+)-(\d+)-(\d+)', section_content)
                        if number_match:
                            header_info["section_number"] = f"{number_match.group(1)}{number_match.group(2)}{number_match.group(3)}"
                        else:
                            # Format 3: "260500" -> "260500"
                            number_match = re.search(r'(\d{6})', section_content)
                            if number_match:
                                header_info["section_number"] = number_match.group(1)
                            else:
                                # If no number found, use the entire content as section number
                                header_info["section_number"] = section_content.replace(" ", "").replace("-", "")
            
            # Second paragraph should be section title
            title_text = doc.paragraphs[1].text.strip()
            if title_text and not title_text.upper().startswith("SECTION"):
                header_info["section_title"] = title_text
        
        return header_info
    
    def extract_content(self, docx_path: str) -> Dict[str, Any]:
        """Extract all content from the Word document using hierarchical processing"""
        print(f"Extracting content from: {docx_path}")
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            self.add_error("File Error", f"Could not open document: {e}")
            return {}
        
        # Initialize extraction data
        self.extracted_data = {
            "header": self.extract_header_info(doc),
            "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
            "margins": {},
            "comments": [],
            "section_number": "",
            "section_title": "",
            "parts": []
        }
        
        # Extract margins
        try:
            sec = doc.sections[0]
            self.extracted_data["margins"] = {
                "top_margin": sec.top_margin.inches if sec.top_margin else 1.0,
                "bottom_margin": sec.bottom_margin.inches if sec.bottom_margin else 1.0,
                "left_margin": sec.left_margin.inches if sec.left_margin else 1.0,
                "right_margin": sec.right_margin.inches if sec.right_margin else 1.0,
                "header_distance": sec.header_distance.inches if sec.header_distance else 1.0,
                "footer_distance": sec.footer_distance.inches if sec.footer_distance else 1.0
            }
        except Exception as e:
            self.add_error("Margin Error", f"Could not extract margins: {e}")
        
        # First pass: Collect all paragraphs
        self.all_paragraphs = [para.text.strip() for para in doc.paragraphs]
        
        # Extract section header and title
        self.extracted_data["section_number"], self.extracted_data["section_title"] = self.extract_section_header_and_title(self.all_paragraphs)
        
        # Find parts and process them hierarchically
        parts = self.find_part_boundaries(self.all_paragraphs)
        
        for part_info in parts:
            part_data = self.process_part_hierarchically(self.all_paragraphs, part_info)
            self.extracted_data["parts"].append(part_data)
        
        return self.extracted_data
    
    def generate_error_report(self) -> str:
        """Generate a detailed error report"""
        if not self.errors:
            return "No errors detected during extraction."
        
        report = f"ERROR REPORT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += "=" * 60 + "\n\n"
        
        # Group errors by type
        error_types = {}
        for error in self.errors:
            if error.error_type not in error_types:
                error_types[error.error_type] = []
            error_types[error.error_type].append(error)
        
        for error_type, errors in error_types.items():
            report += f"{error_type} ({len(errors)} errors):\n"
            report += "-" * (len(error_type) + 10) + "\n"
            
            for error in errors:
                report += f"Line {error.line_number}: {error.message}\n"
                if error.context:
                    report += f"  Context: {error.context}\n"
                if error.expected and error.found:
                    report += f"  Expected: {error.expected}, Found: {error.found}\n"
                report += "\n"
        
        return report
    
    def save_to_json(self, data: Dict[str, Any], output_path: str):
        """Save extracted data to JSON file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            print(f"Content saved to: {output_path}")
        except Exception as e:
            self.add_error("Save Error", f"Could not save JSON file: {e}")
    
    def save_error_report(self, report: str, output_path: str):
        """Save error report to text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report)
            print(f"Error report saved to: {output_path}")
        except Exception as e:
            print(f"Could not save error report: {e}")

    def extract_all_numbered_elements(self, paragraphs: List[str]) -> List[Dict]:
        """
        Phase 1: Extract all numbered elements from the document
        Returns a list of all numbered elements with their positions and levels
        """
        numbered_elements = []
        
        for i, text in enumerate(paragraphs):
            text = text.strip()
            if not text:
                continue
                
            level_type, number, content = self.classify_paragraph_level(text)
            
            if level_type in ["part", "subsection", "item", "list", "sub_list"]:
                numbered_elements.append({
                    "index": i,
                    "level_type": level_type,
                    "number": number,
                    "content": content,
                    "full_text": text,
                    "processed": False
                })
        
        return numbered_elements
    
    def reconstruct_broken_lists(self, numbered_elements: List[Dict]) -> List[Dict]:
        """
        Phase 1: Reconstruct broken lists by analyzing numbering patterns
        """
        reconstructed_elements = []
        
        # Group elements by their hierarchical level
        parts = [e for e in numbered_elements if e["level_type"] == "part"]
        subsections = [e for e in numbered_elements if e["level_type"] == "subsection"]
        items = [e for e in numbered_elements if e["level_type"] == "item"]
        lists = [e for e in numbered_elements if e["level_type"] == "list"]
        sub_lists = [e for e in numbered_elements if e["level_type"] == "sub_list"]
        
        # Reconstruct parts
        for part in parts:
            part_number = part["number"]
            # Find all subsections that should belong to this part
            part_subsections = [s for s in subsections if s["number"].startswith(part_number.split('.')[0] + '.')]
            
            # Check for missing subsections
            expected_subsection_numbers = []
            part_num = int(part_number.split('.')[0])
            for i in range(1, 10):  # Check for subsections 1.01 through 1.09
                expected_subsection_numbers.append(f"{part_num}.{i:02d}")
            
            # Find missing subsections by looking for orphaned content
            for expected_num in expected_subsection_numbers:
                if not any(s["number"] == expected_num for s in part_subsections):
                    # Look for content that might be a missing subsection
                    for elem in numbered_elements:
                        if (elem["level_type"] == "content" and 
                            not elem["processed"] and
                            elem["index"] > part["index"]):
                            # Check if this content looks like a subsection title
                            if self.looks_like_subsection_title(elem["content"]):
                                # Create a reconstructed subsection
                                reconstructed_elements.append({
                                    "index": elem["index"],
                                    "level_type": "subsection",
                                    "number": expected_num,
                                    "content": elem["content"],
                                    "full_text": elem["full_text"],
                                    "processed": True,
                                    "reconstructed": True
                                })
                                elem["processed"] = True
                                break
        
        # Reconstruct items within subsections
        for subsection in subsections:
            subsection_number = subsection["number"]
            # Find all items that should belong to this subsection
            subsection_items = [i for i in items if i["index"] > subsection["index"]]
            
            # Check for missing items (A, B, C, etc.)
            expected_item_letters = [chr(ord('A') + i) for i in range(26)]
            found_letters = [i["number"] for i in subsection_items]
            
            for expected_letter in expected_item_letters:
                if expected_letter not in found_letters:
                    # Look for content that might be a missing item
                    for elem in numbered_elements:
                        if (elem["level_type"] == "content" and 
                            not elem["processed"] and
                            elem["index"] > subsection["index"]):
                            # Check if this content looks like an item
                            if self.looks_like_item_content(elem["content"]):
                                # Create a reconstructed item
                                reconstructed_elements.append({
                                    "index": elem["index"],
                                    "level_type": "item",
                                    "number": expected_letter,
                                    "content": elem["content"],
                                    "full_text": elem["full_text"],
                                    "processed": True,
                                    "reconstructed": True
                                })
                                elem["processed"] = True
                                break
        
        # Reconstruct lists within items
        for item in items:
            item_letter = item["number"]
            # Find all lists that should belong to this item
            item_lists = [l for l in lists if l["index"] > item["index"]]
            
            # Check for missing list numbers
            expected_list_numbers = list(range(1, 21))  # Check for lists 1-20
            found_numbers = [int(l["number"]) for l in item_lists if l["number"].isdigit()]
            
            for expected_num in expected_list_numbers:
                if expected_num not in found_numbers:
                    # Look for content that might be a missing list item
                    for elem in numbered_elements:
                        if (elem["level_type"] == "content" and 
                            not elem["processed"] and
                            elem["index"] > item["index"]):
                            # Check if this content looks like a list item
                            if self.looks_like_list_content(elem["content"]):
                                # Create a reconstructed list item
                                reconstructed_elements.append({
                                    "index": elem["index"],
                                    "level_type": "list",
                                    "number": str(expected_num),
                                    "content": elem["content"],
                                    "full_text": elem["full_text"],
                                    "processed": True,
                                    "reconstructed": True
                                })
                                elem["processed"] = True
                                break
        
        # Combine original and reconstructed elements
        all_elements = numbered_elements + reconstructed_elements
        all_elements.sort(key=lambda x: x["index"])
        
        return all_elements
    
    def scan_content_for_embedded_numbering(self, paragraphs: List[str], numbered_elements: List[Dict]) -> List[Dict]:
        """
        Phase 2: Scan content for embedded numbering patterns
        """
        embedded_elements = []
        
        for elem in numbered_elements:
            if elem["level_type"] in ["item", "list"]:
                content = elem["content"]
                
                # Look for embedded item patterns (C., D., etc.)
                item_matches = re.finditer(r'\b([A-Z])\.\s+', content)
                for match in item_matches:
                    letter = match.group(1)
                    # Extract content after the letter
                    start_pos = match.end()
                    end_pos = content.find('\n', start_pos)
                    if end_pos == -1:
                        end_pos = len(content)
                    
                    embedded_content = content[start_pos:end_pos].strip()
                    if embedded_content:
                        embedded_elements.append({
                            "parent_index": elem["index"],
                            "parent_type": elem["level_type"],
                            "parent_number": elem["number"],
                            "level_type": "embedded_item",
                            "number": letter,
                            "content": embedded_content,
                            "full_text": f"{letter}. {embedded_content}",
                            "processed": False,
                            "embedded": True
                        })
                
                # Look for embedded list patterns (1., 2., etc.)
                list_matches = re.finditer(r'\b(\d+)\.\s+', content)
                for match in list_matches:
                    number = match.group(1)
                    # Extract content after the number
                    start_pos = match.end()
                    end_pos = content.find('\n', start_pos)
                    if end_pos == -1:
                        end_pos = len(content)
                    
                    embedded_content = content[start_pos:end_pos].strip()
                    if embedded_content:
                        embedded_elements.append({
                            "parent_index": elem["index"],
                            "parent_type": elem["level_type"],
                            "parent_number": elem["number"],
                            "level_type": "embedded_list",
                            "number": number,
                            "content": embedded_content,
                            "full_text": f"{number}. {embedded_content}",
                            "processed": False,
                            "embedded": True
                        })
        
        return embedded_elements
    
    def looks_like_subsection_title(self, text: str) -> bool:
        """Check if text looks like a subsection title"""
        subsection_titles = [
            "SCOPE", "EXISTING CONDITIONS", "CODES AND REGULATIONS", "DEFINITIONS",
            "DRAWINGS AND SPECIFICATIONS", "SITE VISIT", "DEVIATIONS",
            "STANDARDS FOR MATERIALS AND WORKMANSHIP", "SHOP DRAWINGS AND SUBMITTAL",
            "RECORD (AS-BUILT) DRAWINGS AND MAINTENANCE MANUALS",
            "COORDINATION", "PROTECTION OF MATERIALS", "TESTS, DEMONSTRATION AND INSTRUCTIONS",
            "GUARANTEE"
        ]
        
        text_upper = text.upper().strip()
        return any(title.upper() in text_upper or text_upper in title.upper() 
                  for title in subsection_titles)
    
    def looks_like_item_content(self, text: str) -> bool:
        """Check if text looks like item content"""
        # Item content is typically a complete sentence or paragraph
        # and doesn't start with common list patterns
        text = text.strip()
        if not text:
            return False
        
        # Should not start with common list patterns
        if re.match(r'^[A-Z]\.\s+', text):
            return False
        if re.match(r'^\d+\.\s+', text):
            return False
        if re.match(r'^[a-z]\.\s+', text):
            return False
        
        # Should be substantial content (not just a few words)
        if len(text.split()) < 3:
            return False
        
        # Should look like a complete thought
        return text.endswith('.') or text.endswith(':') or len(text) > 50
    
    def looks_like_list_content(self, text: str) -> bool:
        """Check if text looks like list content"""
        # List content is typically shorter and more specific
        text = text.strip()
        if not text:
            return False
        
        # Should not start with common list patterns
        if re.match(r'^[A-Z]\.\s+', text):
            return False
        if re.match(r'^\d+\.\s+', text):
            return False
        if re.match(r'^[a-z]\.\s+', text):
            return False
        
        # Should be specific content (not too long)
        if len(text) > 200:
            return False
        
        # Should look like a list item
        return len(text.split()) >= 2

def main():
    """Main function to process Word documents"""
    if len(sys.argv) < 2:
        print("Usage: python extract_spec_content_final.py <docx_file> [output_dir]")
        print("Example: python extract_spec_content_final.py 'SECTION 26 05 00.docx'")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "."
    
    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)
    
    # Create extractor and process document
    extractor = SpecContentExtractor()
    data = extractor.extract_content(docx_path)
    
    if not data:
        print("Error: No data extracted from document")
        sys.exit(1)
    
    # Generate output filenames
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    json_path = os.path.join(output_dir, f"{base_name}_content.json")
    error_path = os.path.join(output_dir, f"{base_name}_errors.txt")
    
    # Save results
    extractor.save_to_json(data, json_path)
    
    # Generate and save error report
    error_report = extractor.generate_error_report()
    extractor.save_error_report(error_report, error_path)
    
    # Print summary
    print(f"\nExtraction Summary:")
    print(f"  Document: {docx_path}")
    print(f"  Parts found: {len(data.get('parts', []))}")
    print(f"  Total subsections: {sum(len(part.get('subsections', [])) for part in data.get('parts', []))}")
    print(f"  Total items: {sum(len(sub.get('items', [])) for part in data.get('parts', []) for sub in part.get('subsections', []))}")
    print(f"  Errors detected: {len(extractor.errors)}")
    print(f"  Output files:")
    print(f"    - {json_path}")
    print(f"    - {error_path}")

if __name__ == "__main__":
    main() 