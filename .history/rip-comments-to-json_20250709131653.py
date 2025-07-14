from docx import Document
import json
from datetime import datetime

def extract_comments(docx_path):
    doc = Document(docx_path)
    comments = []
    for c in doc.part.comments_part.comments:
        # Assemble the full comment text (may be multiple paragraphs)
        full_text = "\n".join(p.text for p in c.paragraphs).strip()
        comments.append({
            "id":       c.id,
            "author":   c.author,
            "date":     c.date.isoformat() if isinstance(c.date, datetime) else str(c.date),
            "text":     full_text,
            # You’ll need to map this comment to your JSON structure
            # e.g. {"part":0,"subsection":1,"item":2}, based on where
            # the commentRangeStart is located in doc.paragraphs
            "ref":      None  
        })
    return comments

if __name__ == "__main__":
    comments = extract_comments("SECTION 26 05 00.docx")
    print(json.dumps({"comments": comments}, indent=2))
