from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_basic_numbering(paragraph, level=0):
    """Apply basic numbering to a paragraph using paragraph properties"""
    # Get paragraph properties
    pPr = paragraph._p.get_or_add_pPr()
    
    # Add numbering properties
    numPr = pPr.get_or_add_numPr()
    
    # Set numbering ID (this will use Word's default numbering)
    numId = numPr.get_or_add_numId()
    numId.val = 1
    
    # Set the level
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = level

# Open the template file
try:
    doc = Document('test_template.docx')
    print("Using existing template file")
except ValueError as e:
    print("Template file error, creating new document...")
    doc = Document()

# Add content with basic numbering
# Level 1 (Parts)
p1 = doc.add_paragraph('PART 1', style='Heading 1')
apply_basic_numbering(p1, 0)

# Level 2 (Subsections)
p2 = doc.add_paragraph('Test Subsection', style='Heading 2')
apply_basic_numbering(p2, 1)

# Level 3 (Items)
p3 = doc.add_paragraph('Test Item', style='Heading 3')
apply_basic_numbering(p3, 2)

# Level 1 again (auto resets Level 2)
p4 = doc.add_paragraph('PART 2', style='Heading 1')
apply_basic_numbering(p4, 0)

p5 = doc.add_paragraph('Another Subsection', style='Heading 2')
apply_basic_numbering(p5, 1)

# Add some regular content
doc.add_paragraph('This is regular paragraph text without numbering.')

doc.save('generated_spec.docx')
print("Document saved as 'generated_spec.docx' with basic numbering!")
print("Note: You may need to apply multilevel list formatting manually in Word")
print("or use a template that already has the multilevel list defined.")
