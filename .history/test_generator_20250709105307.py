from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clear_document(doc):
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)


def apply_basic_numbering(paragraph, level=0):
    """Apply basic numbering to a paragraph using paragraph properties"""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numId = numPr.get_or_add_numId()
    numId.val = 1
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = level

try:
    doc = Document('test_template.docx')
    print("Using existing template file")
except ValueError as e:
    print("Template file error, creating new document...")
    doc = Document()

# Clear all existing content from the document
clear_document(doc)

# Add content with basic numbering
p1 = doc.add_paragraph('PART 1', style='Heading 1')
apply_basic_numbering(p1, 0)

p2 = doc.add_paragraph('Test Subsection', style='Heading 2')
apply_basic_numbering(p2, 1)

p3 = doc.add_paragraph('Test Item', style='Heading 3')
apply_basic_numbering(p3, 2)

p4 = doc.add_paragraph('PART 2', style='Heading 1')
apply_basic_numbering(p4, 0)

p5 = doc.add_paragraph('Another Subsection', style='Heading 2')
apply_basic_numbering(p5, 1)

doc.add_paragraph('This is regular paragraph text without numbering.')

doc.save('generated_spec.docx')
print("Document saved as 'generated_spec.docx' with basic numbering!")
print("Note: You may need to apply multilevel list formatting manually in Word")
print("or use a template that already has the multilevel list defined.")
