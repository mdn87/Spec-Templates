from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def create_multilevel_list(doc):
    """Create and apply multilevel list numbering to the document"""
    
    # Define the multilevel list structure
    list_def = parse_xml(f'''
    <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNum w:abstractNumId="0">
            <w:nsid w:val="12345678"/>
            <w:multiLevelType w:val="hybridMultilevel"/>
            
            <!-- Level 1 (Parts) -->
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.0"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:b/>
                    <w:sz w:val="24"/>
                </w:rPr>
            </w:lvl>
            
            <!-- Level 2 (Subsections) -->
            <w:lvl w:ilvl="1">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimalZero"/>
                <w:lvlText w:val="%1.%2"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="360" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:b/>
                    <w:sz w:val="22"/>
                </w:rPr>
            </w:lvl>
            
            <!-- Level 3 (Items) -->
            <w:lvl w:ilvl="2">
                <w:start w:val="1"/>
                <w:numFmt w:val="upperLetter"/>
                <w:lvlText w:val="%3."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:sz w:val="20"/>
                </w:rPr>
            </w:lvl>
            
            <!-- Level 4 (Lists) -->
            <w:lvl w:ilvl="3">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%4."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="1080" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:sz w:val="18"/>
                </w:rPr>
            </w:lvl>
            
            <!-- Level 5 (Sub-items) -->
            <w:lvl w:ilvl="4">
                <w:start w:val="1"/>
                <w:numFmt w:val="lowerLetter"/>
                <w:lvlText w:val="%5."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="1440" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:sz w:val="18"/>
                </w:rPr>
            </w:lvl>
            
            <!-- Level 6 (Sub-sub-lists) -->
            <w:lvl w:ilvl="5">
                <w:start w:val="1"/>
                <w:numFmt w:val="lowerRoman"/>
                <w:lvlText w:val="%6."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="1800" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:sz w:val="18"/>
                </w:rPr>
            </w:lvl>
        </w:abstractNum>
        
        <w:num w:numId="1">
            <w:abstractNumId w:val="0"/>
        </w:num>
    </w:numbering>
    ''')
    
    # Add the numbering definition to the document
    # Check if numbering part already exists
    try:
        numbering_part = doc.part.get_or_add_part('word/numbering.xml', 'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml')
    except AttributeError:
        # If the method doesn't exist, try a different approach
        numbering_part = doc.part.get_or_add_part('word/numbering.xml')
    
    numbering_part._element = list_def
    
    return doc

def apply_numbering_to_paragraph(paragraph, level=0):
    """Apply multilevel numbering to a paragraph"""
    # Set the paragraph to use the multilevel list
    paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 1
    paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

# Open the template file
try:
    doc = Document('test_template.docx')
except ValueError as e:
    print("Template file error, creating new document...")
    doc = Document()

# Create the multilevel list structure
doc = create_multilevel_list(doc)

# Add content with automatic numbering
# Level 1 (Parts)
p1 = doc.add_paragraph('PART 1', style='Heading 1')
apply_numbering_to_paragraph(p1, 0)

# Level 2 (Subsections)
p2 = doc.add_paragraph('Test Subsection', style='Heading 2')
apply_numbering_to_paragraph(p2, 1)

# Level 3 (Items)
p3 = doc.add_paragraph('Test Item', style='Heading 3')
apply_numbering_to_paragraph(p3, 2)

# Level 4 (Lists)
p4 = doc.add_paragraph('Test List Item', style='List Paragraph')
apply_numbering_to_paragraph(p4, 3)

# Level 5 (Sub-items)
p5 = doc.add_paragraph('Test Sub-item', style='List Paragraph')
apply_numbering_to_paragraph(p5, 4)

# Level 6 (Sub-sub-lists)
p6 = doc.add_paragraph('Test Sub-sub-item', style='List Paragraph')
apply_numbering_to_paragraph(p6, 5)

# Level 1 again (auto resets Level 2)
p7 = doc.add_paragraph('PART 2', style='Heading 1')
apply_numbering_to_paragraph(p7, 0)

p8 = doc.add_paragraph('Another Subsection', style='Heading 2')
apply_numbering_to_paragraph(p8, 1)

doc.save('generated_spec.docx')
print("Document saved as 'generated_spec.docx' with automatic multilevel numbering!")
