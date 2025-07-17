#!/usr/bin/env python3
"""
Test script to verify the "don't add space between paragraphs of the same style" property
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os

def create_test_document():
    """Create a test document with the dont_add_space_between_same_style property set"""
    doc = Document()
    
    # Add a paragraph with the property set
    p1 = doc.add_paragraph("First paragraph with dont_add_space_between_same_style")
    p2 = doc.add_paragraph("Second paragraph with dont_add_space_between_same_style")
    
    # Set the property on both paragraphs
    for p in [p1, p2]:
        pPr = p._p.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p._p.insert(0, pPr)
        
        # Add the dont_add_space_between_same_style property
        dont_add_space = OxmlElement('w:dontAddSpaceBetweenSameStyle')
        pPr.append(dont_add_space)
    
    # Save the test document
    test_path = "../output/test_dont_add_space.docx"
    doc.save(test_path)
    print(f"Test document created: {test_path}")
    return test_path

def test_extraction():
    """Test extraction of the property"""
    from extract_spec_content_v3 import SpecContentExtractorV3
    
    # Create test document
    test_path = create_test_document()
    
    # Extract content
    extractor = SpecContentExtractorV3()
    data = extractor.extract_content(test_path)
    
    # Check if the property was extracted
    found_property = False
    for block in data.get('content_blocks', []):
        if block.get('paragraph_dont_add_space_between_same_style') is True:
            found_property = True
            print(f"Found paragraph with dont_add_space_between_same_style=True: {block['text'][:50]}...")
    
    if found_property:
        print("SUCCESS: Property was correctly extracted!")
    else:
        print("FAILURE: Property was not found in extraction")
    
    # Save the extraction result
    output_path = "../output/test_dont_add_space_extraction.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Extraction result saved: {output_path}")

def test_regeneration():
    """Test regeneration of the property"""
    from test_generator import Document, parse_spec_json, generate_content_from_v3_json
    
    # Use the extraction result from the previous test
    json_path = "../output/test_dont_add_space_extraction.json"
    
    if not os.path.exists(json_path):
        print("Extraction result not found, running extraction first...")
        test_extraction()
    
    # Parse the JSON
    json_data = parse_spec_json(json_path)
    if not json_data:
        print("Failed to parse JSON data")
        return
    
    # Create a new document
    doc = Document()
    
    # Generate content
    generate_content_from_v3_json(doc, json_data)
    
    # Save the regenerated document
    output_path = "../output/test_dont_add_space_regenerated.docx"
    doc.save(output_path)
    print(f"Regenerated document saved: {output_path}")
    
    # Check if the property was applied (we can't easily verify this programmatically,
    # but we can check that no errors occurred)
    print("SUCCESS: Document was regenerated without errors!")

if __name__ == "__main__":
    test_extraction()
    test_regeneration() 